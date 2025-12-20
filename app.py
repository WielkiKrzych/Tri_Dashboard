import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import numpy as np
from numba import jit
import io
from io import BytesIO
from scipy import stats
import neurokit2 as nk
from pathlib import Path

try:
    import sweatpy as sw
    SWEATP_AVAILABLE = True
except Exception:
    SWEATP_AVAILABLE = False
    MLX_AVAILABLE = False
try:
    import mlx.core as mx
    import mlx.nn as nn
    import mlx.optimizers as optim
    MLX_AVAILABLE = True
except ImportError:
    pass

# --- MLX MODULE (APPLE SILICON NEURAL NETWORK) ---
import os
import json
import time

MLX_AVAILABLE = False
try:
    import mlx.core as mx
    import mlx.nn as nn
    import mlx.optimizers as optim
    MLX_AVAILABLE = True
except ImportError:
    pass

MODEL_FILE = "cycling_brain_weights.npz"
HISTORY_FILE = "brain_evolution_history.json"

if MLX_AVAILABLE:
    class PhysioNet(nn.Module):
        def __init__(self):
            super().__init__()
            self.l1 = nn.Linear(3, 64)
            self.l2 = nn.Linear(64, 64)
            self.l3 = nn.Linear(64, 1)

        def __call__(self, x):
            x = nn.relu(self.l1(x))
            x = nn.relu(self.l2(x))
            return self.l3(x)

    def save_model(model, filepath):
        flattened_params = {}
        for k, v in model.parameters().items():
            if isinstance(v, dict):
                 for sub_k, sub_v in v.items():
                     flattened_params[f"{k}.{sub_v}"] = sub_v
            else:
                flattened_params[k] = v
        mx.savez(filepath, **dict(mx.tree_flatten(model.parameters())))

    def load_model(model, filepath):
        if os.path.exists(filepath):
            try:
                weights = mx.load(filepath)
                
                try:
                    model.update(weights)
                except Exception:
                    current_params = model.parameters()
                    new_params = {}
                    
                    for k, v in weights.items():
                        parts = k.split('.')
                        if len(parts) == 2:
                            layer, param = parts
                            if layer not in new_params: new_params[layer] = {}
                            new_params[layer][param] = v
                    
                    model.update(new_params)
                
                return True
            except Exception as e:
                st.sidebar.error(f"⚠️ Błąd AI: {e}")
                print(f"DEBUG ERROR: {e}")
                return False
        return False

    def update_history(hr_base, hr_thresh):
        """Zapisuje historię Baza/Próg do JSON z obsługą None"""
        history = []
        if os.path.exists(HISTORY_FILE):
            try:
                with open(HISTORY_FILE, 'r') as f:
                    history = json.load(f)
            except: pass
        
        entry = {
            "timestamp": time.time(),
            "date": time.strftime("%Y-%m-%d %H:%M"),
            "hr_base": float(hr_base) if hr_base is not None else None,
            "hr_thresh": float(hr_thresh) if hr_thresh is not None else None
        }
        history.append(entry)
        
        with open(HISTORY_FILE, 'w') as f:
            json.dump(history, f)
        return history

    def predict_only(df):
        """Tylko predykcja (bez treningu) - dla automatycznego wykresu"""
        if not os.path.exists(MODEL_FILE):
            return None
            
        w = df['watts_smooth'].values / 500.0
        c = df['cadence_smooth'].values / 120.0 if 'cadence_smooth' in df else np.zeros_like(w)
        t = df['time_min'].values / df['time_min'].max()
        
        X_np = np.column_stack((w, c, t)).astype(np.float32)
        X_np = np.nan_to_num(X_np, copy=False) 
        
        X = mx.array(X_np)
        
        model = PhysioNet()
        if load_model(model, MODEL_FILE):
            y_pred_scaled = model(X)
            return np.array(y_pred_scaled).flatten() * 200.0
        return None
    
    def filter_and_prepare(df, target_watts, tolerance=15, min_samples=30):
        mask = (df['watts_smooth'] >= target_watts - tolerance) & \
            (df['watts_smooth'] <= target_watts + tolerance)
        
        if mask.sum() < min_samples:
            return None, None

        df_filtered = df[mask].copy()
        w = df_filtered['watts_smooth'].values / 500.0
        c = df_filtered['cadence_smooth'].values / 120.0 if 'cadence_smooth' in df_filtered else np.zeros_like(w)
        t = df_filtered['time_min'].values / df['time_min'].max()
        y = df_filtered['heartrate_smooth'].values / 200.0

        X_np = np.column_stack((w, c, t)).astype(np.float32)
        X_np = np.nan_to_num(X_np, copy=False)
        
        y_np = y.astype(np.float32).reshape(-1, 1)
        y_np = np.nan_to_num(y_np, copy=False)

        X = mx.array(X_np)
        Y = mx.array(y_np)
        return X, Y

    def train_cycling_brain(df, epochs=200):
        model = PhysioNet()
        mx.eval(model.parameters())
        
        loaded = load_model(model, MODEL_FILE)
        
        def mse_loss(pred, target): return mx.mean((pred - target) ** 2)
        optimizer = optim.Adam(learning_rate=0.02)
        def train_step(model, X, y):
            loss = mse_loss(model(X), y)
            return loss
        loss_and_grad_fn = nn.value_and_grad(model, train_step)

        status_container = st.empty()
        bar = st.progress(0)
        
        results = {"base": None, "thresh": None}
        targets = [("base", 280), ("thresh", 360)]
        
        status_container.info("Trenowanie modelu ogólnego (cały plik)...")
        w_all = df['watts_smooth'].values / 500.0
        c_all = df['cadence_smooth'].values / 120.0 if 'cadence_smooth' in df else np.zeros_like(w_all)
        t_all = df['time_min'].values / df['time_min'].max()
        y_all = df['heartrate_smooth'].values / 200.0
        
        X_all_np = np.column_stack((w_all, c_all, t_all)).astype(np.float32)
        X_all_np = np.nan_to_num(X_all_np, copy=False)
        
        Y_all_np = y_all.astype(np.float32).reshape(-1, 1)
        Y_all_np = np.nan_to_num(Y_all_np, copy=False)

        X_all = mx.array(X_all_np)
        Y_all = mx.array(Y_all_np)
        
        for i in range(100): 
            loss, grads = loss_and_grad_fn(model, X_all, Y_all)
            optimizer.update(model, grads)
            mx.eval(model.parameters(), optimizer.state)
        
        y_pred_full = np.array(model(X_all)).flatten() * 200.0
        save_model(model, MODEL_FILE) 

        step = 0
        total_steps = len(targets) * epochs
        
        for name, watts in targets:
            status_container.info(f"Kalibracja strefy: {watts}W...")
            
            X_chunk, y_chunk = filter_and_prepare(df, watts)
            
            if X_chunk is not None:
                for i in range(epochs):
                    loss, grads = loss_and_grad_fn(model, X_chunk, y_chunk)
                    optimizer.update(model, grads)
                    mx.eval(model.parameters(), optimizer.state)
                    if i % 10 == 0: 
                        step += 10
                        bar.progress(min(step / total_steps, 1.0))
                
                in_vec = mx.array([[watts/500.0, 80.0/120.0, 0.5]]) 
                pred = float(model(in_vec)[0][0]) * 200.0
                results[name] = pred
            else:
                results[name] = None
                step += epochs
                
        bar.empty(); status_container.empty()

        history = update_history(results["base"], results["thresh"])

        return y_pred_full, results["base"], results["thresh"], loaded, history

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import zipfile
from io import BytesIO

from fpdf import FPDF
import base64

# ===== TRAINING NOTES SYSTEM =====
class TrainingNotes:
    """Zarządzanie notatkami do treningów"""
    
    NOTES_DIR = Path('training_notes')
    
    def __init__(self):
        self.NOTES_DIR.mkdir(exist_ok=True)
    
    def get_notes_file(self, training_file):
        """Pobierz plik notatek dla danego treningu"""
        base_name = Path(training_file).stem
        return self.NOTES_DIR / f"{base_name}_notes.json"
    
    def load_notes(self, training_file):
        """Załaduj notatki z JSON"""
        notes_file = self.get_notes_file(training_file)
        if notes_file.exists():
            with open(notes_file, 'r', encoding='utf-8') as f:
                return json.load(f)
        return {"training_file": training_file, "notes": []}
    
    def save_notes(self, training_file, notes_data):
        """Zapisz notatki do JSON"""
        notes_file = self.get_notes_file(training_file)
        with open(notes_file, 'w', encoding='utf-8') as f:
            json.dump(notes_data, f, indent=2, ensure_ascii=False)
    
    def add_note(self, training_file, time_minute, metric, text):
        """Dodaj notatkę"""
        notes_data = self.load_notes(training_file)
        
        note = {
            "time_minute": float(time_minute),
            "metric": metric,
            "text": text,
            "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        }
        
        notes_data["notes"].append(note)
        self.save_notes(training_file, notes_data)
        return note
    
    def get_notes_for_metric(self, training_file, metric):
        """Pobierz notatki dla konkretnej metryki"""
        notes_data = self.load_notes(training_file)
        return [n for n in notes_data["notes"] if n["metric"] == metric]
    
    def delete_note(self, training_file, note_index):
        """Usuń notatkę"""
        notes_data = self.load_notes(training_file)
        if 0 <= note_index < len(notes_data["notes"]):
            notes_data["notes"].pop(note_index)
            self.save_notes(training_file, notes_data)
            return True
        return False

# Inicjalizuj
training_notes = TrainingNotes()
# ===== KONIEC NOTES =====

# ===== FUNKCJA GENEROWANIA RAPORTU DOCX (ZMODYFIKOWANA) =====
def generate_docx_report(metrics, df_plot, df_plot_resampled, uploaded_file, cp_input, 
                        vt1_watts, vt2_watts, rider_weight, vt1_vent, vt2_vent, w_prime_input):
    """
    Generuje raport .docx z rozszerzonymi KPI, W', VO2max i czystymi notatkami.
    """
    
    doc = Document()
    
    # --- OBLICZENIA DODATKOWE DLA RAPORTU ---
    # 1. Normalized Power (NP)
    if 'watts' in df_plot.columns:
        rolling_30s = df_plot['watts'].rolling(window=30, min_periods=1).mean()
        np_val = np.power(np.mean(np.power(rolling_30s, 4)), 0.25)
    else:
        np_val = 0

    # 2. Praca (Total Work in kJ)
    total_work_kj = df_plot['watts'].sum() / 1000 if 'watts' in df_plot.columns else 0

    # 3. Pulse Power & EF (Efficiency Factor)
    avg_pp = 0
    if 'watts' in df_plot.columns and 'heartrate' in df_plot.columns:
        # Filtrujemy zera
        mask = (df_plot['watts'] > 10) & (df_plot['heartrate'] > 40)
        if mask.sum() > 0:
            avg_pp = (df_plot.loc[mask, 'watts'] / df_plot.loc[mask, 'heartrate']).mean()
    
    # 4. Temperatura i HSI
    max_core = df_plot['core_temperature'].max() if 'core_temperature' in df_plot.columns else 0
    avg_core = df_plot['core_temperature'].mean() if 'core_temperature' in df_plot.columns else 0
    max_hsi = df_plot['hsi'].max() if 'hsi' in df_plot.columns else 0

    # 5. RMSSD (Jeśli dostępne w kolumnach lub obliczone wcześniej, tutaj uproszczona ekstrakcja)
    avg_rmssd = 0
    if 'rmssd' in df_plot.columns:
        avg_rmssd = df_plot['rmssd'].mean()
    elif 'hrv' in df_plot.columns:
        avg_rmssd = df_plot['hrv'].mean() # Często HRV w plikach to RMSSD

    # 6. Spalone Węglowodany (Logika z tab_nutrition)
    carbs_total = 0
    if 'watts' in df_plot.columns:
        # Zakładamy efficiency 22%
        energy_kcal_sec = (df_plot['watts'] / 0.22) / 4184.0
        
        # Frakcje węgli wg stref (uproszczone dla raportu)
        conditions = [
            (df_plot['watts'] < vt1_watts),
            (df_plot['watts'] >= vt1_watts) & (df_plot['watts'] < vt2_watts),
            (df_plot['watts'] >= vt2_watts)
        ]
        choices = [0.3, 0.8, 1.1] 
        carb_fraction = np.select(conditions, choices, default=1.0)
        
        carbs_burned_sec = (energy_kcal_sec * carb_fraction) / 4.0
        carbs_total = carbs_burned_sec.sum()

    # 7. VO2max (Estymacja z 5 min MMP)
    vo2_max_est = 0
    if 'watts' in df_plot.columns:
        mmp_5m = df_plot['watts'].rolling(300).mean().max()
        if not pd.isna(mmp_5m):
            vo2_max_est = (10.8 * mmp_5m / rider_weight) + 7

    # --- KONIEC OBLICZEŃ ---

    # STYLE
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    
    # NAGŁÓWEK
    title = doc.add_heading('Pro Athlete Dashboard - Raport Treningowy', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    date_para = doc.add_paragraph(f"Data: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.runs[0].font.color.rgb = RGBColor(100, 100, 100)
    
    source_para = doc.add_paragraph(f"Plik: {uploaded_file.name if uploaded_file else 'Brak'}")
    source_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph() 
    
    # SEKCJA 1: PODSUMOWANIE KPI (ROZSZERZONE)
    doc.add_heading('1. Podsumowanie KPI', level=1)
    
    # Tabela z dużą ilością danych (18 wierszy)
    kpi_data = [
        ("Średnia Moc", f"{metrics.get('avg_watts', 0):.0f} W"),
        ("Normalized Power (NP)", f"{np_val:.0f} W"),
        ("Praca Całkowita", f"{total_work_kj:.0f} kJ"),
        ("Średnie Tętno", f"{metrics.get('avg_hr', 0):.0f} bpm"),
        ("Średnia Kadencja", f"{metrics.get('avg_cadence', 0):.0f} rpm"),
        ("Średnia Wentylacja (VE)", f"{metrics.get('avg_vent', 0):.1f} L/min"),
        ("Średnie Oddechy (RR)", f"{metrics.get('avg_rr', 0):.1f} /min"),
        ("Średnie SmO2", f"{df_plot['smo2'].mean() if 'smo2' in df_plot.columns else 0:.1f}%"),
        ("Min SmO2", f"{df_plot['smo2'].min() if 'smo2' in df_plot.columns else 0:.1f}%"),
        ("Max SmO2", f"{df_plot['smo2'].max() if 'smo2' in df_plot.columns else 0:.1f}%"),
        ("Power/HR", f"{metrics.get('power_hr', 0):.2f}"),
        ("Efficiency Factor (EF)", f"{metrics.get('ef_factor', 0):.2f}"),
        ("Średnie Pulse Power", f"{avg_pp:.2f} W/bpm"),
        ("Średnie RMSSD", f"{avg_rmssd:.1f} ms"),
        ("Max HSI (Indeks Ciepła)", f"{max_hsi:.1f}"),
        ("Max Temperatura Ciała", f"{max_core:.2f} °C"),
        ("Średnia Temperatura Ciała", f"{avg_core:.2f} °C"),
        ("Spalone Węglowodany", f"{carbs_total:.0f} g")
    ]

    table = doc.add_table(rows=len(kpi_data)+1, cols=2)
    table.style = 'Light Grid Accent 1'
    
    # Nagłówki tabeli
    table.rows[0].cells[0].text = 'Metryka'
    table.rows[0].cells[1].text = 'Wartość'
    
    # Wypełnianie danymi
    for i, (label, val) in enumerate(kpi_data):
        row_cells = table.rows[i+1].cells
        row_cells[0].text = label
        row_cells[1].text = val
    
    doc.add_paragraph()
    
    # SEKCJA 2: PROGI TRENINGOWE (DODANO W' i VO2max)
    doc.add_heading('2. Progi i Parametry Fizjologiczne', level=1)
    
    p = doc.add_paragraph()
    p.add_run(f"VT1 (Próg Tlenowy): ").bold = True
    p.add_run(f"{vt1_watts} W @ {vt1_vent} L/min\n")
    
    p.add_run(f"VT2 (Próg Beztlenowy): ").bold = True
    p.add_run(f"{vt2_watts} W @ {vt2_vent} L/min\n")
    
    p.add_run(f"CP (Moc Krytyczna): ").bold = True
    p.add_run(f"{cp_input} W\n")
    
    p.add_run(f"W' (Pojemność Beztlenowa): ").bold = True
    p.add_run(f"{w_prime_input} J\n")
    
    p.add_run(f"Szacunkowe VO2max (MMP 5'): ").bold = True
    p.add_run(f"{vo2_max_est:.1f} ml/kg/min\n")
    
    p.add_run(f"Waga zawodnika: ").bold = True
    p.add_run(f"{rider_weight} kg")
    
    doc.add_paragraph()
    
    # SEKCJA 3: STREFY MOCY (POPRAWKA Z6)
    doc.add_heading('3. Czas w Strefach Mocy', level=1)
    
    if 'watts' in df_plot.columns:
        # Definicja stref
        bins = [0, 0.55*cp_input, 0.75*cp_input, 0.90*cp_input, 1.05*cp_input, 1.20*cp_input, 10000]
        labels = ['Z1 Recovery', 'Z2 Endurance', 'Z3 Tempo', 'Z4 Threshold', 'Z5 VO2Max', 'Z6 Anaerobic']
        
        # Obliczenia
        dfz = df_plot.copy()
        dfz['Zone'] = pd.cut(dfz['watts'], bins=bins, labels=labels, right=False)
        
        table2 = doc.add_table(rows=7, cols=3)
        table2.style = 'Light Grid Accent 1'
        
        table2.rows[0].cells[0].text = 'Strefa'
        table2.rows[0].cells[1].text = 'Zakres'
        table2.rows[0].cells[2].text = 'Czas'
        
        for i, (label, low, high) in enumerate(zip(labels, bins[:-1], bins[1:])):
            count = len(dfz[dfz['Zone'] == label])
            time_min = count / 60
            
            # POPRAWKA: Wyświetlanie 2000 W zamiast 10000 W dla estetyki
            display_high = 2000 if high == 10000 else int(high)
            display_low = int(low)
            
            table2.rows[i+1].cells[0].text = label
            table2.rows[i+1].cells[1].text = f"{display_low}-{display_high} W"
            table2.rows[i+1].cells[2].text = f"{time_min:.1f} min"
    
    doc.add_paragraph()
    
    # SEKCJA 4: NOTATKI (WYCZYSZCZONE)
    doc.add_heading('4. Notatki Trenera / Zawodnika', level=1)
    
    # Pusty paragraf jako przestrzeń edytowalna w Pages/Word
    note_p = doc.add_paragraph("[Miejsce na Twoje notatki...]")
    note_p.runs[0].font.italic = True
    note_p.runs[0].font.color.rgb = RGBColor(150, 150, 150)
    
    # Dodajemy kilka pustych linii, żeby wizualnie zrobić miejsce, ale bez "___"
    for _ in range(5):
        doc.add_paragraph("")
    
    # STOPKA
    doc.add_paragraph("---")
    footer = doc.add_paragraph("Raport wygenerowany przez Pro Athlete Dashboard | Streamlit App")
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    return doc
# ===== KONIEC FUNKCJI DOCX =====

# ===== PNG BATCH EXPORT (FIXED VARIABLE SHADOWING) =====
def export_all_charts_as_png(df_plot, df_plot_resampled, cp_input, vt1_watts, vt2watts,
                            metrics, rider_weight, uploaded_file):
    """
    Export wykresów PNG z pełną legendą statystyczną (Ghost Traces).
    Poprawiono błąd 'stats' variable shadowing.
    """
    
    zip_buffer = BytesIO()
    
    # Konfiguracja wizualna
    layout_args = dict(
        template='plotly_dark',
        height=600,
        width=1200,
        font=dict(family="Inter", size=14),
        margin=dict(l=50, r=50, t=80, b=50),
        legend=dict(font=dict(size=12))
    )

    # Helper do dodawania statystyk do legendy
    def add_stats_to_legend(fig, stats_list):
        for stat in stats_list:
            fig.add_trace(go.Scatter(
                x=[None], y=[None], mode='markers',
                marker=dict(color='rgba(0,0,0,0)'),
                name=stat, hoverinfo='none'
            ))

    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zipf:
        
        # --- 1. POWER ---
        if 'watts_smooth' in df_plot_resampled.columns:
            fig = go.Figure()
            fig.add_trace(go.Scatter(x=df_plot_resampled['time_min'], y=df_plot_resampled['watts_smooth'],
                                   name='Power', fill='tozeroy', line=dict(color='#00cc96', width=1.5)))
            
            # Statystyki
            avg_p = df_plot_resampled['watts_smooth'].mean()
            max_p = df_plot_resampled['watts_smooth'].max()
            norm_p = np.power(np.mean(np.power(df_plot_resampled['watts_smooth'], 4)), 0.25)
            
            legend_stats = [
                f"⚡ Avg: {avg_p:.0f} W",
                f"🔥 Max: {max_p:.0f} W",
                f"📈 NP (est): {norm_p:.0f} W",
                f"⚖️ W/kg: {avg_p/rider_weight:.2f}"
            ]
            add_stats_to_legend(fig, legend_stats)
            
            fig.update_layout(title='1. Power Profile (W)', xaxis_title='Time (min)', yaxis_title='Power (W)', **layout_args)
            zipf.writestr('01_Power.png', fig.to_image(format='png', width=1200, height=600))
        
        # --- 2. HR ---
        if 'heartrate_smooth' in df_plot_resampled.columns:
            fig = go.Figure()
            fig.add_trace(go.Scatter(x=df_plot_resampled['time_min'], y=df_plot_resampled['heartrate_smooth'],
                                   name='HR', fill='tozeroy', line=dict(color='#ef553b', width=1.5)))
            
            # Statystyki
            avg_hr = df_plot_resampled['heartrate_smooth'].mean()
            max_hr = df_plot_resampled['heartrate_smooth'].max()
            min_hr = df_plot_resampled[df_plot_resampled['heartrate_smooth'] > 40]['heartrate_smooth'].min()
            
            legend_stats = [
                f"❤️ Avg: {avg_hr:.0f} bpm",
                f"🔥 Max: {max_hr:.0f} bpm",
                f"💤 Min: {min_hr:.0f} bpm"
            ]
            add_stats_to_legend(fig, legend_stats)

            fig.update_layout(title='2. Heart Rate (bpm)', xaxis_title='Time (min)', yaxis_title='HR (bpm)', **layout_args)
            zipf.writestr('02_HeartRate.png', fig.to_image(format='png', width=1200, height=600))

        # --- 3. SmO2 ---
        if 'smo2_smooth' in df_plot_resampled.columns:
            fig = go.Figure()
            fig.add_trace(go.Scatter(x=df_plot_resampled['time_min'], y=df_plot_resampled['smo2_smooth'],
                                   name='SmO2', line=dict(color='#ab63fa', width=2)))
            
            # Statystyki
            avg_smo2 = df_plot_resampled['smo2_smooth'].mean()
            min_smo2 = df_plot_resampled['smo2_smooth'].min()
            max_smo2 = df_plot_resampled['smo2_smooth'].max()
            
            legend_stats = [
                f"📊 Avg: {avg_smo2:.1f}%",
                f"🔻 Min: {min_smo2:.1f}%",
                f"🔺 Max: {max_smo2:.1f}%"
            ]
            add_stats_to_legend(fig, legend_stats)

            fig.update_layout(title='3. Muscle Oxygenation (SmO2)', xaxis_title='Time (min)', yaxis_title='SmO2 (%)', 
                            yaxis=dict(range=[0, 100]), **layout_args)
            zipf.writestr('03_SmO2.png', fig.to_image(format='png', width=1200, height=600))

        # --- 4. VE + RR (Dual Axis) ---
        if 'tymeventilation_smooth' in df_plot_resampled.columns:
            fig = go.Figure()
            # VE
            fig.add_trace(go.Scatter(x=df_plot_resampled['time_min'], y=df_plot_resampled['tymeventilation_smooth'],
                                   name='VE', line=dict(color='#ffa15a', width=2)))
            
            legend_stats = []
            avg_ve = df_plot_resampled['tymeventilation_smooth'].mean()
            max_ve = df_plot_resampled['tymeventilation_smooth'].max()
            legend_stats.append(f"🫁 Avg VE: {avg_ve:.1f} L/min")
            legend_stats.append(f"🔥 Max VE: {max_ve:.1f} L/min")

            # RR (Prawa oś)
            if 'tymebreathrate_smooth' in df_plot_resampled.columns:
                fig.add_trace(go.Scatter(x=df_plot_resampled['time_min'], y=df_plot_resampled['tymebreathrate_smooth'],
                                       name='RR', line=dict(color='#19d3f3', width=2, dash='dot'), yaxis='y2'))
                avg_rr = df_plot_resampled['tymebreathrate_smooth'].mean()
                legend_stats.append(f"💨 Avg RR: {avg_rr:.1f} /min")
            
            add_stats_to_legend(fig, legend_stats)

            fig.update_layout(title='4. Ventilation (VE) & Respiratory Rate (RR)', 
                            xaxis_title='Time (min)', yaxis=dict(title='VE (L/min)'),
                            yaxis2=dict(title='RR (bpm)', overlaying='y', side='right'), **layout_args)
            zipf.writestr('04_Ventilation_RR.png', fig.to_image(format='png', width=1200, height=600))

        # --- 5. PULSE POWER ---
        if 'watts_smooth' in df_plot_resampled.columns and 'heartrate_smooth' in df_plot_resampled.columns:
            mask = (df_plot_resampled['watts_smooth'] > 50) & (df_plot_resampled['heartrate_smooth'] > 90)
            df_pp = df_plot_resampled[mask].copy()
            if not df_pp.empty:
                df_pp['pp'] = df_pp['watts_smooth'] / df_pp['heartrate_smooth']
                df_pp['pp_smooth'] = df_pp['pp'].rolling(window=30, center=True).mean()
                
                fig = go.Figure()
                fig.add_trace(go.Scatter(x=df_pp['time_min'], y=df_pp['pp_smooth'],
                                       name='Pulse Power', line=dict(color='#FFD700', width=2)))
                
                # Tu był błąd! Teraz używamy 'stats' biblioteki (bo lista nazywa się 'legend_stats')
                slope, intercept, _, _, _ = stats.linregress(df_pp['time_min'], df_pp['pp'])
                trend = intercept + slope * df_pp['time_min']
                fig.add_trace(go.Scatter(x=df_pp['time_min'], y=trend, name='Trend', line=dict(color='white', dash='dash')))

                # Statystyki
                avg_eff = df_pp['pp'].mean()
                total_drift = slope * (df_pp['time_min'].iloc[-1] - df_pp['time_min'].iloc[0])
                drift_pct = (total_drift / intercept) * 100 if intercept != 0 else 0
                
                legend_stats = [
                    f"🔋 Avg EF: {avg_eff:.2f} W/bpm",
                    f"📉 Drift: {drift_pct:.1f}%"
                ]
                add_stats_to_legend(fig, legend_stats)

                fig.update_layout(title='5. Pulse Power (Watts / Heart Beat)', xaxis_title='Time (min)', 
                                yaxis_title='Efficiency (W/bpm)', **layout_args)
                zipf.writestr('05_PulsePower.png', fig.to_image(format='png', width=1200, height=600))

        # --- 6. HRV TIME (Alpha-1) ---
        df_dfa, _ = calculate_dynamic_dfa(df_plot, window_sec=120, step_sec=30)
        if df_dfa is not None and not df_dfa.empty:
            df_dfa['time_min'] = df_dfa['time'] / 60.0
            fig = go.Figure()
            fig.add_trace(go.Scatter(x=df_dfa['time_min'], y=df_dfa['alpha1'],
                                   name='DFA Alpha-1', line=dict(color='#00cc96', width=2)))
            
            fig.add_hline(y=0.75, line_dash="dash", line_color="red", annotation_text="VT1 (0.75)")
            
            # Statystyki
            avg_a1 = df_dfa['alpha1'].mean()
            min_a1 = df_dfa['alpha1'].min()
            rmssd_avg = df_dfa['rmssd'].mean() if 'rmssd' in df_dfa else 0
            
            legend_stats = [
                f"🧠 Avg Alpha-1: {avg_a1:.2f}",
                f"⚠️ Min Alpha-1: {min_a1:.2f}",
                f"💓 Avg RMSSD: {rmssd_avg:.0f} ms"
            ]
            add_stats_to_legend(fig, legend_stats)

            fig.update_layout(title='6. HRV Variability (DFA Alpha-1)', xaxis_title='Time (min)', 
                            yaxis=dict(title='Alpha-1', range=[0.2, 1.6]), **layout_args)
            zipf.writestr('06_HRV_Time.png', fig.to_image(format='png', width=1200, height=600))
        
        # --- 7. POINCARE PLOT ---
        rr_col = next((c for c in df_plot.columns if any(x in c.lower() for x in ['rr', 'hrv', 'ibi', 'r-r'])), None)
        if rr_col:
            rr_vals = df_plot[rr_col].dropna().values
            if rr_vals.mean() < 2.0: rr_vals *= 1000 
            rr_vals = rr_vals[(rr_vals > 300) & (rr_vals < 2000)]
            
            if len(rr_vals) > 100:
                rr_n = rr_vals[:-1]
                rr_n1 = rr_vals[1:]
                
                # Obliczenia SD1/SD2
                diff_rr = np.diff(rr_vals)
                sd1 = np.std(diff_rr) / np.sqrt(2)
                sd2 = np.sqrt(2 * np.std(rr_vals)**2 - 0.5 * np.std(diff_rr)**2)
                ratio = sd2 / sd1 if sd1 > 0 else 0

                fig = go.Figure()
                fig.add_trace(go.Scatter(x=rr_n, y=rr_n1, mode='markers', 
                                       marker=dict(size=3, color='#00cc96', opacity=0.5), name='RR Intervals'))
                fig.add_trace(go.Scatter(x=[min(rr_vals), max(rr_vals)], y=[min(rr_vals), max(rr_vals)],
                                       mode='lines', line=dict(color='white', dash='dash'), name='Identity Line'))
                
                legend_stats = [
                    f"🟢 SD1 (Fast): {sd1:.1f} ms",
                    f"🔵 SD2 (Slow): {sd2:.1f} ms",
                    f"⚖️ Ratio: {ratio:.2f}"
                ]
                add_stats_to_legend(fig, legend_stats)

                fig.update_layout(title='7. Poincaré Plot (RR Intervals)', xaxis_title='RR(n) [ms]', 
                                yaxis_title='RR(n+1) [ms]', width=800, height=800, template='plotly_dark')
                zipf.writestr('07_Poincare.png', fig.to_image(format='png', width=800, height=800))

        # --- 8. TORQUE vs SmO2 ---
        if 'torque' in df_plot.columns and 'smo2' in df_plot.columns:
            df_bins = df_plot.copy()
            df_bins['Torque_Bin'] = (df_bins['torque'] // 2 * 2).astype(int)
            bin_stats = df_bins.groupby('Torque_Bin')['smo2'].agg(['mean', 'std', 'count']).reset_index()
            bin_stats = bin_stats[bin_stats['count'] > 10]
            
            if not bin_stats.empty:
                fig = go.Figure()
                fig.add_trace(go.Scatter(x=bin_stats['Torque_Bin'], y=bin_stats['mean']+bin_stats['std'],
                                       mode='lines', line=dict(width=0), showlegend=False))
                fig.add_trace(go.Scatter(x=bin_stats['Torque_Bin'], y=bin_stats['mean']-bin_stats['std'],
                                       mode='lines', line=dict(width=0), fill='tonexty', fillcolor='rgba(255,75,75,0.2)', showlegend=False))
                fig.add_trace(go.Scatter(x=bin_stats['Torque_Bin'], y=bin_stats['mean'],
                                       mode='lines+markers', name='Mean SmO2', line=dict(color='#FF4B4B', width=3)))
                
                # Statystyki
                max_t_idx = bin_stats['Torque_Bin'].idxmax()
                max_t = bin_stats.loc[max_t_idx, 'Torque_Bin']
                smo2_at_max = bin_stats.loc[max_t_idx, 'mean']
                
                legend_stats = [
                    f"💪 Max Torque: {max_t:.0f} Nm",
                    f"🩸 SmO2 @ Max: {smo2_at_max:.1f}%"
                ]
                add_stats_to_legend(fig, legend_stats)

                fig.update_layout(title='8. Mechanical Impact: Torque vs SmO2', xaxis_title='Torque (Nm)', 
                                yaxis_title='SmO2 (%)', **layout_args)
                zipf.writestr('08_Torque_SmO2.png', fig.to_image(format='png', width=1200, height=600))

        # --- 9. SmO2 ANALYSIS (FULL + SELECTION + STATS) ---
        s_sec = st.session_state.get('smo2_start_sec')
        e_sec = st.session_state.get('smo2_end_sec')
        
        if 'smo2_smooth' in df_plot_resampled.columns:
            fig = go.Figure()
            fig.add_trace(go.Scatter(x=df_plot_resampled['time_min'], y=df_plot_resampled['smo2_smooth'],
                name='SmO2 (Full)', line=dict(color='#FF4B4B', width=1.5)))
            
            if 'watts' in df_plot_resampled.columns:
                 fig.add_trace(go.Scatter(x=df_plot_resampled['time_min'], y=df_plot_resampled['watts_smooth'],
                     name='Power', line=dict(color='#1f77b4', width=1), opacity=0.3, yaxis='y2'))

            if s_sec is not None and e_sec is not None:
                s_min = s_sec / 60.0
                e_min = e_sec / 60.0
                fig.add_vrect(x0=s_min, x1=e_min, fillcolor="green", opacity=0.15, layer="below", line_width=0, annotation_text="ANALYSIS")
                
                mask = (df_plot['time'] >= s_sec) & (df_plot['time'] <= e_sec)
                df_sel = df_plot.loc[mask]
                
                if not df_sel.empty:
                    duration = e_sec - s_sec
                    avg_w_sel = df_sel['watts_smooth'].mean() if 'watts_smooth' in df_sel else 0
                    avg_s_sel = df_sel['smo2_smooth'].mean()
                    min_s_sel = df_sel['smo2_smooth'].min()
                    max_s_sel = df_sel['smo2_smooth'].max()
                    
                    # Użycie biblioteki stats (działa bo nazwa listy to legend_stats)
                    slope, intercept, _, _, _ = stats.linregress(df_sel['time'], df_sel['smo2_smooth'])
                    
                    x_trend_min = df_sel['time'] / 60.0
                    y_trend = intercept + slope * df_sel['time']
                    fig.add_trace(go.Scatter(x=x_trend_min, y=y_trend, name='Trend', line=dict(color='yellow', dash='solid', width=3)))

                    m_dur, s_dur = divmod(int(duration), 60)
                    legend_stats = [
                        f"⏱️ Time: {m_dur:02d}:{s_dur:02d}",
                        f"⚡ Avg W: {avg_w_sel:.0f} W",
                        f"📉 Slope: {slope:.4f} %/s",
                        f"📊 Avg SmO2: {avg_s_sel:.1f}%",
                        f"🔻 Min: {min_s_sel:.1f}%",
                        f"🔺 Max: {max_s_sel:.1f}%"
                    ]
                    add_stats_to_legend(fig, legend_stats)

            fig.update_layout(title='9. SmO2 Kinetics Analysis', xaxis_title='Time (min)', yaxis=dict(title='SmO2 (%)'),
                yaxis2=dict(title='Power (W)', overlaying='y', side='right', showgrid=False), **layout_args)
            zipf.writestr('09_SmO2_Analysis.png', fig.to_image(format='png', width=1200, height=600))

        # --- 10. VENT ANALYSIS (FULL + SELECTION + STATS) ---
        s_v_sec = st.session_state.get('vent_start_sec')
        e_v_sec = st.session_state.get('vent_end_sec')
        
        if 'tymeventilation_smooth' in df_plot_resampled.columns:
            fig = go.Figure()
            fig.add_trace(go.Scatter(x=df_plot_resampled['time_min'], y=df_plot_resampled['tymeventilation_smooth'],
                name='VE (Full)', line=dict(color='#ffa15a', width=1.5)))
            
            if 'watts' in df_plot_resampled.columns:
                 fig.add_trace(go.Scatter(x=df_plot_resampled['time_min'], y=df_plot_resampled['watts_smooth'],
                     name='Power', line=dict(color='#1f77b4', width=1), opacity=0.3, yaxis='y2'))
            
            if s_v_sec is not None and e_v_sec is not None:
                s_v_min = s_v_sec / 60.0
                e_v_min = e_v_sec / 60.0
                fig.add_vrect(x0=s_v_min, x1=e_v_min, fillcolor="orange", opacity=0.15, layer="below", line_width=0, annotation_text="ANALYSIS")
                
                mask_v = (df_plot['time'] >= s_v_sec) & (df_plot['time'] <= e_v_sec)
                df_v = df_plot.loc[mask_v]
                
                if not df_v.empty:
                    duration_v = e_v_sec - s_v_sec
                    avg_w_v = df_v['watts_smooth'].mean() if 'watts_smooth' in df_v else 0
                    avg_ve = df_v['tymeventilation_smooth'].mean()
                    min_ve = df_v['tymeventilation_smooth'].min()
                    max_ve = df_v['tymeventilation_smooth'].max()
                    
                    slope_v, intercept_v, _, _, _ = stats.linregress(df_v['time'], df_v['tymeventilation_smooth'])
                    
                    x_trend_v_min = df_v['time'] / 60.0
                    y_trend_v = intercept_v + slope_v * df_v['time']
                    fig.add_trace(go.Scatter(x=x_trend_v_min, y=y_trend_v, name='Trend', line=dict(color='white', dash='solid', width=3)))

                    m_dur_v, s_dur_v = divmod(int(duration_v), 60)
                    legend_stats = [
                        f"⏱️ Time: {m_dur_v:02d}:{s_dur_v:02d}",
                        f"⚡ Avg W: {avg_w_v:.0f} W",
                        f"📈 Slope: {slope_v:.4f} L/s",
                        f"🫁 Avg VE: {avg_ve:.1f} L/min",
                        f"🔻 Min: {min_ve:.1f}",
                        f"🔺 Max: {max_ve:.1f}"
                    ]
                    add_stats_to_legend(fig, legend_stats)

            fig.update_layout(title='10. Ventilation Threshold Analysis', xaxis_title='Time (min)', yaxis=dict(title='VE (L/min)'),
                yaxis2=dict(title='Power (W)', overlaying='y', side='right', showgrid=False), **layout_args)
            zipf.writestr('10_Vent_Analysis.png', fig.to_image(format='png', width=1200, height=600))

        # README
        readme = f"""RAPORT WYKRESÓW Z PEŁNĄ ANALIZĄ
Data: {datetime.now().strftime('%Y-%m-%d %H:%M')}
Plik: {uploaded_file.name}

Wszystkie wykresy zawierają statystyki w legendzie (Avg, Min, Max, Trends).
Wykresy 9 i 10 zawierają analizę odcinków wybranych w aplikacji.
"""
        zipf.writestr('00_README.txt', readme)
    
    zip_buffer.seek(0)
    return zip_buffer.getvalue()
# ===== KONIEC PNG EXPORT =====

class PDFReport(FPDF):
    def header(self):
        self.set_font('Arial', 'B', 15)
        self.cell(0, 10, 'Pro Athlete Dashboard - Raport Treningowy', 0, 1, 'C')
        self.ln(10)

    def footer(self):
        self.set_y(-15)
        self.set_font('Arial', 'I', 8)
        self.cell(0, 10, f'Strona {self.page_no()}', 0, 0, 'C')

def create_download_link(val, filename):
    b64 = base64.b64encode(val)  # val looks like b'...'
    return f'<a href="data:application/octet-stream;base64,{b64.decode()}" download="{filename}.pdf">📥 Pobierz Raport PDF</a>'

st.set_page_config(page_title="Pro Athlete Dashboard", layout="wide", page_icon="⚡")

st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;600&family=Rajdhani:wght@500;600;700&display=swap');

    /* GLOBALNE TŁO */
    .stApp {
        background: radial-gradient(circle at 10% 20%, #1a1f25 0%, #0e1117 90%);
        color: #e0e0e0;
        font-family: 'Inter', sans-serif;
    }

    /* TYPOGRAFIA */
    h1, h2, h3, h4, h5 {
        font-family: 'Rajdhani', sans-serif !important;
        text-transform: uppercase;
        letter-spacing: 1px;
        color: #ffffff !important;
        text-shadow: 0 2px 4px rgba(0,0,0,0.5);
    }
    h1 { font-weight: 700; color: #00cc96 !important; }

    /* SIDEBAR */
    [data-testid="stSidebar"] {
        background-color: #050505;
        border-right: 1px solid #30363d;
    }

    /* METRYKI */
    [data-testid="stMetric"] {
        background: rgba(255, 255, 255, 0.03);
        backdrop-filter: blur(10px);
        -webkit-backdrop-filter: blur(10px);
        border-radius: 12px;
        border: 1px solid rgba(255, 255, 255, 0.1);
        padding: 15px !important;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.3);
        transition: all 0.3s ease;
    }
    [data-testid="stMetric"]:hover {
        transform: translateY(-3px);
        border-color: #00cc96;
        box-shadow: 0 8px 15px rgba(0, 204, 150, 0.2);
    }
    
    /* Kolor etykiety metryki */
    [data-testid="stMetricLabel"] {
        font-size: 0.9rem !important;
        color: #8b949e !important;
        font-family: 'Rajdhani', sans-serif;
    }
    /* Kolor wartości metryki */
    [data-testid="stMetricValue"] {
        font-family: 'Rajdhani', sans-serif;
        font-weight: 600;
        color: #f0f6fc !important;
    }

    /* ZAKŁADKI */
    .stTabs [data-baseweb="tab-list"] {
        gap: 8px;
        background-color: transparent;
        padding-bottom: 10px;
    }
    .stTabs [data-baseweb="tab"] {
        height: 40px;
        border-radius: 8px;
        background-color: rgba(255,255,255,0.05);
        color: #c9d1d9;
        border: none;
        font-family: 'Rajdhani', sans-serif;
        font-weight: 600;
        padding: 0 20px;
    }
    .stTabs [aria-selected="true"] {
        background-color: #00cc96 !important;
        color: #000000 !important;
    }

    /* INFO BOXY */
    .stAlert {
        background-color: rgba(22, 27, 34, 0.9);
        border: 1px solid #30363d;
        border-left: 5px solid #58a6ff;
        border-radius: 8px;
        color: #c9d1d9;
    }
    </style>
""", unsafe_allow_html=True)

# --- 2. KONFIGURACJA STAŁYCH ---
class Config:
    SMOOTH_WINDOW = 30
    SMOOTH_WINDOW_SHORT = 5
    COLOR_POWER = '#00cc96'
    COLOR_HR = '#ef553b'
    COLOR_SMO2 = '#ab63fa'
    COLOR_VE = '#ffa15a'
    COLOR_RR = '#19d3f3'
    COLOR_THB = '#e377c2'
    COLOR_TORQUE = '#e377c2'

# --- 3. FUNKCJE POMOCNICZE ---

def parse_time_input(t_str):
    try:
        parts = list(map(int, t_str.split(':')))
        if len(parts) == 3: return parts[0]*3600 + parts[1]*60 + parts[2]
        if len(parts) == 2: return parts[0]*60 + parts[1]
        if len(parts) == 1: return parts[0]
    except: return None
    return None

def _serialize_df_to_parquet_bytes(df):
    bio = io.BytesIO()
    try:
        df.to_parquet(bio, index=False)
        return bio.getvalue()
    except Exception:
        bio = io.BytesIO()
        df.to_csv(bio, index=False)
        return bio.getvalue()

# --- NOWY SILNIK NUMBA (WKLEJ TO) ---
@jit(nopython=True)
def calculate_w_prime_fast(watts, time, cp, w_prime_cap):
    n = len(watts)
    w_bal = np.empty(n, dtype=np.float64)
    curr_w = w_prime_cap
    
    # Obliczamy różnice czasu (dt) wewnątrz Numby dla szybkości
    # Dla pierwszego punktu zakładamy 1 sekundę, dla reszty różnicę
    dt = np.empty(n, dtype=np.float64)
    dt[0] = 1.0 
    for i in range(1, n):
        val = time[i] - time[i-1]
        # Zabezpieczenie przed zerowym czasem lub ujemnym (błędy w pliku)
        if val <= 0:
            dt[i] = 1.0
        else:
            dt[i] = val

    for i in range(n):
        # Logika: CP - Moc = delta. 
        # Jeśli delta > 0 (jedziesz lekko) -> regeneracja.
        # Jeśli delta < 0 (jedziesz mocno) -> spalanie.
        delta = (cp - watts[i]) * dt[i]
        curr_w += delta
        
        # Nie możemy mieć więcej niż 100% baterii
        if curr_w > w_prime_cap:
            curr_w = w_prime_cap
        # Nie możemy mieć mniej niż 0% baterii
        elif curr_w < 0:
            curr_w = 0.0
            
        w_bal[i] = curr_w
        
    return w_bal
# --- KONIEC NOWEGO SILNIKA ---

@st.cache_data
def _calculate_w_prime_balance_cached(df_bytes: bytes, cp: float, w_prime: float):
    try:
        # 1. Wczytanie danych z bajtów (tak jak było)
        bio = io.BytesIO(df_bytes)
        try:
            df_pd = pd.read_parquet(bio)
        except Exception:
            bio.seek(0)
            df_pd = pd.read_csv(bio)

        if 'watts' not in df_pd.columns:
            df_pd['w_prime_balance'] = np.nan
            return df_pd

        # 2. Przygotowanie tablic dla Numby (musi dostać czyste tablice numpy)
        watts_arr = df_pd['watts'].to_numpy(dtype=np.float64)
        
        if 'time' in df_pd.columns:
            time_arr = df_pd['time'].to_numpy(dtype=np.float64)
        else:
            # Jak nie ma czasu, zakładamy co 1 sekundę
            time_arr = np.arange(len(watts_arr), dtype=np.float64)

        # 3. Uruchomienie TURBO SILNIKA!
        # Tu dzieje się magia - to trwa milisekundy zamiast sekund
        w_bal = calculate_w_prime_fast(watts_arr, time_arr, float(cp), float(w_prime))

        # 4. Zapisanie wyniku
        df_pd['w_prime_balance'] = w_bal
        return df_pd

    except Exception as e:
        # Awaryjnie zwróć pusty wynik, żeby apka się nie wywaliła
        print(f"Błąd obliczeń W': {e}")
        try:
            bio = io.BytesIO(df_bytes)
            try:
                df_pd = pd.read_parquet(bio)
            except:
                bio.seek(0)
                df_pd = pd.read_csv(bio)
            df_pd['w_prime_balance'] = 0.0
            return df_pd
        except:
            return pd.DataFrame({'w_prime_balance': []})

def calculate_w_prime_balance(_df_pl_active, cp: float, w_prime: float):
    if isinstance(_df_pl_active, dict):
        df_pd = pd.DataFrame(_df_pl_active)
    elif hasattr(_df_pl_active, 'to_pandas'):
        df_pd = _df_pl_active.to_pandas()
    else:
        df_pd = _df_pl_active.copy()
    if 'time' not in df_pd.columns:
        df_pd['time'] = np.arange(len(df_pd), dtype=float)
    df_bytes = _serialize_df_to_parquet_bytes(df_pd)
    result_df = _calculate_w_prime_balance_cached(df_bytes, float(cp), float(w_prime))
    return result_df

def load_data(file):
    try:
        file.seek(0)
        df_pd = pd.read_csv(file, low_memory=False) 
    except:
        file.seek(0)
        df_pd = pd.read_csv(file, sep=';', low_memory=False)

    df_pd.columns = [str(c).lower().strip() for c in df_pd.columns]
    rename_map = {}
    if 've' in df_pd.columns and 'tymeventilation' not in df_pd.columns: rename_map['ve'] = 'tymeventilation'
    if 'ventilation' in df_pd.columns and 'tymeventilation' not in df_pd.columns: rename_map['ventilation'] = 'tymeventilation'
    if 'total_hemoglobin' in df_pd.columns and 'thb' not in df_pd.columns: rename_map['total_hemoglobin'] = 'thb'
    if rename_map: 
        df_pd = df_pd.rename(columns=rename_map)

    if 'hrv' in df_pd.columns:
        df_pd['hrv'] = df_pd['hrv'].astype(str)
        def clean_hrv_hardcore(val):
            val = val.strip().lower()
            if val == 'nan' or val == '': 
                return np.nan
            if ':' in val:
                try:
                    parts = [float(x) for x in val.split(':') if x]
                    return np.mean(parts) if parts else np.nan
                except:
                    return np.nan
            try:
                return float(val)
            except:
                return np.nan

        df_pd['hrv'] = df_pd['hrv'].apply(clean_hrv_hardcore)
        df_pd['hrv'] = pd.to_numeric(df_pd['hrv'], errors='coerce')
        df_pd['hrv'] = df_pd['hrv'].interpolate(method='linear').ffill().bfill()

    if 'time' not in df_pd.columns:
        df_pd['time'] = np.arange(len(df_pd)).astype(float)

    numeric_cols = ['watts', 'heartrate', 'cadence', 'smo2', 'thb', 'temp', 'torque', 'core_temperature', 
                    'skin_temperature', 'velocity_smooth', 'tymebreathrate', 'tymeventilation', 'rr', 'rr_interval', 'hrv', 'ibi', 'time', 'skin_temp', 'core_temp', 'power']
    
    for col in numeric_cols:
        if col in df_pd.columns:
            df_pd[col] = pd.to_numeric(df_pd[col], errors='coerce')

    return df_pd

def normalize_columns_pandas(df_pd):
    mapping = {}
    cols = [c.lower() for c in df_pd.columns]
    if 've' in cols and 'tymeventilation' not in cols:
        mapping[[c for c in df_pd.columns if c.lower() == 've'][0]] = 'tymeventilation'
    if 'ventilation' in cols and 'tymeventilation' not in cols:
        mapping[[c for c in df_pd.columns if c.lower() == 'ventilation'][0]] = 'tymeventilation'
    if 'total_hemoglobin' in cols and 'thb' not in cols:
        mapping[[c for c in df_pd.columns if c.lower() == 'total_hemoglobin'][0]] = 'thb'
    df_pd = df_pd.rename(columns=mapping)
    df_pd.columns = [c.lower() for c in df_pd.columns]
    return df_pd

def process_data(df):
    df_pd = df.to_pandas() if hasattr(df, "to_pandas") else df.copy()

    if 'time' not in df_pd.columns:
        df_pd['time'] = np.arange(len(df_pd)).astype(float)
    df_pd['time'] = pd.to_numeric(df_pd['time'], errors='coerce')
    
    # Usuń wiersze z NaN w kolumnie time przed utworzeniem indeksu
    df_pd = df_pd.dropna(subset=['time'])
    
    # Wypełnij brakujące wartości time sekwencyjnie jeśli są duplikaty lub luki
    if df_pd['time'].isna().any() or len(df_pd) == 0:
        df_pd['time'] = np.arange(len(df_pd)).astype(float)

    df_pd = df_pd.sort_values('time').reset_index(drop=True)
    df_pd['time_dt'] = pd.to_timedelta(df_pd['time'], unit='s')
    
    # Upewnij się, że indeks nie ma NaN
    df_pd = df_pd[df_pd['time_dt'].notna()]
    df_pd = df_pd.set_index('time_dt')

    num_cols = df_pd.select_dtypes(include=['float64', 'int64']).columns.tolist()
    if num_cols:
        # Użyj metody 'linear' zamiast 'time' dla większej niezawodności
        df_pd[num_cols] = df_pd[num_cols].interpolate(method='linear').ffill().bfill()

    try:
        df_numeric = df_pd.select_dtypes(include=[np.number])
        df_resampled = df_numeric.resample('1S').mean()
        df_resampled = df_resampled.interpolate(method='linear').ffill().bfill()
    except Exception:
        df_resampled = df_pd
    df_resampled['time'] = df_resampled.index.total_seconds()
    df_resampled['time_min'] = df_resampled['time'] / 60.0

    window_long = '30s'
    window_short = '5s'
    smooth_cols = ['watts', 'heartrate', 'cadence', 'smo2', 'torque', 'core_temperature',
                   'skin_temperature', 'velocity_smooth', 'tymebreathrate', 'tymeventilation', 'thb']
    
    for col in smooth_cols:
        if col in df_resampled.columns:
            df_resampled[f'{col}_smooth'] = df_resampled[col].rolling(window=window_long, min_periods=1).mean()
            df_resampled[f'{col}_smooth_5s'] = df_resampled[col].rolling(window=window_short, min_periods=1).mean()

    df_resampled = df_resampled.reset_index(drop=True)

    return df_resampled

def calculate_metrics(df_pl, cp_val):
    cols = df_pl.columns
    avg_watts = df_pl['watts'].mean() if 'watts' in cols else 0
    avg_hr = df_pl['heartrate'].mean() if 'heartrate' in cols else 0
    avg_cadence = df_pl['cadence'].mean() if 'cadence' in cols else 0
    avg_vent = df_pl['tymeventilation'].mean() if 'tymeventilation' in cols else 0
    avg_rr = df_pl['tymebreathrate'].mean() if 'tymebreathrate' in cols else 0
    power_hr = (avg_watts / avg_hr) if avg_hr > 0 else 0
    np_est = avg_watts * 1.05
    ef_factor = (np_est / avg_hr) if avg_hr > 0 else 0
    work_above_cp_kj = 0.0
    if 'watts' in cols:
        try:
            if hasattr(df_pl, "select"):
                t = df_pl['time'].to_numpy().astype(float)
                w = df_pl['watts'].to_numpy().astype(float)
            else:
                t = df_pl['time'].values.astype(float)
                w = df_pl['watts'].values.astype(float)
            dt = np.diff(t, prepend=t[0])
            if len(dt) > 1:
                dt[0] = dt[1] if dt[1] > 0 else np.median(dt[1:]) if len(dt)>2 else 1.0
            else:
                dt = np.ones_like(w)
            excess = np.maximum(w - cp_val, 0.0)
            energy_j = np.sum(excess * dt)  # w·s = J
            work_above_cp_kj = energy_j / 1000.0
        except Exception:
            df_above_cp = df_pl[df_pl['watts'] > cp_val] if 'watts' in df_pl.columns else pd.DataFrame()
            work_above_cp_kj = (df_above_cp['watts'].sum() / 1000) if len(df_above_cp)>0 else 0.0
    return {
        'avg_watts': avg_watts, 'avg_hr': avg_hr, 'avg_cadence': avg_cadence,
        'avg_vent': avg_vent, 'avg_rr': avg_rr, 'power_hr': power_hr,
        'np_est': np_est, 'ef_factor': ef_factor, 'work_above_cp_kj': work_above_cp_kj
    }

def calculate_dynamic_dfa(df_pl, window_sec=300, step_sec=30):
    """
    Oblicza metryki HRV (RMSSD, SDNN) w oknie przesuwnym.
    Działa z danymi resamplowanymi (1 Hz) i surowymi R-R.
    Zwraca pseudo-DFA bazujący na zmienności HRV.
    """

    df = df_pl.to_pandas() if hasattr(df_pl, "to_pandas") else df_pl.copy()
    
    rr_col = next((c for c in ['rr', 'rr_interval', 'hrv', 'ibi', 'r-r', 'rr_ms'] if c in df.columns), None)
    
    if rr_col is None:
        return None, "Brak kolumny z danymi R-R/HRV"

    rr_data = df[['time', rr_col]].dropna()
    rr_data = rr_data[rr_data[rr_col] > 0]
    
    if len(rr_data) < 100:
        return None, f"Za mało danych R-R ({len(rr_data)} < 100)"

    # Automatyczna detekcja jednostek
    mean_val = rr_data[rr_col].mean()
    if mean_val < 2.0:  # Prawdopodobnie sekundy
        rr_data = rr_data.copy()
        rr_data[rr_col] = rr_data[rr_col] * 1000
    elif mean_val > 2000:  # Prawdopodobnie mikrosekundy
        rr_data = rr_data.copy()
        rr_data[rr_col] = rr_data[rr_col] / 1000

    rr_values = rr_data[rr_col].values
    time_values = rr_data['time'].values

    results = []
    
    max_time = time_values[-1]
    curr_time = time_values[0] + window_sec

    while curr_time < max_time:
        mask = (time_values >= (curr_time - window_sec)) & (time_values <= curr_time)
        window_rr = rr_values[mask]
        
        if len(window_rr) >= 30:
            try:
                # Usuwamy outliers
                q1, q3 = np.percentile(window_rr, [25, 75])
                iqr = q3 - q1
                mask_valid = (window_rr > q1 - 1.5*iqr) & (window_rr < q3 + 1.5*iqr)
                clean_rr = window_rr[mask_valid]
                
                if len(clean_rr) >= 20:
                    # Oblicz RMSSD (różnice kolejnych interwałów)
                    diffs = np.diff(clean_rr)
                    rmssd = np.sqrt(np.mean(diffs**2))
                    sdnn = np.std(clean_rr)
                    mean_rr = np.mean(clean_rr)
                    
                    # Pseudo-Alpha1: normalizacja RMSSD/SDNN do skali 0.5-1.5
                    # Wysoki RMSSD/SDNN = wysoka zmienność = wysoki alpha (stan zrelaksowany)
                    # Niski RMSSD/SDNN = niska zmienność = niski alpha (stres)
                    cv = (rmssd / mean_rr) * 100  # Coefficient of variation
                    
                    # Mapowanie CV do alpha1 (empiryczne)
                    # CV ~1-2% = niska zmienność = alpha ~0.5 (stres)
                    # CV ~5-10% = wysoka zmienność = alpha ~1.0 (relaks)
                    alpha1 = 0.4 + (cv / 15.0)  # Skalowanie
                    alpha1 = np.clip(alpha1, 0.3, 1.5)
                    
                    results.append({
                        'time': curr_time, 
                        'alpha1': alpha1,
                        'rmssd': rmssd,
                        'sdnn': sdnn,
                        'mean_rr': mean_rr
                    })
            except Exception:
                pass 
        
        curr_time += step_sec

    if not results:
        return None, f"Nie udało się obliczyć HRV. Dane: {len(rr_data)} próbek"

    return pd.DataFrame(results), None

def calculate_advanced_kpi(df_pl):
    df = df_pl.to_pandas() if hasattr(df_pl, "to_pandas") else df_pl.copy()
    if 'watts_smooth' not in df.columns or 'heartrate_smooth' not in df.columns:
        return 0.0, 0.0
    df_active = df[(df['watts_smooth'] > 100) & (df['heartrate_smooth'] > 80)]
    if len(df_active) < 600: return 0.0, 0.0
    mid = len(df_active) // 2
    p1, p2 = df_active.iloc[:mid], df_active.iloc[mid:]
    hr1 = p1['heartrate_smooth'].mean()
    hr2 = p2['heartrate_smooth'].mean()
    if hr1 == 0 or hr2 == 0: return 0.0, 0.0
    ef1 = p1['watts_smooth'].mean() / hr1
    ef2 = p2['watts_smooth'].mean() / hr2
    if ef1 == 0: return 0.0, 0.0
    return ((ef1 - ef2) / ef1) * 100, (df_active['watts_smooth'] / df_active['heartrate_smooth']).mean()

def calculate_z2_drift(df_pl, cp):
    df = df_pl.to_pandas() if hasattr(df_pl, "to_pandas") else df_pl.copy()
    if 'watts_smooth' not in df.columns or 'heartrate_smooth' not in df.columns:
        return 0.0
    df_z2 = df[(df['watts_smooth'] >= 0.55*cp) & (df['watts_smooth'] <= 0.75*cp) & (df['heartrate_smooth'] > 60)]
    if len(df_z2) < 300: return 0.0
    mid = len(df_z2) // 2
    p1, p2 = df_z2.iloc[:mid], df_z2.iloc[mid:]
    hr1 = p1['heartrate_smooth'].mean()
    hr2 = p2['heartrate_smooth'].mean()
    if hr1 == 0 or hr2 == 0: return 0.0
    ef1 = p1['watts_smooth'].mean() / hr1
    ef2 = p2['watts_smooth'].mean() / hr2
    return ((ef1 - ef2) / ef1) * 100 if ef1 != 0 else 0.0

def calculate_heat_strain_index(df_pl):
    df = df_pl.to_pandas() if hasattr(df_pl, "to_pandas") else df_pl.copy()
    core_col = 'core_temperature_smooth' if 'core_temperature_smooth' in df.columns else None
    if not core_col or 'heartrate_smooth' not in df.columns:
        df['hsi'] = None
        return df
    df['hsi'] = ((5 * (df[core_col] - 37.0) / 2.5) + (5 * (df['heartrate_smooth'] - 60.0) / 120.0)).clip(0.0, 10.0)
    return df

def calculate_vo2max(mmp_5m, rider_weight):
    if mmp_5m is None or pd.isna(mmp_5m) or rider_weight <= 0: return 0.0
    return (10.8 * mmp_5m / rider_weight) + 7

def calculate_trend(x, y):
    try:
        idx = np.isfinite(x) & np.isfinite(y)
        if np.sum(idx) < 2: return None
        z = np.polyfit(x[idx], y[idx], 1)
        p = np.poly1d(z)
        return p(x)
    except: return None

def apply_chart_style(fig, title=None):
    fig.update_layout(
        template="plotly_dark",
        plot_bgcolor='rgba(0,0,0,0)',
        paper_bgcolor='rgba(0,0,0,0)',
        title=dict(
            text=title, 
            font=dict(family="Rajdhani", size=24, color="#f0f6fc")
        ) if title else None,
        font=dict(family="Inter", size=12, color="#c9d1d9"),
        xaxis=dict(showgrid=False, zeroline=False, showline=True, linecolor='#30363d'),
        yaxis=dict(showgrid=True, gridcolor='rgba(255,255,255,0.05)', zeroline=False),
        margin=dict(l=20, r=20, t=50, b=20),
        hovermode="x unified"
    )
    return fig

# --- APP START ---

st.title("⚡ Pro Athlete Dashboard")

st.sidebar.header("Ustawienia Zawodnika")
with st.sidebar.expander("⚙️ Parametry Fizyczne", expanded=True):
    rider_weight = st.number_input("Waga Zawodnika [kg]", value=95.0, step=0.5, min_value=30.0, max_value=200.0, key="weight")
    rider_height = st.number_input("Wzrost [cm]", value=180, step=1, min_value=100, max_value=250, key="height")
    rider_age = st.number_input("Wiek [lata]", value=30, step=1, min_value=10, max_value=100, key="age")
    is_male = st.checkbox("Mężczyzna?", value=True, key="gender_m")
    
    st.markdown("---")
    vt1_watts = st.number_input("VT1 (Próg Tlenowy) [W]", value=280, min_value=0, key="vt1_w")
    vt2_watts = st.number_input("VT2 (Próg Beztlenowy/FTP) [W]", value=400, min_value=0, key="vt2_w")
    
    st.divider()
    st.markdown("### 🫁 Wentylacja [L/min]")
    vt1_vent = st.number_input("VT1 (Próg Tlenowy) [L/min]", value=79.0, min_value=0.0, key="vt1_v")
    vt2_vent = st.number_input("VT2 (Próg Beztlenowy) [L/min]", value=136.0, min_value=0.0, key="vt2_v")

st.sidebar.divider()
cp_input = st.sidebar.number_input("Moc Krytyczna (CP/FTP) [W]", value=410, min_value=1, key="cp_in")
w_prime_input = st.sidebar.number_input("W' (W Prime) [J]", value=31000, min_value=0, key="wp_in")
st.sidebar.divider()
crank_length = st.sidebar.number_input("Długość korby [mm]", value=160.0, key="crank")
uploaded_file = st.sidebar.file_uploader("Wgraj plik (CSV / TXT)", type=['csv', 'txt'])

if rider_weight <= 0 or cp_input <= 0:
    st.error("Błąd: Waga i CP muszą być większe od zera.")
    st.stop()

if uploaded_file is not None:
    with st.spinner('Przetwarzanie danych...'):
        try:
            df_raw = load_data(uploaded_file)
            df_clean_pl = process_data(df_raw)
            metrics = calculate_metrics(df_clean_pl, cp_input)
            df_w_prime = calculate_w_prime_balance(df_clean_pl, cp_input, w_prime_input)
            decoupling_percent, ef_factor = calculate_advanced_kpi(df_clean_pl)
            drift_z2 = calculate_z2_drift(df_clean_pl, cp_input)
            df_with_hsi = calculate_heat_strain_index(df_w_prime)
            df_plot = df_with_hsi.copy()
            
            if 'smo2' in df_plot.columns:
                 df_plot['smo2_smooth_ultra'] = df_plot['smo2'].rolling(window=60, center=True, min_periods=1).mean()
            df_plot_resampled = df_plot.iloc[::5, :] if len(df_plot) > 10000 else df_plot
            
            # --- SEKCJA AI / MLX ---
            if MLX_AVAILABLE and os.path.exists(MODEL_FILE):
                try:
                    # Próbujemy odpalić predykcję
                    auto_pred = predict_only(df_plot_resampled)
                    
                    if auto_pred is not None:
                        df_plot_resampled['ai_hr'] = auto_pred
                    else:
                        st.sidebar.warning("⚠️ AI zwróciło pusty wynik (None). Sprawdź load_model.")
                except Exception as e:
                    st.sidebar.error(f"💥 Krytyczny błąd w Auto-Inference: {e}")
            elif not os.path.exists(MODEL_FILE):
                # Tylko info, nie błąd - użytkownik może jeszcze nie trenował
                pass 
            # ----------------------------------

        except Exception as e:  # <--- TEGO BRAKOWAŁO!
            st.error(f"Błąd wczytywania pliku: {e}")
            st.stop()

        # --- HEADER METRICS ---
        if 'watts' in df_plot.columns:
            rolling_30s_header = df_plot['watts'].rolling(window=30, min_periods=1).mean()
            np_header = np.power(np.mean(np.power(rolling_30s_header, 4)), 0.25)
            if pd.isna(np_header): np_header = metrics['avg_watts']
        else:
            np_header = 0

        if cp_input > 0:
            if_header = np_header / cp_input
            duration_sec = len(df_plot)
            tss_header = (duration_sec * np_header * if_header) / (cp_input * 3600) * 100
        else:
            tss_header = 0; if_header = 0

        # ===== STICKY HEADER - PANEL Z KLUCZOWYMI METRYKAMI =====
        st.markdown("""
        <style>
        .sticky-metrics {
            position: sticky;
            top: 60px;
            z-index: 999;
            background: linear-gradient(135deg, #1a1f25 0%, #0e1117 100%);
            padding: 15px;
            border-radius: 10px;
            border: 1px solid #30363d;
            box-shadow: 0 4px 12px rgba(0, 0, 0, 0.5);
            margin-bottom: 20px;
            backdrop-filter: blur(10px);
        }
        .sticky-metrics h4 {
            margin: 0 0 10px 0;
            color: #00cc96;
            font-size: 14px;
            text-transform: uppercase;
            letter-spacing: 1px;
        }
        .metric-row {
            display: flex;
            justify-content: space-around;
            flex-wrap: wrap;
            gap: 10px;
        }
        .metric-box {
            flex: 1;
            min-width: 120px;
            background: rgba(255, 255, 255, 0.03);
            padding: 10px;
            border-radius: 8px;
            text-align: center;
            border: 1px solid rgba(255, 255, 255, 0.1);
        }
        .metric-box .label {
            font-size: 11px;
            color: #8b949e;
            text-transform: uppercase;
        }
        .metric-box .value {
            font-size: 20px;
            font-weight: 700;
            color: #f0f6fc;
            margin-top: 5px;
        }
        .metric-box .unit {
            font-size: 12px;
            color: #8b949e;
        }
        </style>
        """, unsafe_allow_html=True)

        # Oblicz metryki dla sticky panelu
        avg_power = metrics.get('avg_watts', 0)
        avg_hr = metrics.get('avg_hr', 0)
        avg_smo2 = df_plot['smo2'].mean() if 'smo2' in df_plot.columns else 0
        avg_cadence = metrics.get('avg_cadence', 0)
        avg_ve = metrics.get('avg_vent', 0)
        duration_min = len(df_plot) / 60 if len(df_plot) > 0 else 0

        st.markdown(f"""
        <div class="sticky-metrics">
            <h4>⚡ Live Training Summary</h4>
            <div class="metric-row">
                <div class="metric-box">
                    <div class="label">Avg Power</div>
                    <div class="value">{avg_power:.0f} <span class="unit">W</span></div>
                </div>
                <div class="metric-box">
                    <div class="label">Avg HR</div>
                    <div class="value">{avg_hr:.0f} <span class="unit">bpm</span></div>
                </div>
                <div class="metric-box">
                    <div class="label">Avg SmO2</div>
                    <div class="value">{avg_smo2:.1f} <span class="unit">%</span></div>
                </div>
                <div class="metric-box">
                    <div class="label">Cadence</div>
                    <div class="value">{avg_cadence:.0f} <span class="unit">rpm</span></div>
                </div>
                <div class="metric-box">
                    <div class="label">Avg VE</div>
                    <div class="value">{avg_ve:.0f} <span class="unit">L/min</span></div>
                </div>
                <div class="metric-box">
                    <div class="label">Duration</div>
                    <div class="value">{duration_min:.0f} <span class="unit">min</span></div>
                </div>
            </div>
        </div>
        """, unsafe_allow_html=True)
        # ===== KONIEC STICKY HEADER =====

        m1, m2, m3 = st.columns(3)
        m1.metric("NP (Norm. Power)", f"{np_header:.0f} W", help="Normalized Power (Coggan Formula)")
        m2.metric("TSS", f"{tss_header:.0f}", help=f"IF: {if_header:.2f}")
        m3.metric("Praca [kJ]", f"{df_plot['watts'].sum()/1000:.0f}")
        
        # --- ZAKŁADKI ---
        tab_raport, tab_kpi, tab_power, tab_hrv, tab_biomech, tab_thermal, tab_trends, tab_nutrition, tab_smo2, tab_hemo, tab_vent, tab_limiters, tab_model, tab_ai = st.tabs(
            ["Raport", "KPI", "Power", "HRV", "Biomech", "Thermal", "Trends", "Nutrition", "SmO2 Analysis", "Hematology Analysis", "Ventilation Analysis", "Limiters Analysis", "Model Analysis", "AI Coach"]
        )
                
       # --- TAB RAPORT ---
        with tab_raport:
            st.header("Executive Summary")
            
            st.subheader("Przebieg Treningu")
            fig_exec = go.Figure()
            
            if 'watts_smooth' in df_plot.columns:
                fig_exec.add_trace(go.Scatter(x=df_plot['time_min'], y=df_plot['watts_smooth'], name='Moc', fill='tozeroy', line=dict(color=Config.COLOR_POWER, width=1), hovertemplate="Moc: %{y:.0f} W<extra></extra>"))
            if 'heartrate_smooth' in df_plot.columns:
                fig_exec.add_trace(go.Scatter(x=df_plot['time_min'], y=df_plot['heartrate_smooth'], name='HR', line=dict(color=Config.COLOR_HR, width=2), yaxis='y2', hovertemplate="HR: %{y:.0f} bpm<extra></extra>"))
            if 'smo2_smooth' in df_plot.columns:
                fig_exec.add_trace(go.Scatter(x=df_plot['time_min'], y=df_plot['smo2_smooth'], name='SmO2', line=dict(color=Config.COLOR_SMO2, width=2, dash='dot'), yaxis='y3', hovertemplate="SmO2: %{y:.1f}%<extra></extra>"))
            if 'tymeventilation_smooth' in df_plot.columns:
                fig_exec.add_trace(go.Scatter(x=df_plot['time_min'], y=df_plot['tymeventilation_smooth'], name='VE', line=dict(color=Config.COLOR_VE, width=2, dash='dash'), yaxis='y4', hovertemplate="VE: %{y:.1f} L/min<extra></extra>"))

            fig_exec.update_layout(
                template="plotly_dark", height=500,
                yaxis=dict(title="Moc [W]"),
                yaxis2=dict(title="HR", overlaying='y', side='right', showgrid=False),
                yaxis3=dict(title="SmO2", overlaying='y', side='right', showgrid=False, showticklabels=False, range=[0, 100]),
                yaxis4=dict(title="VE", overlaying='y', side='right', showgrid=False, showticklabels=False),
                legend=dict(orientation="h", y=1.05, x=0), hovermode="x unified"
            )
            st.plotly_chart(fig_exec, use_container_width=True)

            st.markdown("---")
            col_dist1, col_dist2 = st.columns(2)
            with col_dist1:
                st.subheader("Czas w Strefach (Moc)")
                if 'watts' in df_plot.columns:
                    bins = [0, 0.55*cp_input, 0.75*cp_input, 0.90*cp_input, 1.05*cp_input, 1.20*cp_input, 10000]
                    labels = ['Z1', 'Z2', 'Z3', 'Z4', 'Z5', 'Z6']
                    colors = ['#808080', '#32CD32', '#FFD700', '#FF8C00', '#FF4500', '#8B0000']
                    df_z = df_plot.copy()
                    df_z['Zone'] = pd.cut(df_z['watts'], bins=bins, labels=labels, right=False)
                    pcts = (df_z['Zone'].value_counts().sort_index() / len(df_z) * 100).round(1)
                    fig_hist = go.Figure(go.Bar(x=pcts.values, y=labels, orientation='h', marker_color=colors, text=pcts.apply(lambda x: f"{x}%"), textposition='auto'))
                    fig_hist.update_layout(template="plotly_dark", height=250, xaxis=dict(visible=False), yaxis=dict(showgrid=False), margin=dict(t=20, b=20))
                    st.plotly_chart(fig_hist, use_container_width=True)
            
            with col_dist2:
                st.subheader("Rozkład Tętna")
                if 'heartrate' in df_plot.columns:
                    hr_counts = df_plot['heartrate'].dropna().round().astype(int).value_counts().sort_index()
                    fig_hr = go.Figure(go.Bar(x=hr_counts.index, y=hr_counts.values, marker_color=Config.COLOR_HR, hovertemplate="<b>%{x} BPM</b><br>Czas: %{y} s<extra></extra>"))
                    fig_hr.update_layout(template="plotly_dark", height=250, xaxis_title="BPM", yaxis=dict(visible=False), bargap=0.1, margin=dict(t=20, b=20))
                    st.plotly_chart(fig_hr, use_container_width=True)

            st.markdown("---")
            c_bot1, c_bot2 = st.columns(2)
            with c_bot1:
                st.subheader("🏆 Peak Power")
                mmp_windows = {'5s': 5, '1m': 60, '5m': 300, '20m': 1200, '60m': 3600}
                cols = st.columns(5)
                if 'watts' in df_plot.columns:
                    for c, (l, s) in zip(cols, mmp_windows.items()):
                        val = df_plot['watts'].rolling(s).mean().max()
                        with c:
                            if not pd.isna(val): st.metric(l, f"{val:.0f} W", f"{val/rider_weight:.1f} W/kg")
                            else: st.metric(l, "--")
            
            with c_bot2:
                st.subheader("🎯 Strefy (wg CP)")
                z2_l, z2_h = int(0.56*cp_input), int(0.75*cp_input)
                z3_l, z3_h = int(0.76*cp_input), int(0.90*cp_input)
                z4_l, z4_h = int(0.91*cp_input), int(1.05*cp_input)
                z5_l, z5_h = int(1.06*cp_input), int(1.20*cp_input)
                st.info(f"**Z2 (Baza):** {z2_l}-{z2_h} W | **Z3 (Tempo):** {z3_l}-{z3_h} W | **Z4 (Próg):** {z4_l}-{z4_h} W")

        # --- TAB KPI ---
        with tab_kpi:
            st.header("Kluczowe Wskaźniki Wydajności (KPI)")
            c1, c2, c3, c4, c5 = st.columns(5)
            c1.metric("Średnia Moc", f"{metrics['avg_watts']:.0f} W")
            c2.metric("Średnie Tętno", f"{metrics['avg_hr']:.0f} BPM")
            c3.metric("Średnie SmO2", f"{df_plot['smo2'].mean() if 'smo2' in df_plot.columns else 0:.1f} %")
            c4.metric("Kadencja", f"{metrics['avg_cadence']:.0f} RPM")
            vo2max_est = calculate_vo2max(df_plot['watts'].rolling(window=300).mean().max() if 'watts' in df_plot.columns else 0, rider_weight)
            c5.metric("Szac. VO2max", f"{vo2max_est:.1f}", help="Estymowane na podstawie mocy 5-minutowej (ACSM).")
                     
            st.divider()
            c5, c6, c7, c8 = st.columns(4)
            c5.metric("Power/HR", f"{metrics['power_hr']:.2f}")
            c6.metric("Efficiency (EF)", f"{metrics['ef_factor']:.2f}")
            c7.metric("Praca > CP", f"{metrics['work_above_cp_kj']:.0f} kJ")
            c8.metric("Wentylacja (VE)", f"{metrics['avg_vent']:.1f} L/min")
            st.divider()
            c9, c10, c11, c12 = st.columns(4)
            c9.metric("Dryf (Pa:Hr)", f"{decoupling_percent:.1f} %", delta_color="inverse" if decoupling_percent<5 else "normal")
            c10.metric("Dryf Z2", f"{drift_z2:.1f} %", delta_color="inverse" if drift_z2<5 else "normal")
            max_hsi = df_plot['hsi'].max() if 'hsi' in df_plot.columns else 0
            c11.metric("Max HSI", f"{max_hsi:.1f}", delta_color="normal" if max_hsi>5 else "inverse")
            c12.metric("Oddechy (RR)", f"{metrics['avg_rr']:.1f} /min")

            st.subheader("Wizualizacja Dryfu i Zmienności")
            if 'watts_smooth' in df_plot.columns:
                fig_dec = go.Figure()
                fig_dec.add_trace(go.Scatter(x=df_plot_resampled['time_min'], y=df_plot_resampled['watts_smooth'], name='Moc', line=dict(color=Config.COLOR_POWER, width=1.5), hovertemplate="Moc: %{y:.0f} W<extra></extra>"))
                if 'heartrate_smooth' in df_plot.columns: fig_dec.add_trace(go.Scatter(x=df_plot_resampled['time_min'], y=df_plot_resampled['heartrate_smooth'], name='HR', yaxis='y2', line=dict(color=Config.COLOR_HR, width=1.5), hovertemplate="HR: %{y:.0f} BPM<extra></extra>"))
                if 'smo2_smooth' in df_plot.columns: fig_dec.add_trace(go.Scatter(x=df_plot_resampled['time_min'], y=df_plot_resampled['smo2_smooth'], name='SmO2', yaxis='y3', line=dict(color=Config.COLOR_SMO2, dash='dot', width=1.5), hovertemplate="SmO2: %{y:.1f}%<extra></extra>"))
                
                fig_dec.update_layout(template="plotly_dark", title="Dryf Mocy, Tętna i SmO2 w Czasie", hovermode="x unified",
                    yaxis=dict(title="Moc [W]"),
                    yaxis2=dict(title="HR [bpm]", overlaying='y', side='right', showgrid=False),
                    yaxis3=dict(title="SmO2 [%]", overlaying='y', side='right', showgrid=False, showticklabels=False, range=[0, 100]),
                    legend=dict(orientation="h", y=1.1, x=0))
                st.plotly_chart(fig_dec, use_container_width=True)
                
                st.info("""
                **💡 Interpretacja: Fizjologia Zmęczenia (Triada: Moc - HR - SmO2)**

                Ten wykres pokazuje "koszt fizjologiczny" utrzymania zadanej mocy w czasie.

                **1. Stan Idealny (Brak Dryfu):**
                * **Moc (Zielony):** Linia płaska (stałe obciążenie).
                * **Tętno (Czerwony):** Linia płaska (równoległa do mocy).
                * **SmO2 (Fiolet):** Stabilne.
                * **Wniosek:** Jesteś w pełnej równowadze tlenowej. Możesz tak jechać godzinami.

                **2. Dryf Sercowo-Naczyniowy (Cardiac Drift):**
                * **Moc:** Stała.
                * **Tętno:** Powoli rośnie (rozjeżdża się z linią mocy).
                * **SmO2:** Stabilne.
                * **Przyczyna:** Odwodnienie (spadek objętości osocza) lub przegrzanie (krew ucieka do skóry). Serce musi bić szybciej, by pompować tę samą ilość tlenu.

                **3. Zmęczenie Metaboliczne (Metabolic Fatigue):**
                * **Moc:** Stała.
                * **Tętno:** Stabilne lub lekko rośnie.
                * **SmO2:** **Zaczyna spadać.**
                * **Przyczyna:** Mięśnie tracą wydajność (rekrutacja włókien szybkokurczliwych II typu, które zużywają więcej tlenu). To pierwszy sygnał nadchodzącego "odcięcia".

                **4. "Zgon" (Bonking/Failure):**
                * **Moc:** Zaczyna spadać (nie jesteś w stanie jej utrzymać).
                * **Tętno:** Może paradoksalnie spadać (zmęczenie układu nerwowego) lub rosnąć (panika organizmu).
                * **SmO2:** Gwałtowny spadek lub chaotyczne skoki.
                """)

        # --- TAB POWER ---
        with tab_power:
            st.subheader("Wykres Mocy i W'")
            fig_pw = go.Figure()
            fig_pw.add_trace(go.Scatter(x=df_plot_resampled['time_min'], y=df_plot_resampled['watts_smooth'], name="Moc", fill='tozeroy', line=dict(color=Config.COLOR_POWER, width=1), hovertemplate="Moc: %{y:.0f} W<extra></extra>"))
            fig_pw.add_trace(go.Scatter(x=df_plot_resampled['time_min'], y=df_plot_resampled['w_prime_balance'], name="W' Bal", yaxis="y2", line=dict(color=Config.COLOR_HR, width=2), hovertemplate="W' Bal: %{y:.0f} J<extra></extra>"))
            fig_pw.update_layout(template="plotly_dark", title="Zarządzanie Energią (Moc vs W')", hovermode="x unified", yaxis=dict(title="Moc [W]"), yaxis2=dict(title="W' Balance [J]", overlaying="y", side="right", showgrid=False))
            st.plotly_chart(fig_pw, use_container_width=True)
            
            st.info("""
            **💡 Interpretacja: Energia Beztlenowa (W' Balance)**

            Ten wykres pokazuje, ile "zapałek" masz jeszcze w pudełku.

            * **Czerwona Linia (W' Bal):** Poziom energii beztlenowej w Dżulach [J].
            * **Moc Krytyczna (CP):** To Twoja granica tlenowa (jak FTP, ale fizjologicznie precyzyjniejsza).

            **Jak to działa?**
            * **Moc < CP (Strefa Tlenowa):** Nie spalasz W'. Jeśli jechałeś mocno wcześniej, bateria się ładuje (czerwona linia rośnie).
            * **Moc > CP (Strefa Beztlenowa):** Zaczynasz "palić zapałki". Czerwona linia spada. Im mocniej depczesz, tym szybciej spada.
            * **W' = 0 J (Wyczerpanie):** "Odcina prąd". Nie jesteś w stanie utrzymać mocy powyżej CP ani sekundy dłużej. Musisz zwolnić, żeby zregenerować.

            **Scenariusze:**
            1.  **Interwały:** W' powinno spadać w trakcie powtórzenia (wysiłek) i rosnąć w przerwie (regeneracja). Jeśli nie wraca do 100% przed kolejnym startem, kumulujesz zmęczenie.
            2.  **Finisz:** Idealnie rozegrany wyścig to taki, gdzie W' spada do zera dokładnie na linii mety. Jeśli zostało Ci 10kJ, mogłeś finiszować mocniej. Jeśli spadło do zera 500m przed metą - przeszarżowałeś.
            3.  **Błędne CP:** Jeśli podczas spokojnej jazdy W' ciągle spada, Twoje CP jest ustawione za wysoko. Jeśli finiszujesz "w trupa", a W' pokazuje wciąż 50% - Twoje CP lub W' są niedoszacowane.
            """)

            st.subheader("Czas w Strefach Mocy (Time in Zones)")
            if 'watts' in df_plot.columns:
                bins = [0, 0.55*cp_input, 0.75*cp_input, 0.90*cp_input, 1.05*cp_input, 1.20*cp_input, 10000]
                labels = ['Z1: Regeneracja', 'Z2: Wytrzymałość', 'Z3: Tempo', 'Z4: Próg', 'Z5: VO2Max', 'Z6: Beztlenowa']
                colors = ['#A0A0A0', '#32CD32', '#FFD700', '#FF8C00', '#FF4500', '#8B0000']
                df_z = df_plot.copy()
                df_z['Zone'] = pd.cut(df_z['watts'], bins=bins, labels=labels, right=False)
                pcts = (df_z['Zone'].value_counts().sort_index() / len(df_z) * 100).round(1)
                fig_z = px.bar(x=pcts.values, y=labels, orientation='h', text=pcts.apply(lambda x: f"{x}%"), color=labels, color_discrete_sequence=colors)
                fig_z.update_layout(template="plotly_dark", showlegend=False)
                st.plotly_chart(apply_chart_style(fig_z), use_container_width=True)

                st.info("""
                **💡 Interpretacja Treningowa:**
                * **Polaryzacja:** Dobry plan często ma dużo Z1/Z2 (baza) i trochę Z5/Z6 (bodziec), a mało "śmieciowych kilometrów" w Z3. Strefa Z3 to "szara strefa", która męczy, ale nie daje dużych korzyści adaptacyjnych, jednakże zużywa dużo glikogenu. Mimo tego, w triathlonie Z3 ma swoje miejsce (jazda na czas) i warto ją stosować taktycznie.
                * **Długie Wyścigi (Triathlon):** Większość czasu powinna być w Z2, z akcentami w Z4 (próg mleczanowy) i Z5 (VO2Max) dla poprawy wydolności. Spędzanie czasu w Z3 powinno być ograniczone ale taktyczne (np. jazda na czas).
                * **Sprinty i Criterium:** Więcej czasu w Z4/Z5/Z6, ale z odpowiednią regeneracją w Z1. Dużo interwałów wysokiej intensywności. Ważne jest, aby nie zaniedbywać Z2 dla budowy bazy tlenowej.
                * **Regeneracja:** Z1 to strefa regeneracyjna, idealna na dni odpoczynku lub bardzo lekkie sesje. Może pomóc w usuwaniu metabolitów i poprawie krążenia bez dodatkowego stresu. "Nie trenować" to też trening.
                * **Adaptacje Fizjologiczne:**
                * **Z1 (Szary):** Regeneracja i krążenie.
                * **Z2 (Zielony):** Kluczowe dla budowania mitochondriów i spalania tłuszczu. Podstawa wytrzymałości.
                * **Z3 (Żółty):** Mieszana strefa, poprawia ekonomię jazdy i tolerancję na wysiłek, ale może prowadzić do zmęczenia bez odpowiedniej regeneracji.
                * **Z4/Z5 (Pomarańczowy/Czerwony):** Budują tolerancję na mleczan i VO2Max, ale wymagają długiej regeneracji. Nie powinny dominować w planie treningowym.
                """)

                st.markdown("### 📚 Kompendium Fizjologii Stref (Deep Dive)")
                with st.expander("🟩 Z1/Z2: Fundament Tlenowy (< 75% CP)", expanded=True):
                    st.markdown("""
                    * **Metabolizm:** Dominacja Wolnych Kwasów Tłuszczowych (WKT). RER ~0.7-0.85. Oszczędność glikogenu.
                    * **Fizjologia:**
                        * Biogeneza mitochondriów (więcej "pieców" energetycznych).
                        * Angiogeneza (tworzenie nowych naczyń włosowatych).
                        * Wzrost aktywności enzymów oksydacyjnych.
                    * **Biomechanika:** Rekrutacja głównie włókien wolnokurczliwych (Typ I).
                    * **SmO2:** Stabilne, wysokie wartości (Równowaga Podaż=Popyt).
                    * **Oddech (VT):** Poniżej VT1. Pełna konwersacja.
                    * **Typowy Czas:** 1.5h - 6h+.
                    """)

                with st.expander("🟨 Z3: Tempo / Sweet Spot (76-90% CP)"):
                    st.markdown("""
                    * **Metabolizm:** Miks węglowodanów i tłuszczów (RER ~0.85-0.95). Zaczyna się znaczne zużycie glikogenu.
                    * **Fizjologia:** "Strefa Szara". Bodziec tlenowy, ale już z narastającym zmęczeniem.
                    * **Zastosowanie:** Trening specyficzny pod 70.3 / Ironman (długie utrzymanie mocy).
                    * **SmO2:** Stabilne, ale niższe niż w Z2. Możliwy powolny trend spadkowy.
                    * **Oddech (VT):** Okolice VT1. Głęboki, rytmiczny oddech.
                    * **Typowy Czas:** 45 min - 2.5h.
                    """)

                with st.expander("🟧 Z4: Próg Mleczanowy (91-105% CP)"):
                    st.markdown("""
                    * **Metabolizm:** Dominacja glikogenu (RER ~1.0). Produkcja mleczanu równa się jego utylizacji (MLSS).
                    * **Fizjologia:** Poprawa tolerancji na kwasicę. Zwiększenie magazynów glikogenu.
                    * **Biomechanika:** Rekrutacja włókien pośrednich (Typ IIa).
                    * **SmO2:** Granica równowagi. Utrzymuje się na stałym, niskim poziomie.
                    * **Oddech (VT):** Pomiędzy VT1 a VT2. Oddech mocny, utrudniona mowa.
                    * **Typowy Czas:** Interwały 8-30 min (łącznie do 60-90 min w sesji).
                    """)

                with st.expander("🟥 Z5/Z6: VO2Max i Beztlenowa (> 106% CP)"):
                    st.markdown("""
                    * **Metabolizm:** Wyłącznie glikogen + Fosfokreatyna (PCr). RER > 1.1.
                    * **Fizjologia:** Maksymalny pobór tlenu (pułap tlenowy). Szybkie narastanie długu tlenowego.
                    * **Biomechanika:** Pełna rekrutacja wszystkich włókien (Typ IIx). Duży moment siły.
                    * **SmO2:** Gwałtowny spadek (Desaturacja).
                    * **Oddech (VT):** Powyżej VT2 (RCP). Hiperwentylacja.
                    * **Typowy Czas:** Z5: 3-8 min. Z6: < 2 min.
                    """)
            
            st.divider()
            st.subheader("🔥 Symulator 'Spalania Zapałek' (W' Attack)")
            st.markdown("Sprawdź, jak konkretny atak wpłynie na Twoje rezerwy energii.")

            c_sim1, c_sim2 = st.columns(2)
            with c_sim1:
                sim_watts = st.slider("Moc Ataku [W]", min_value=int(cp_input), max_value=int(cp_input*2.5), value=int(cp_input*1.2), step=10)
                sim_dur = st.slider("Czas Trwania [sek]", min_value=10, max_value=300, value=60, step=10)

                if sim_watts > cp_input:
                    w_burned = (sim_watts - cp_input) * sim_dur
                    w_rem = w_prime_input - w_burned
                    w_rem_pct = (w_rem / w_prime_input) * 100
                else:
                    w_burned = 0; w_rem = w_prime_input; w_rem_pct = 100
                if w_rem < 0: w_rem = 0; w_rem_pct = 0
                st.markdown(f"**Spalone:** {w_burned:.0f} J\n**Pozostało:** {w_rem:.0f} J ({w_rem_pct:.1f}%)")
            with c_sim2:
                fig_g = go.Figure(go.Indicator(
                    mode = "gauge+number",
                    value = w_rem,
                    domain = {'x': [0, 1], 'y': [0, 1]},
                    title = {'text': "Pozostałe W'"},
                    gauge = {
                        'axis': {'range': [0, w_prime_input], 'tickwidth': 1},
                        'bar': {'color': "darkblue"},
                        'steps': [
                            {'range': [0, w_prime_input*0.25], 'color': "red"},
                            {'range': [w_prime_input*0.25, w_prime_input*0.5], 'color': "orange"},
                            {'range': [w_prime_input*0.5, w_prime_input], 'color': "green"}],
                    }
                ))
                # fig_g.update_layout(height=250, margin=dict(l=20,r=20,t=30,b=20), template="plotly_dark")
                st.plotly_chart(apply_chart_style(fig_g), use_container_width=True)
            
            if w_rem_pct == 0:
                st.error("💀 **TOTAL FAILURE!** Ten atak wyczerpie Cię całkowicie. Nie dojedziesz.")
            elif w_rem_pct < 25:
                st.warning("⚠️ **KRYTYCZNIE:** Bardzo ryzykowny atak. Zostaniesz na oparach.")
            else:
                st.success("✅ **BEZPIECZNIE:** Masz zapas na taki ruch.")

            # --- PULSE POWER (EFICIENCY) ---
            st.divider()
            st.subheader("🫀 Pulse Power (Moc na Uderzenie Serca)")
            
            if 'watts_smooth' in df_plot_resampled.columns and 'heartrate_smooth' in df_plot_resampled.columns:
                
                mask_pp = (df_plot_resampled['watts_smooth'] > 50) & (df_plot_resampled['heartrate_smooth'] > 90)
                df_pp = df_plot_resampled[mask_pp].copy()
                
                if not df_pp.empty:
                    df_pp['pulse_power'] = df_pp['watts_smooth'] / df_pp['heartrate_smooth']
                    
                    df_pp['pp_smooth'] = df_pp['pulse_power'].rolling(window=12, center=True).mean() 
                    x_pp = df_pp['time_min']
                    y_pp = df_pp['pulse_power']
                    valid_idx = np.isfinite(x_pp) & np.isfinite(y_pp)
                    
                    if valid_idx.sum() > 100:
                        slope_pp, intercept_pp, _, _, _ = stats.linregress(x_pp[valid_idx], y_pp[valid_idx])
                        trend_line_pp = intercept_pp + slope_pp * x_pp
                        total_drop = (trend_line_pp.iloc[-1] - trend_line_pp.iloc[0]) / trend_line_pp.iloc[0] * 100
                    else:
                        slope_pp = 0; total_drop = 0; trend_line_pp = None

                    avg_pp = df_pp['pulse_power'].mean()
                    
                    c_pp1, c_pp2, c_pp3 = st.columns(3)
                    c_pp1.metric("Średnie Pulse Power", f"{avg_pp:.2f} W/bpm", help="Ile watów generuje jedno uderzenie serca.")
                    
                    drift_color = "normal"
                    if total_drop < -5: drift_color = "inverse"
                    
                    c_pp2.metric("Zmiana Efektywności (Trend)", f"{total_drop:.1f}%", delta_color=drift_color)
                    c_pp3.metric("Interpretacja", "Stabilna Wydolność" if total_drop > -5 else "Dryf / Zmęczenie")

                    fig_pp = go.Figure()
                    
                    fig_pp.add_trace(go.Scatter(
                        x=df_pp['time_min'], 
                        y=df_pp['pp_smooth'], 
                        name='Pulse Power (W/bpm)', 
                        mode='lines',
                        line=dict(color='#FFD700', width=2), # Złoty kolor
                        hovertemplate="Pulse Power: %{y:.2f} W/bpm<extra></extra>"
                    ))
                    
                    if trend_line_pp is not None:
                        fig_pp.add_trace(go.Scatter(
                            x=x_pp, y=trend_line_pp,
                            name='Trend',
                            mode='lines',
                            line=dict(color='white', width=1.5, dash='dash'),
                            hoverinfo='skip'
                        ))
                    
                    fig_pp.add_trace(go.Scatter(
                        x=df_pp['time_min'], y=df_pp['watts_smooth'],
                        name='Moc (tło)',
                        yaxis='y2',
                        line=dict(width=0),
                        fill='tozeroy',
                        fillcolor='rgba(255,255,255,0.05)',
                        hoverinfo='skip'
                    ))

                    fig_pp.update_layout(
                        template="plotly_dark",
                        title="Pulse Power: Koszt Energetyczny Serca",
                        hovermode="x unified",
                        xaxis=dict(title="Czas [min]"),
                        yaxis=dict(title="Pulse Power [W / bpm]"),
                        yaxis2=dict(overlaying='y', side='right', showgrid=False, visible=False),
                        margin=dict(l=10, r=10, t=40, b=10),
                        legend=dict(orientation="h", y=1.05, x=0),
                        height=450
                    )
                    
                    st.plotly_chart(fig_pp, use_container_width=True)
                    
                    st.info("""
                    **💡 Jak to czytać?**
                    
                    * **Pulse Power (W/bpm)** mówi nam o objętości wyrzutowej serca i ekstrakcji tlenu. Im wyżej, tym lepiej.
                    * **Trend Płaski:** Idealnie. Twoje serce pracuje tak samo wydajnie w 1. minucie jak w 60. minucie. Jesteś dobrze nawodniony i chłodzony.
                    * **Trend Spadkowy (Dryf):** Serce musi bić coraz szybciej, żeby utrzymać te same waty.
                        * **Spadek < 5%:** Norma fizjologiczna.
                        * **Spadek > 10%:** Odwodnienie, przegrzanie lub wyczerpanie zapasów glikogenu w mięśniach. Czas zjeść i pić!
                    """)
                else:
                    st.warning("Zbyt mało danych (jazda poniżej 50W lub HR poniżej 90bpm), aby obliczyć wiarygodne Pulse Power.")
            else:
                st.error("Brak danych mocy lub tętna.")
                
            # --- GROSS EFFICIENCY ---
            st.divider()
            st.subheader("⚙️ Gross Efficiency (GE%) - Estymacja")
            st.caption("Stosunek mocy generowanej (Waty) do spalanej energii (Metabolizm). Typowo: 18-23%.")

            # 1. Sprawdzamy, czy mamy potrzebne dane
            if 'watts_smooth' in df_plot_resampled.columns and 'heartrate_smooth' in df_plot_resampled.columns:
                
                # 2. Obliczamy Moc Metaboliczną (Wzór Keytela na podstawie HR)
                # Wzór zwraca kJ/min. Zamieniamy to na Waty (J/s).
                # P_met [W] = (kJ/min * 1000) / 60
                
                # Współczynniki Keytela
                gender_factor = -55.0969 if is_male else -20.4022
                
                # Obliczenie wydatku energetycznego (EE) w kJ/min dla każdej sekundy
                # Używamy wygładzonego HR, żeby uniknąć skoków
                ee_kj_min = gender_factor + \
                            (0.6309 * df_plot_resampled['heartrate_smooth']) + \
                            (0.1988 * rider_weight) + \
                            (0.2017 * rider_age)
                
                # Konwersja na Waty Metaboliczne (P_met)
                # Uwaga: EE nie może być ujemne ani zerowe (serce bije)
                p_metabolic = (ee_kj_min * 1000) / 60
                p_metabolic = p_metabolic.replace(0, np.nan) # Unikamy dzielenia przez zero
                
                # 3. Obliczamy Gross Efficiency (GE)
                # GE = (Moc Mechaniczna / Moc Metaboliczna) * 100
                # Filtrujemy momenty, gdzie nie pedałujesz (Moc < 10W), bo wtedy GE=0
                
                ge_series = (df_plot_resampled['watts_smooth'] / p_metabolic) * 100
                
                # Filtrujemy dane nierealistyczne i "zimny start"
                # 1. Watts > 40 (żeby nie dzielić przez zero na postojach)
                # 2. GE między 5% a 30% (wszystko powyżej 30% to błąd pomiaru lub HR Lag)
                # 3. HR > 100 bpm (Wzór Keytela bardzo słabo działa dla niskiego tętna!)
                
                mask_ge = (df_plot_resampled['watts_smooth'] > 100) & \
                        (ge_series > 5) & (ge_series < 30) & \
                        (df_plot_resampled['heartrate_smooth'] > 110) 
                
                # Zerujemy błędne wartości (zamieniamy na NaN, żeby nie rysowały się na wykresie)
                df_ge = pd.DataFrame({
                    'time_min': df_plot_resampled['time_min'],
                    'ge': ge_series,
                    'watts': df_plot_resampled['watts_smooth']
                })
                df_ge.loc[~mask_ge, 'ge'] = np.nan
                
                # 4. Czyszczenie danych (Realistyczne ramy fizjologiczne)
                # GE rzadko przekracza 30% (chyba że zjeżdżasz z góry i HR spada szybciej niż waty)
                # GE poniżej 0% to błąd.
                mask_ge = (df_plot_resampled['watts_smooth'] > 40) & \
                        (ge_series > 5) & (ge_series < 35)
                
                df_ge = pd.DataFrame({
                    'time_min': df_plot_resampled['time_min'],
                    'ge': ge_series,
                    'watts': df_plot_resampled['watts_smooth']
                })
                # Zerujemy nierealistyczne wartości do wykresu
                df_ge.loc[~mask_ge, 'ge'] = np.nan

                if not df_ge['ge'].isna().all():
                    avg_ge = df_ge['ge'].mean()
                    
                    # KOLUMNY Z WYNIKAMI
                    cg1, cg2, cg3 = st.columns(3)
                    cg1.metric("Średnie GE", f"{avg_ge:.1f}%", help="Pro: 23%+, Amator: 18-21%")
                    
                    # Trend GE (czy spada w czasie?)
                    valid_ge = df_ge.dropna(subset=['ge'])
                    if len(valid_ge) > 100:
                        slope_ge, _, _, _, _ = stats.linregress(valid_ge['time_min'], valid_ge['ge'])
                        total_drift_ge = slope_ge * (valid_ge['time_min'].iloc[-1] - valid_ge['time_min'].iloc[0])
                        cg2.metric("Zmiana GE (Trend)", f"{total_drift_ge:.1f}%", delta_color="inverse" if total_drift_ge < 0 else "normal")
                    else:
                        cg2.metric("Zmiana GE", "-")

                    cg3.info("Wartości powyżej 25% mogą wynikać z opóźnienia tętna względem mocy (np. krótkie interwały). Analizuj trendy na długich odcinkach.")

                    # WYKRES GE
                    fig_ge = go.Figure()
                    
                    # Linia GE
                    fig_ge.add_trace(go.Scatter(
                        x=df_ge['time_min'], 
                        y=df_ge['ge'],
                        mode='lines',
                        name='Gross Efficiency (%)',
                        line=dict(color='#00cc96', width=1.5),
                        connectgaps=False, # Nie łączymy przerw (postojów)
                        hovertemplate="GE: %{y:.1f}%<extra></extra>"
                    ))
                    
                    # Tło (Moc)
                    fig_ge.add_trace(go.Scatter(
                        x=df_ge['time_min'], 
                        y=df_ge['watts'],
                        mode='lines',
                        name='Moc (Tło)',
                        yaxis='y2',
                        line=dict(color='rgba(255,255,255,0.1)', width=1),
                        fill='tozeroy',
                        fillcolor='rgba(255,255,255,0.05)',
                        hoverinfo='skip'
                    ))
                    
                    # Linia Trendu GE
                    if len(valid_ge) > 100:
                        trend_line = np.poly1d(np.polyfit(valid_ge['time_min'], valid_ge['ge'], 1))(valid_ge['time_min'])
                        fig_ge.add_trace(go.Scatter(
                            x=valid_ge['time_min'],
                            y=trend_line,
                            mode='lines',
                            name='Trend GE',
                            line=dict(color='white', width=2, dash='dash')
                        ))

                    fig_ge.update_layout(
                        template="plotly_dark",
                        title="Efektywność Brutto (GE%) w Czasie",
                        hovermode="x unified",
                        yaxis=dict(title="GE [%]", range=[10, 30]),
                        yaxis2=dict(title="Moc [W]", overlaying='y', side='right', showgrid=False),
                        height=400,
                        margin=dict(l=10, r=10, t=40, b=10),
                        legend=dict(orientation="h", y=1.1, x=0)
                    )
                    
                    st.plotly_chart(fig_ge, use_container_width=True)
                    
                    with st.expander("🧠 Jak interpretować GE?", expanded=False):
                        st.markdown("""
                        **Fizjologia GE:**
                        * **< 18%:** Niska wydajność. Dużo energii tracisz na ciepło i nieskoordynowane ruchy (kołysanie biodrami). Częste u początkujących.
                        * **19-21%:** Standard amatorski. Dobrze wytrenowany kolarz klubowy.
                        * **22-24%:** Poziom ELITE / PRO. Twoje mięśnie to maszyny.
                        * **> 25%:** Podejrzane (chyba że jesteś zwycięzcą Tour de France). Często wynika z błędów pomiaru (np. miernik mocy zawyża, tętno zaniżone, jazda w dół).

                        **Dlaczego GE spada w czasie?**
                        Gdy się męczysz, rekrutujesz włókna mięśniowe typu II (szybkokurczliwe), które są mniej wydajne tlenowo. Dodatkowo rośnie temperatura ciała (Core Temp), co kosztuje energię. Spadek GE pod koniec długiego treningu to doskonały wskaźnik zmęczenia metabolicznego.
                        """)
                else:
                    st.warning("Brak wystarczających danych do obliczenia GE (zbyt krótkie odcinki stabilnej jazdy).")
            else:
                st.error("Do obliczenia GE potrzebujesz danych Mocy (Watts) oraz Tętna (HR).")

        # --- TAB HRV ---
        with tab_hrv:
            st.header("Analiza Zmienności Rytmu Serca (HRV)")

            # 1. Inicjalizacja "Pamięci" (Session State)
            if 'df_dfa' not in st.session_state:
                st.session_state.df_dfa = None
            if 'dfa_error' not in st.session_state:
                st.session_state.dfa_error = None

            # 2. Obsługa Przycisku
            if st.session_state.df_dfa is None and st.session_state.dfa_error is None:
                st.info("💡 Analiza DFA Alpha-1 wymaga zaawansowanych obliczeń fraktalnych.")
                st.markdown("Kliknij przycisk poniżej, aby uruchomić algorytm. Może to zająć od kilku do kilkunastu sekund.")
                
                if st.button("🚀 Oblicz HRV i DFA Alpha-1"):
                    with st.spinner("Analiza geometrii rytmu serca... Proszę czekać..."):
                        try:
                            result_df, error_msg = calculate_dynamic_dfa(df_clean_pl)
                            
                            st.session_state.df_dfa = result_df
                            st.session_state.dfa_error = error_msg
                            
                            st.rerun()
                        except Exception as e:
                            st.error(f"Wystąpił błąd krytyczny algorytmu: {e}")

            # 3. Pobranie danych z pamięci do zmiennych lokalnych
            df_dfa = st.session_state.df_dfa
            dfa_error = st.session_state.dfa_error
           
            if df_dfa is not None and not df_dfa.empty:
                
                df_dfa = df_dfa.sort_values('time')
                orig_times = df_clean_pl['time'].values
                orig_watts = df_clean_pl['watts_smooth'].values if 'watts_smooth' in df_clean_pl.columns else np.zeros(len(orig_times))
                orig_hr = df_clean_pl['heartrate_smooth'].values if 'heartrate_smooth' in df_clean_pl.columns else np.zeros(len(orig_times))
                df_dfa['watts'] = np.interp(df_dfa['time'], orig_times, orig_watts)
                df_dfa['hr'] = np.interp(df_dfa['time'], orig_times, orig_hr)
                df_dfa['time_min'] = df_dfa['time'] / 60.0

                # Metryki podsumowujące
                col1, col2, col3, col4 = st.columns(4)
                col1.metric("Śr. RMSSD", f"{df_dfa['rmssd'].mean():.1f} ms" if 'rmssd' in df_dfa.columns else "N/A")
                col2.metric("Śr. SDNN", f"{df_dfa['sdnn'].mean():.1f} ms" if 'sdnn' in df_dfa.columns else "N/A")
                col3.metric("Śr. RR", f"{df_dfa['mean_rr'].mean():.0f} ms" if 'mean_rr' in df_dfa.columns else "N/A")
                col4.metric("Śr. HR (z RR)", f"{60000/df_dfa['mean_rr'].mean():.0f} bpm" if 'mean_rr' in df_dfa.columns else "N/A")

                st.subheader("Indeks Zmienności HRV (Pseudo-Alpha)")
                st.caption("Wyższe wartości = większa zmienność = lepszy stan regeneracji. Niższe = stres metaboliczny.")
                
                fig_dfa = go.Figure()
                fig_dfa.add_trace(go.Scatter(
                    x=df_dfa['time_min'], 
                    y=df_dfa['alpha1'],
                    name='Indeks HRV',
                    mode='lines',
                    line=dict(color='#00cc96', width=2),
                    hovertemplate="Indeks: %{y:.2f}<extra></extra>"
                ))

                fig_dfa.add_trace(go.Scatter(
                    x=df_dfa['time_min'], 
                    y=df_dfa['watts'],
                    name='Moc',
                    yaxis='y2',
                    fill='tozeroy',
                    line=dict(width=0.5, color='rgba(255,255,255,0.1)'),
                    hovertemplate="Moc: %{y:.0f} W<extra></extra>"
                ))

                fig_dfa.add_hline(y=0.75, line_dash="solid", line_color="#ef553b", line_width=2, 
                                annotation_text="Próg stresu (0.75)", annotation_position="top left")
                
                fig_dfa.add_hline(y=0.50, line_dash="dot", line_color="#ab63fa", line_width=1, 
                                annotation_text="Wysoki stres (0.50)", annotation_position="bottom left")

                fig_dfa.update_layout(
                    template="plotly_dark",
                    title="Indeks Zmienności HRV vs Czas",
                    hovermode="x unified",
                    xaxis=dict(title="Czas [min]"),
                    yaxis=dict(title="Indeks HRV", range=[0.2, 1.6]),
                    yaxis2=dict(title="Moc [W]", overlaying='y', side='right', showgrid=False),
                    height=500,
                    margin=dict(l=10, r=10, t=40, b=10),
                    legend=dict(orientation="h", y=1.05, x=0)
                )

                st.plotly_chart(fig_dfa, use_container_width=True)

                # Wykres RMSSD jeśli dostępny
                if 'rmssd' in df_dfa.columns:
                    st.subheader("RMSSD w czasie")
                    fig_rmssd = go.Figure()
                    fig_rmssd.add_trace(go.Scatter(
                        x=df_dfa['time_min'], 
                        y=df_dfa['rmssd'],
                        name='RMSSD',
                        mode='lines',
                        line=dict(color='#636efa', width=2),
                        hovertemplate="RMSSD: %{y:.1f} ms<extra></extra>"
                    ))
                    fig_rmssd.add_trace(go.Scatter(
                        x=df_dfa['time_min'], 
                        y=df_dfa['watts'],
                        name='Moc',
                        yaxis='y2',
                        fill='tozeroy',
                        line=dict(width=0.5, color='rgba(255,255,255,0.1)'),
                        hovertemplate="Moc: %{y:.0f} W<extra></extra>"
                    ))
                    fig_rmssd.update_layout(
                        template="plotly_dark",
                        title="RMSSD (Root Mean Square of Successive Differences)",
                        hovermode="x unified",
                        xaxis=dict(title="Czas [min]"),
                        yaxis=dict(title="RMSSD [ms]"),
                        yaxis2=dict(title="Moc [W]", overlaying='y', side='right', showgrid=False),
                        height=400,
                        margin=dict(l=10, r=10, t=40, b=10),
                        legend=dict(orientation="h", y=1.05, x=0)
                    )
                    st.plotly_chart(fig_rmssd, use_container_width=True)

                # --- WYKRES POINCARE (Lorenz Plot) ---
                st.markdown("---")
                st.subheader("Wykres Poincaré (Geometria Rytmu)")
                
                rr_values = df_dfa['mean_rr'].values 
                
                rr_col_raw = next((c for c in df_clean_pl.columns if any(x in c.lower() for x in ['rr', 'hrv', 'ibi', 'r-r'])), None)
                
                if rr_col_raw:
                    raw_rr_series = df_clean_pl[rr_col_raw].dropna().values
                    if raw_rr_series.mean() < 2.0: raw_rr_series = raw_rr_series * 1000
                    raw_rr_series = raw_rr_series[(raw_rr_series > 300) & (raw_rr_series < 2000)]
                    if len(raw_rr_series) > 10:
                        rr_n = raw_rr_series[:-1]
                        rr_n1 = raw_rr_series[1:]
                        
                        diff_rr = rr_n1 - rr_n
                        sd1 = np.std(diff_rr) / np.sqrt(2)
                        sd2 = np.sqrt(2 * np.std(raw_rr_series)**2 - 0.5 * np.std(diff_rr)**2)
                        ratio_sd = sd2 / sd1 if sd1 > 0 else 0

                        fig_poincare = go.Figure()
                        
                        fig_poincare.add_trace(go.Scatter(
                            x=rr_n, y=rr_n1,
                            mode='markers',
                            name='Interwały R-R',
                            marker=dict(
                                size=3,
                                color='rgba(0, 204, 150, 0.5)', 
                                line=dict(width=0)
                            ),
                            hovertemplate="RR(n): %{x:.0f} ms<br>RR(n+1): %{y:.0f} ms<extra></extra>"
                        ))
                        
                        min_rr, max_rr = min(raw_rr_series), max(raw_rr_series)
                        fig_poincare.add_trace(go.Scatter(
                            x=[min_rr, max_rr], y=[min_rr, max_rr],
                            mode='lines',
                            name='Linia tożsamości',
                            line=dict(color='white', width=1, dash='dash'),
                            hoverinfo='skip'
                        ))

                        fig_poincare.update_layout(
                            template="plotly_dark",
                            title=f"Poincaré Plot (SD1: {sd1:.1f}ms, SD2: {sd2:.1f}ms, Ratio: {ratio_sd:.2f})",
                            xaxis=dict(title="RR [n] (ms)", scaleanchor="y", scaleratio=1),
                            yaxis=dict(title="RR [n+1] (ms)"),
                            width=600, height=600, # Kwadratowy wykres
                            showlegend=False,
                            margin=dict(l=20, r=20, t=40, b=20)
                        )
                        
                        c_p1, c_p2 = st.columns([2, 1])
                        with c_p1:
                            st.plotly_chart(fig_poincare, use_container_width=True)
                        with c_p2:
                            st.info(f"""
                            **📊 Interpretacja Kliniczna:**
                            
                            * **Kształt "Komety" / "Rakiety":** Fizjologiczna norma u sportowca. Długa oś (SD2) to ogólna zmienność, krótka oś (SD1) to nagłe zmiany (parasympatyka).
                            * **Kształt "Kulisty":** Wysoki stres, dominacja współczulna (Fight or Flight) lub... bardzo równe tempo (metronom).
                            * **SD1 ({sd1:.1f} ms):** Czysta aktywność nerwu błędnego (regeneracja). Im więcej, tym lepiej.
                            * **SD2 ({sd2:.1f} ms):** Długoterminowa zmienność (rytm dobowy + termoregulacja).
                            
                            *Punkty daleko od głównej chmury to zazwyczaj ektopie (dodatkowe skurcze) lub błędy pomiaru.*
                            """)
                    else:
                        st.warning("Za mało danych R-R po filtracji artefaktów.")
                else:
                    st.warning("Brak surowych danych R-R do wygenerowania wykresu Poincaré.")    

                mask_threshold = (df_dfa['time_min'] > 5) & (df_dfa['alpha1'] < 0.75)
                
                if mask_threshold.any():
                    row = df_dfa[mask_threshold].iloc[0]
                    vt1_est_power = row['watts']
                    vt1_est_hr = row['hr']
                    vt1_time = row['time_min']
                    
                    c1, c2, c3 = st.columns(3)
                    c1.metric("Estymowane VT1 (Moc)", f"{vt1_est_power:.0f} W", help="Moc w momencie przecięcia linii 0.75")
                    c2.metric("Estymowane VT1 (HR)", f"{vt1_est_hr:.0f} bpm", help="Tętno w momencie przecięcia linii 0.75")
                    c3.metric("Czas przecięcia", f"{vt1_time:.0f} min")
                    
                    if vt1_est_power < 100:
                        st.warning("⚠️ Wykryto bardzo niskie VT1. Sprawdź jakość danych HRV (artefakty mogą zaniżać wynik).")
                else:
                    st.info("Nie przekroczono progu 0.75 w trakcie tego treningu (cały czas praca tlenowa lub krótkie dane).")

                # --- TEORIA ---
                with st.expander("🧠 O co chodzi z DFA Alpha-1?", expanded=True):
                    st.markdown(r"""
                    **Detrended Fluctuation Analysis ($\alpha_1$)** mierzy tzw. korelacje fraktalne w odstępach między uderzeniami serca.
                    
                    * **$\alpha_1 \approx 1.0$ (Szum Różowy):** Stan zdrowy, wypoczęty. Serce bije w sposób złożony, elastyczny. Organizuje się samo.
                    * **$\alpha_1 \approx 0.5$ (Szum Biały/Losowy):** Silny stres metaboliczny. Układ nerwowy "bombarduje" węzeł zatokowy, rytm staje się nieskorelowany.
                    
                    **Dlaczego 0.75?**
                    Badania (m.in. Rogers et al.) wykazały, że przejście przez wartość **0.75** idealnie pokrywa się z **Pierwszym Progiem Wentylacyjnym (VT1)**. Jest to punkt, w którym zaczynasz tracić "luz tlenowy", a organizm zaczyna rekrutować więcej włókien szybkokurczliwych.
                    """)

            else:
                # Debugowanie - pokaż dostępne kolumny
                hrv_cols = [c for c in df_clean_pl.columns if any(x in c.lower() for x in ['rr', 'hrv', 'ibi', 'r-r'])]
                if hrv_cols:
                    st.info(f"🔍 Znaleziono kolumny HRV: {hrv_cols}")
                    for col in hrv_cols:
                        col_data = df_clean_pl[col].dropna()
                        valid_count = (col_data > 0).sum()
                        st.write(f"  - {col}: {valid_count} wartości > 0, średnia: {col_data.mean():.2f}, zakres: {col_data.min():.2f} - {col_data.max():.2f}")
                else:
                    st.info(f"🔍 Dostępne kolumny: {list(df_clean_pl.columns)}")
                
                if dfa_error:
                    st.error(f"❌ Błąd DFA: {dfa_error}")
                
                st.warning("⚠️ **Brak wystarczających danych R-R (Inter-Beat Intervals).**")
                st.markdown("""
                Aby analiza DFA zadziałała, plik musi zawierać surowe dane o każdym uderzeniu serca, a nie tylko uśrednione tętno.
                * Sprawdź, czy Twój pas HR obsługuje HRV (np. Polar H10, Garmin HRM-Pro).
                * Upewnij się, że włączyłeś zapis zmienności tętna w zegarku/komputerze (często opcja "Log HRV").
                * Wymagane jest minimum 300 próbek z interwałami R-R > 0.
                """)
            
            st.divider()
            
            c1, c2 = st.columns(2)
            
            # LEWA KOLUMNA: SmO2 + TREND
            with c1:
                st.subheader("SmO2")
                # Szukamy odpowiedniej kolumny
                col_smo2 = 'smo2_smooth_ultra' if 'smo2_smooth_ultra' in df_plot.columns else ('smo2_smooth' if 'smo2_smooth' in df_plot.columns else None)
                
                if col_smo2:
                    fig_s = go.Figure()
                    
                    # 1. SmO2 (Linia)
                    fig_s.add_trace(go.Scatter(
                        x=df_plot_resampled['time_min'], 
                        y=df_plot_resampled[col_smo2], 
                        name='SmO2', 
                        line=dict(color='#ab63fa', width=2), 
                        hovertemplate="SmO2: %{y:.1f}%<extra></extra>"
                    ))
                    
                    # 2. Trend (Linia przerywana)
                    trend_y = calculate_trend(df_plot_resampled['time_min'].values, df_plot_resampled[col_smo2].values)
                    if trend_y is not None:
                        fig_s.add_trace(go.Scatter(
                            x=df_plot_resampled['time_min'], 
                            y=trend_y, 
                            name='Trend', 
                            line=dict(color='white', dash='dash', width=1.5), 
                            hovertemplate="Trend: %{y:.1f}%<extra></extra>"
                        ))
                    
                    # Layout "Pro"
                    fig_s.update_layout(
                        template="plotly_dark",
                        title="Lokalna Oksydacja (SmO2)",
                        hovermode="x unified", # <--- To robi robotę
                        yaxis=dict(title="SmO2 [%]", range=[0, 100]), # Sztywna skala dla czytelności
                        legend=dict(orientation="h", y=1.1, x=0),
                        margin=dict(l=10, r=10, t=40, b=10),
                        height=400
                    )
                    
                    st.plotly_chart(fig_s, use_container_width=True)
                    
                    st.info("""
                    **💡 Hemodynamika Mięśniowa (SmO2) - Lokalny Monitoring:**
                    
                    SmO2 to "wskaźnik paliwa" bezpośrednio w pracującym mięśniu (zazwyczaj czworogłowym uda).
                    * **Równowaga (Linia Płaska):** Podaż tlenu = Zapotrzebowanie. To stan zrównoważony (Steady State).
                    * **Desaturacja (Spadek):** Popyt > Podaż. Wchodzisz w dług tlenowy. Jeśli dzieje się to przy stałej mocy -> zmęczenie metaboliczne.
                    * **Reoksygenacja (Wzrost):** Odpoczynek. Szybkość powrotu do normy to doskonały wskaźnik wytrenowania (regeneracji).
                    """)
                else:
                     st.info("Brak danych SmO2")

            # PRAWA KOLUMNA: TĘTNO (HR)
            with c2:
                st.subheader("Tętno")
                
                # Przepisane na go.Figure dla spójności stylu z resztą aplikacji
                fig_h = go.Figure()
                fig_h.add_trace(go.Scatter(
                    x=df_plot_resampled['time_min'], 
                    y=df_plot_resampled['heartrate_smooth'], 
                    name='HR', 
                    fill='tozeroy', # Ładne wypełnienie pod wykresem
                    line=dict(color='#ef553b', width=2), 
                    hovertemplate="HR: %{y:.0f} BPM<extra></extra>"
                ))
                
                fig_h.update_layout(
                    template="plotly_dark",
                    title="Odpowiedź Sercowa (HR)",
                    hovermode="x unified", # <--- To robi robotę
                    yaxis=dict(title="HR [bpm]"),
                    margin=dict(l=10, r=10, t=40, b=10),
                    height=400
                )
                
                st.plotly_chart(fig_h, use_container_width=True)
                
                st.info("""
                **💡 Reakcja Sercowo-Naczyniowa (HR) - Globalny System:**
                
                Serce to pompa centralna. Jego reakcja jest **opóźniona** względem wysiłku.
                * **Lag (Opóźnienie):** W krótkich interwałach (np. 30s) tętno nie zdąży wzrosnąć, mimo że moc jest max. Nie steruj sprintami na tętno!
                * **Decoupling (Rozjazd):** Jeśli moc jest stała, a tętno rośnie (dryfuje) -> organizm walczy z przegrzaniem lub odwodnieniem.
                * **Recovery HR:** Jak szybko tętno spada po wysiłku? Szybki spadek = sprawne przywspółczulne układu nerwowego (dobra forma).
                """)

            st.divider()

            st.subheader("Wentylacja (VE) i Oddechy (RR)")
            
            fig_v = go.Figure()
            
            # 1. WENTYLACJA (Oś Lewa)
            if 'tymeventilation_smooth' in df_plot_resampled.columns:
                fig_v.add_trace(go.Scatter(
                    x=df_plot_resampled['time_min'], 
                    y=df_plot_resampled['tymeventilation_smooth'], 
                    name="VE", 
                    line=dict(color='#ffa15a', width=2), 
                    hovertemplate="VE: %{y:.1f} L/min<extra></extra>"
                ))
                
                # Trend VE
                trend_ve = calculate_trend(df_plot_resampled['time_min'].values, df_plot_resampled['tymeventilation_smooth'].values)
                if trend_ve is not None:
                     fig_v.add_trace(go.Scatter(
                         x=df_plot_resampled['time_min'], 
                         y=trend_ve, 
                         name="Trend VE", 
                         line=dict(color='#ffa15a', dash='dash', width=1.5), 
                         hovertemplate="Trend: %{y:.1f} L/min<extra></extra>"
                     ))
            
            # 2. ODDECHY / RR (Oś Prawa)
            if 'tymebreathrate_smooth' in df_plot_resampled.columns:
                fig_v.add_trace(go.Scatter(
                    x=df_plot_resampled['time_min'], 
                    y=df_plot_resampled['tymebreathrate_smooth'], 
                    name="RR", 
                    yaxis="y2", # Druga oś
                    line=dict(color='#19d3f3', dash='dot', width=2), 
                    hovertemplate="RR: %{y:.1f} /min<extra></extra>"
                ))
            
            # Linie Progi Wentylacyjne (Zostawiamy jako stałe linie odniesienia)
            fig_v.add_hline(y=vt1_vent, line_dash="dot", line_color="green", annotation_text="VT1", annotation_position="bottom right")
            fig_v.add_hline(y=vt2_vent, line_dash="dot", line_color="red", annotation_text="VT2", annotation_position="bottom right")

            # LAYOUT (Unified Hover)
            fig_v.update_layout(
                template="plotly_dark",
                title="Mechanika Oddechu (Wydajność vs Częstość)",
                hovermode="x unified", # <--- To łączy dane w jeden dymek
                
                # Oś Lewa
                yaxis=dict(title="Wentylacja [L/min]"),
                
                # Oś Prawa
                yaxis2=dict(
                    title="Kadencja Oddechu [RR]", 
                    overlaying="y", 
                    side="right", 
                    showgrid=False
                ),
                
                legend=dict(orientation="h", y=1.1, x=0),
                margin=dict(l=10, r=10, t=40, b=10),
                height=450
            )
            
            st.plotly_chart(fig_v, use_container_width=True)
            
            st.info("""
            **💡 Interpretacja: Mechanika Oddychania**

            * **Wzorzec Prawidłowy (Efektywność):** Wentylacja (VE) rośnie liniowo wraz z mocą, a częstość (RR) jest stabilna. Oznacza to głęboki, spokojny oddech.
            * **Wzorzec Niekorzystny (Płytki Oddech):** Bardzo wysokie RR (>40-50) przy stosunkowo niskim VE. Oznacza to "dyszenie" - powietrze wchodzi tylko do "martwej strefy" płuc, nie biorąc udziału w wymianie gazowej.
            * **Dryf Wentylacyjny:** Jeśli przy stałej mocy VE ciągle rośnie (rosnący trend pomarańczowej linii), oznacza to narastającą kwasicę (organizm próbuje wydmuchać CO2) lub zmęczenie mięśni oddechowych.
            * **Próg VT2 (RCP):** Punkt załamania, gdzie VE wystrzeliwuje pionowo w górę. To Twoja "czerwona linia" metaboliczna.
            """)
            
            col_vent_full = 'tymeventilation_smooth' if 'tymeventilation_smooth' in df_plot.columns else ('tymeventilation' if 'tymeventilation' in df_plot.columns else None)
            
            if col_vent_full:
                st.markdown("#### Czas w Strefach Wentylacyjnych")
                total_samples = len(df_plot)
                z1_count = len(df_plot[df_plot[col_vent_full] < vt1_vent])
                z2_count = len(df_plot[(df_plot[col_vent_full] >= vt1_vent) & (df_plot[col_vent_full] < vt2_vent)])
                z3_count = len(df_plot[df_plot[col_vent_full] >= vt2_vent])
                
                def format_time(seconds):
                    m, s = divmod(seconds, 60)
                    h, m = divmod(m, 60)
                    if h > 0: return f"{int(h)}h {int(m)}m {int(s)}s"
                    return f"{int(m)}m {int(s)}s"

                z1_time = format_time(z1_count)
                z2_time = format_time(z2_count)
                z3_time = format_time(z3_count)
                
                z1_pct = z1_count / total_samples * 100 if total_samples > 0 else 0
                z2_pct = z2_count / total_samples * 100 if total_samples > 0 else 0
                z3_pct = z3_count / total_samples * 100 if total_samples > 0 else 0
                
                c_z1, c_z2, c_z3 = st.columns(3)
                c_z1.metric(f"Tlenowa (< {vt1_vent} L)", z1_time, f"{z1_pct:.1f}%")
                c_z2.metric(f"Mieszana ({vt1_vent}-{vt2_vent} L)", z2_time, f"{z2_pct:.1f}%")
                c_z3.metric(f"Beztlenowa (> {vt2_vent} L)", z3_time, f"{z3_pct:.1f}%")

            if 'tymeventilation' in df_plot.columns:
                st.markdown("#### Średnie Wartości (10 min)")
                df_s = df_plot.copy()
                df_s['Int'] = (df_s['time_min'] // 10).astype(int)
                grp = df_s.groupby('Int')[['tymeventilation', 'tymebreathrate']].mean().reset_index()
                grp['Czas'] = grp['Int'].apply(lambda x: f"{x*10}-{(x+1)*10} min")
                st.dataframe(grp[['Czas', 'tymeventilation', 'tymebreathrate']].style.format("{:.1f}", subset=['tymeventilation', 'tymebreathrate']), use_container_width=True, hide_index=True)

        # --- TAB BIOMECH ---
        with tab_biomech:
            st.header("Biomechaniczny Stres")
            
            if 'torque_smooth' in df_plot_resampled.columns:
                fig_b = go.Figure()
                
                # 1. MOMENT OBROTOWY (Oś Lewa)
                # Kolor różowy/magenta - symbolizuje napięcie/siłę
                fig_b.add_trace(go.Scatter(
                    x=df_plot_resampled['time_min'], 
                    y=df_plot_resampled['torque_smooth'], 
                    name='Moment (Torque)', 
                    line=dict(color='#e377c2', width=1.5), 
                    hovertemplate="Moment: %{y:.1f} Nm<extra></extra>"
                ))
                
                # 2. KADENCJA (Oś Prawa)
                # Kolor cyan/turkus - symbolizuje szybkość/obroty
                if 'cadence_smooth' in df_plot_resampled.columns:
                    fig_b.add_trace(go.Scatter(
                        x=df_plot_resampled['time_min'], 
                        y=df_plot_resampled['cadence_smooth'], 
                        name='Kadencja', 
                        yaxis="y2", # Druga oś
                        line=dict(color='#19d3f3', width=1.5), 
                        hovertemplate="Kadencja: %{y:.0f} RPM<extra></extra>"
                    ))
                
                # LAYOUT (Unified Hover)
                fig_b.update_layout(
                    template="plotly_dark",
                    title="Analiza Generowania Mocy (Siła vs Szybkość)",
                    hovermode="x unified", # <--- Klucz do sukcesu
                    
                    # Oś Lewa
                    yaxis=dict(title="Moment [Nm]"),
                    
                    # Oś Prawa
                    yaxis2=dict(
                        title="Kadencja [RPM]", 
                        overlaying="y", 
                        side="right", 
                        showgrid=False
                    ),
                    
                    legend=dict(orientation="h", y=1.1, x=0),
                    margin=dict(l=10, r=10, t=40, b=10),
                    height=450
                )
                
                st.plotly_chart(fig_b, use_container_width=True)
                
                # --- ZMIANA: ROZBUDOWANE KOMPENDIUM BIOMECH ---
                st.info("""
                **💡 Kompendium: Moment Obrotowy (Siła) vs Kadencja (Szybkość)**

                Wykres pokazuje, w jaki sposób generujesz moc.
                Pamiętaj: `Moc = Moment x Kadencja`. Tę samą moc (np. 200W) możesz uzyskać "siłowo" (50 RPM) lub "szybkościowo" (100 RPM).

                **1. Interpretacja Stylu Jazdy:**
                * **Grinding (Niska Kadencja < 70, Wysoki Moment):**
                    * **Fizjologia:** Dominacja włókien szybkokurczliwych (beztlenowych). Szybkie zużycie glikogenu.
                    * **Skutek:** "Betonowe nogi" na biegu.
                    * **Ryzyko:** Przeciążenie stawu rzepkowo-udowego (ból kolan) i odcinka lędźwiowego.
                * **Spinning (Wysoka Kadencja > 90, Niski Moment):**
                    * **Fizjologia:** Przeniesienie obciążenia na układ krążenia (serce i płuca). Lepsze ukrwienie mięśni (pompa mięśniowa).
                    * **Skutek:** Świeższe nogi do biegu (T2).
                    * **Wyzwanie:** Wymaga dobrej koordynacji nerwowo-mięśniowej (żeby nie podskakiwać na siodełku).

                **2. Praktyczne Przykłady (Kiedy co stosować?):**
                * **Podjazd:** Naturalna tendencja do spadku kadencji. **Błąd:** "Przepychanie" na twardym biegu. **Korekta:** Zredukuj bieg, utrzymaj 80+ RPM, nawet jeśli prędkość spadnie. Oszczędzisz mięśnie.
                * **Płaski odcinek (TT):** Utrzymuj "Sweet Spot" kadencji (zazwyczaj 85-95 RPM). To balans między zmęczeniem mięśniowym a sercowym.
                * **Finisz / Atak:** Chwilowe wejście w wysoki moment I wysoką kadencję. Kosztowne energetycznie, ale daje max prędkość.

                **3. Możliwe Komplikacje i Sygnały Ostrzegawcze:**
                * **Ból przodu kolana:** Zbyt duży moment obrotowy (za twarde przełożenia). -> Zwiększ kadencję.
                * **Ból bioder / "skakanie":** Zbyt wysoka kadencja przy słabej stabilizacji (core). -> Wzmocnij brzuch lub nieco zwolnij obroty.
                * **Drętwienie stóp:** Często wynik ciągłego nacisku przy niskiej kadencji. Wyższa kadencja poprawia krążenie (faza luzu w obrocie).
                """)
            
            st.divider()
            st.subheader("Wpływ Momentu na Oksydację (Torque vs SmO2)")
            
            if 'torque' in df_plot.columns and 'smo2' in df_plot.columns:
                # Przygotowanie danych (Binning)
                df_bins = df_plot.copy()
                # Grupujemy moment co 2 Nm
                df_bins['Torque_Bin'] = (df_bins['torque'] // 2 * 2).astype(int)
                
                # Liczymy statystyki dla każdego koszyka
                bin_stats = df_bins.groupby('Torque_Bin')['smo2'].agg(['mean', 'std', 'count']).reset_index()
                # Filtrujemy szum (musi być min. 10 próbek dla danej siły)
                bin_stats = bin_stats[bin_stats['count'] > 10]
                
                fig_ts = go.Figure()
                
                # 1. GÓRNA GRANICA (Mean + STD) - Niewidoczna linia, potrzebna do cieniowania
                fig_ts.add_trace(go.Scatter(
                    x=bin_stats['Torque_Bin'], 
                    y=bin_stats['mean'] + bin_stats['std'], 
                    mode='lines', 
                    line=dict(width=0), 
                    showlegend=False, 
                    name='Górny zakres (+1SD)',
                    hovertemplate="Max (zakres): %{y:.1f}%<extra></extra>"
                ))
                
                # 2. DOLNA GRANICA (Mean - STD) - Wypełnienie
                fig_ts.add_trace(go.Scatter(
                    x=bin_stats['Torque_Bin'], 
                    y=bin_stats['mean'] - bin_stats['std'], 
                    mode='lines', 
                    line=dict(width=0), 
                    fill='tonexty', # Wypełnia do poprzedniej ścieżki (Górnej granicy)
                    fillcolor='rgba(255, 75, 75, 0.15)', # Lekka czerwień
                    showlegend=False, 
                    name='Dolny zakres (-1SD)',
                    hovertemplate="Min (zakres): %{y:.1f}%<extra></extra>"
                ))
                
                # 3. ŚREDNIA (Główna Linia)
                fig_ts.add_trace(go.Scatter(
                    x=bin_stats['Torque_Bin'], 
                    y=bin_stats['mean'], 
                    mode='lines+markers', 
                    name='Średnie SmO2', 
                    line=dict(color='#FF4B4B', width=3), 
                    marker=dict(size=6, color='#FF4B4B', line=dict(width=1, color='white')),
                    hovertemplate="<b>Śr. SmO2:</b> %{y:.1f}%<extra></extra>"
                ))
                
                # LAYOUT (Unified Hover)
                fig_ts.update_layout(
                    template="plotly_dark",
                    title="Agregacja: Jak Siła (Moment) wpływa na Tlen (SmO2)?",
                    hovermode="x unified", # <--- Skanujemy w pionie dla konkretnej wartości Nm
                    xaxis=dict(title="Moment Obrotowy [Nm]"),
                    yaxis=dict(title="SmO2 [%]"),
                    legend=dict(orientation="h", y=1.1, x=0),
                    margin=dict(l=10, r=10, t=40, b=10),
                    height=450
                )
                
                st.plotly_chart(fig_ts, use_container_width=True)
                
                st.info("""
                **💡 Fizjologia Okluzji (Analiza Koszykowa):**
                
                **Mechanizm Okluzji:** Kiedy mocno napinasz mięsień (wysoki moment), ciśnienie wewnątrzmięśniowe przewyższa ciśnienie w naczyniach włosowatych. Krew przestaje płynąć, tlen nie dociera, a metabolity (kwas mlekowy) nie są usuwane. To "duszenie" mięśnia od środka.
                
                **Punkt Krytyczny:** Szukaj momentu (na osi X), gdzie czerwona linia gwałtownie opada w dół. To Twój limit siłowy. Powyżej tej wartości generujesz waty 'na kredyt' beztlenowy.
                
                **Praktyczny Wniosek (Scenario):** * Masz do wygenerowania 300W. Możesz to zrobić siłowo (70 RPM, wysoki moment) lub kadencyjnie (90 RPM, niższy moment).
                * Spójrz na wykres: Jeśli przy momencie odpowiadającym 70 RPM Twoje SmO2 spada do 30%, a przy momencie dla 90 RPM wynosi 50% -> **Wybierz wyższą kadencję!** Oszczędzasz nogi (glikogen) kosztem nieco wyższego tętna.
                """)

        # --- TAB THERMAL ---
        with tab_thermal:
            st.header("Wydajność Chłodzenia")
            
            fig_t = go.Figure()
            
            # 1. CORE TEMP (Oś Lewa)
            # Kolor pomarańczowy - symbolizuje ciepło
            if 'core_temperature_smooth' in df_plot.columns:
                fig_t.add_trace(go.Scatter(
                    x=df_plot['time_min'], 
                    y=df_plot['core_temperature_smooth'], 
                    name='Core Temp', 
                    line=dict(color='#ff7f0e', width=2), 
                    hovertemplate="Temp: %{y:.2f}°C<extra></extra>"
                ))
            
            # 2. HSI - HEAT STRAIN INDEX (Oś Prawa)
            # Kolor czerwony przerywany - symbolizuje ryzyko/alarm
            if 'hsi' in df_plot.columns:
                fig_t.add_trace(go.Scatter(
                    x=df_plot['time_min'], 
                    y=df_plot['hsi'], 
                    name='HSI', 
                    yaxis="y2", # Druga oś
                    line=dict(color='#d62728', width=2, dash='dot'), 
                    hovertemplate="HSI: %{y:.1f}<extra></extra>"
                ))
            
            # Linie referencyjne dla temperatury (Strefy)
            fig_t.add_hline(y=38.5, line_dash="dash", line_color="red", opacity=0.5, annotation_text="Krytyczna (38.5°C)", annotation_position="top left")
            fig_t.add_hline(y=37.5, line_dash="dot", line_color="green", opacity=0.5, annotation_text="Optymalna (37.5°C)", annotation_position="bottom left")

            # LAYOUT (Unified Hover)
            fig_t.update_layout(
                template="plotly_dark",
                title="Termoregulacja: Temperatura Głęboka vs Indeks Zmęczenia (HSI)",
                hovermode="x unified", # <--- Skanujemy obie wartości na raz
                
                # Oś Lewa
                yaxis=dict(title="Core Temp [°C]"),
                
                # Oś Prawa
                yaxis2=dict(
                    title="HSI [0-10]", 
                    overlaying="y", 
                    side="right", 
                    showgrid=False,
                    range=[0, 12] # Lekki zapas na skali, żeby wykres nie dotykał sufitu
                ),
                
                legend=dict(orientation="h", y=1.1, x=0),
                margin=dict(l=10, r=10, t=40, b=10),
                height=450
            )
            
            st.plotly_chart(fig_t, use_container_width=True)
            
            st.info("""
            **🌡️ Kompendium Termoregulacji: Fizjologia i Strategia**

            **1. Fizjologiczny Koszt Ciepła (Konkurencja o Krew)**
            Twój układ krążenia to system zamknięty o ograniczonej pojemności (ok. 5L krwi). Podczas wysiłku w upale serce musi obsłużyć dwa konkurencyjne cele:
            * **Mięśnie:** Dostarczenie tlenu i paliwa (priorytet wysiłkowy).
            * **Skóra:** Oddanie ciepła przez pot i konwekcję (priorytet przeżycia).
            * **Efekt:** Mniej krwi trafia do mięśni -> Spadek VO2max -> Wzrost tętna przy tej samej mocy (Cardiac Drift). Dodatkowo, utrata osocza (pot) zagęszcza krew, zmuszając serce do cięższej pracy.

            **2. Strefy Temperaturowe (Core Temp):**
            * **36.5°C - 37.5°C:** Homeostaza. Strefa komfortu i rozgrzewki.
            * **37.5°C - 38.4°C:** **Strefa Wydajności.** Optymalna temperatura pracy mięśni (enzymy działają najszybciej). Tutaj chcesz być podczas wyścigu.
            * **> 38.5°C:** **Strefa Krytyczna ("The Meltdown").** Ośrodkowy Układ Nerwowy (mózg) zaczyna "zaciągać hamulec ręczny", redukując rekrutację jednostek motorycznych, by chronić organy przed ugotowaniem. Odczuwasz to jako nagły brak mocy ("odcięcie").

            **3. HSI (Heat Strain Index 0-10):**
            * **0-3 (Niski):** Pełen komfort. Możesz cisnąć maxa.
            * **4-6 (Umiarkowany):** Fizjologiczny koszt rośnie. Wymagane nawadnianie.
            * **7-9 (Wysoki):** Znaczący spadek wydajności. Skup się na chłodzeniu, nie na watach.
            * **10 (Ekstremalny):** Ryzyko udaru. Zwolnij natychmiast.

            **4. Protokół Chłodzenia (Strategia):**
            * **Internal (Wewnętrzne):** Pij zimne napoje (tzw. ice slurry). Obniża to temp. żołądka i core temp.
            * **External (Zewnętrzne):** Polewaj wodą głowę, kark i **nadgarstki** (duże naczynia krwionośne blisko skóry). Lód w stroju startowym (na karku/klatce) to game-changer.

            **5. Czerwone Flagi (Kiedy przerwać):**
            * Gęsia skórka lub dreszcze w upale (paradoksalna reakcja - mózg "wariuje").
            * Nagły spadek tętna przy utrzymaniu wysiłku.
            * Zaburzenia widzenia lub koordynacji.
            """)

            st.header("Koszt Termiczny Wydajności (Cardiac Drift)")
            
            # Sprawdzamy czy mamy potrzebne kolumny
            temp_col = 'core_temperature_smooth' if 'core_temperature_smooth' in df_plot.columns else 'core_temperature'
            
            if 'watts' in df_plot.columns and temp_col in df_plot.columns and 'heartrate' in df_plot.columns:
                
                # 1. FILTROWANIE DANYCH
                # Wywalamy zera i postoje
                mask = (df_plot['watts'] > 10) & (df_plot['heartrate'] > 60)
                df_clean = df_plot[mask].copy()
                
                # 2. OBLICZENIE EFEKTYWNOŚCI (EF)
                df_clean['eff_raw'] = df_clean['watts'] / df_clean['heartrate']
                
                # 3. USUWANIE OUTLIERÓW
                df_clean = df_clean[df_clean['eff_raw'] < 6.0]

                if not df_clean.empty:
                    # Tworzymy wykres z linią trendu (Lowess - lokalna regresja)
                    fig_te = px.scatter(
                        df_clean, 
                        x=temp_col, 
                        y='eff_raw', 
                        trendline="lowess", 
                        trendline_options=dict(frac=0.3), 
                        trendline_color_override="#FF4B4B", 
                        template="plotly_dark",
                        opacity=0.3 # Przezroczyste punkty, żeby widzieć gęstość
                    )
                    
                    # Formatowanie punktów (Scatter)
                    fig_te.update_traces(
                        selector=dict(mode='markers'),
                        marker=dict(size=5, color='#1f77b4'),
                        hovertemplate="<b>Temp:</b> %{x:.2f}°C<br><b>EF:</b> %{y:.2f} W/bpm<extra></extra>"
                    )
                    
                    # Formatowanie linii trendu
                    fig_te.update_traces(
                        selector=dict(mode='lines'),
                        line=dict(width=4),
                        hovertemplate="<b>Trend:</b> %{y:.2f} W/bpm<extra></extra>"
                    )
                    
                    # LAYOUT (Unified Hover)
                    fig_te.update_layout(
                        title="Spadek Efektywności (W/HR) vs Temperatura",
                        hovermode="x unified", # <--- To robi robotę
                        
                        xaxis=dict(title="Temperatura Głęboka [°C]"),
                        yaxis=dict(title="Efficiency Factor [W/bpm]"),
                        
                        showlegend=False,
                        margin=dict(l=10, r=10, t=40, b=10),
                        height=450
                    )

                    st.plotly_chart(fig_te, use_container_width=True, config={'scrollZoom': False}, key="thermal_eff")
                    
                    st.info("""
                    ℹ️ **Jak to czytać?**
                    Ten wykres pokazuje **Cardiac Drift** w funkcji temperatury.
                    * **Oś Y (W/HR):** Ile watów generujesz z jednego uderzenia serca. Wyższa wartość = lepsza efektywność.
                    * **Oś X (Core Temp):** Twoja temperatura wewnętrzna. Wyższa wartość = większy stres cieplny.
                    * **Trend spadkowy:** Oznacza, że wraz ze wzrostem temperatury Twoje serce musi bić szybciej dla tej samej mocy (krew idzie do skóry na chłodzenie = mniejszy rzut serca dla mięśni).
                    * **Filtracja:** Usunąłem momenty, gdy nie pedałujesz (Moc < 10W), żeby nie zaburzać wyniku.
                    """)
                else:
                    st.warning("Zbyt mało danych po przefiltrowaniu (sprawdź czy masz odczyty mocy i tętna).")
            else:
                st.error("Brak wymaganych kolumn (watts, heartrate, core_temperature).")
                
                st.info("""
                **💡 Interpretacja: Koszt Fizjologiczny Ciepła (Decoupling Termiczny)**

                Ten wykres pokazuje, jak Twoje "serce płaci" za każdy wat mocy w miarę wzrostu temperatury ciała.
                * **Oś X:** Temperatura Centralna (Core Temp).
                * **Oś Y:** Efektywność (Waty na 1 uderzenie serca).
                * **Czerwona Linia:** Trend zmian.

                **🔍 Scenariusze:**
                1.  **Linia Płaska (Idealnie):** Twoja termoregulacja działa świetnie. Mimo wzrostu temperatury, serce pracuje tak samo wydajnie. Jesteś dobrze nawodniony i zaadaptowany do ciepła.
                2.  **Linia Opadająca (Typowe):** Wraz ze wzrostem temp. serce musi bić szybciej, by utrzymać tę samą moc (Dryf). Krew ucieka do skóry, by Cię chłodzić, zamiast napędzać mięśnie.
                3.  **Gwałtowny Spadek:** "Zawał termiczny" wydajności. Zazwyczaj powyżej 38.5°C. W tym momencie walczysz o przetrwanie, a nie o wynik.

                **Wniosek:** Jeśli linia leci mocno w dół, musisz poprawić chłodzenie (polewanie wodą, lód) lub strategię nawadniania przed startem.
                """)

        # --- TAB TRENDS ---
        with tab_trends:
            st.header("Trendy")
            
            if 'watts_smooth' in df_plot.columns and 'heartrate_smooth' in df_plot.columns:
                # Przygotowanie danych do ścieżki (Rolling Average 5 min)
                df_trend = df_plot.copy()
                df_trend['w_trend'] = df_trend['watts'].rolling(window=300, min_periods=60).mean()
                df_trend['hr_trend'] = df_trend['heartrate'].rolling(window=300, min_periods=60).mean()
                
                # Próbkowanie co 60 wierszy (co minutę), żeby nie zamulić wykresu tysiącami kropek
                df_path = df_trend.iloc[::60, :]
                
                fig_d = go.Figure()
                
                fig_d.add_trace(go.Scatter(
                    x=df_path['w_trend'], 
                    y=df_path['hr_trend'], 
                    mode='markers+lines', 
                    name='Ścieżka',
                    # Kolorowanie wg czasu (Gradient)
                    marker=dict(
                        size=8, 
                        color=df_path['time_min'], 
                        colorscale='Viridis', 
                        showscale=True, 
                        colorbar=dict(title="Czas [min]"),
                        line=dict(width=1, color='white')
                    ),
                    line=dict(color='rgba(255,255,255,0.3)', width=1), # Cienka linia łącząca
                    
                    # Bogaty Tooltip (Stylizowany jak w innych zakładkach)
                    hovertemplate="<b>Czas: %{marker.color:.0f} min</b><br>" +
                                  "Moc (5min): %{x:.0f} W<br>" +
                                  "HR (5min): %{y:.0f} BPM<extra></extra>"
                ))
                
                fig_d.update_layout(
                    template="plotly_dark",
                    title="Ścieżka Dryfu: Relacja Moc vs Tętno w Czasie",
                    
                    # Tutaj używamy 'closest', bo oś X to Moc, a nie czas. 
                    # 'x unified' zrobiłoby bałagan pokazując wszystkie momenty z tą samą mocą na raz.
                    hovermode="closest", 
                    
                    xaxis=dict(title="Moc (Średnia 5 min) [W]"),
                    yaxis=dict(title="Tętno (Średnia 5 min) [BPM]"),
                    margin=dict(l=10, r=10, t=40, b=10),
                    height=500
                )
                
                st.plotly_chart(fig_d, use_container_width=True)
                
                st.info("""
                **💡 Interpretacja Ścieżki:**
                * **Pionowo w górę:** Czysty dryf tętna (rosnące zmęczenie przy stałej mocy). Związane jest to z odwodnieniem lub nagromadzeniem ciepła. Zazwyczaj obserwowane w długotrwałych wysiłkach (>60 min) w ciepłych warunkach. Protip: nawadniaj się regularnie i stosuj chłodzenie.
                * **Poziomo w prawo:** Zwiększenie mocy bez wzrostu tętna. Oznacza poprawę efektywności (np. zjazd, lepsza aerodynamika, wiatr w plecy).
                * **Poziomo w lewo:** Spadek mocy przy stałym tętnie. Może wskazywać na zmęczenie mięśniowe lub pogorszenie warunków (podjazd pod wiatr).
                * **W lewo i w dół:** Niekorzystna reakcja organizmu (spadek mocy i tętna) - możliwe początki wyczerpania energetycznego lub przegrzania.
                * **W prawo i w górę:** Zdrowa reakcja na zwiększenie intensywności. Twoje ciało efektywnie dostosowuje się do rosnącego wysiłku. Oznaka odpowiedniego poziomu wytrenowania.
                """)

            st.divider()
            st.subheader("Analiza Kwadrantowa 3D")
            if 'torque' in df_plot.columns and 'cadence' in df_plot.columns and 'watts' in df_plot.columns:
                df_q = df_plot.sample(min(len(df_plot), 5000))
                color_col = 'smo2_smooth' if 'smo2_smooth' in df_q.columns else 'watts'
                title_col = 'SmO2' if 'smo2_smooth' in df_q.columns else 'Moc'
                scale = 'Spectral' if 'smo2_smooth' in df_q.columns else 'Viridis'
                
                fig_3d = px.scatter_3d(df_q, x='cadence', y='torque', z='watts', color=color_col, title=f"3D Quadrant Analysis (Kolor: {title_col})", labels={'cadence': 'Kadencja', 'torque': 'Moment', 'watts': 'Moc'}, color_continuous_scale=scale, template='plotly_dark')
                fig_3d.update_traces(marker=dict(size=3, opacity=0.6), hovertemplate="Kadencja: %{x:.0f}<br>Moment: %{y:.1f}<br>Moc: %{z:.0f}<br>Val: %{marker.color:.1f}<extra></extra>")
                # W 3D używamy wbudowanego w px template, więc tylko update layout dla wysokości
                fig_3d.update_layout(height=700) 
                st.plotly_chart(fig_3d, use_container_width=True)
                
                st.info("""
                **💡 Jak czytać ten wykres 3D? (Instrukcja i Przykłady)**

                Ten wykres to "mapa Twojego silnika". Każdy punkt to jedna sekunda jazdy.
                * **Oś X (Kadencja):** Szybkość obrotu korbą.
                * **Oś Y (Moment):** Siła nacisku na pedał.
                * **Oś Z (Wysokość - Moc):** Wynik końcowy (Siła x Szybkość).
                * **Kolor (SmO2):** Poziom tlenu w mięśniu (Czerwony = Niedotlenienie, Niebieski = Komfort).

                **🔍 Przykłady z Życia (Szukaj tych obszarów na wykresie):**
                1.  **"Młynek" (Prawa Strona, Nisko):** Wysoka kadencja, niski moment. To jazda ekonomiczna (np. na płaskim). Punkty powinny być **niebieskie/zielone** (dobre ukrwienie, "pompa mięśniowa" działa).
                2.  **"Przepychanie" (Lewa Strona, Wysoko):** Niska kadencja, duża siła (np. sztywny podjazd na twardym przełożeniu). Mięśnie są napięte, krew nie dopływa. Punkty mogą być **czerwone** (hipoksja/okluzja). To męczy mięśnie szybciej niż serce.
                3.  **Sprint (Prawy Górny Róg, Wysoko w górę):** Max kadencja i max siła. Generujesz szczytową moc (Oś Z). To stan beztlenowy, punkty szybko zmienią się na **czerwone**.
                4.  **Jazda w Grupie (Środek):** Umiarkowana kadencja i siła. To Twój "Sweet Spot" biomechaniczny.

                **Wniosek:** Jeśli widzisz dużo czerwonych punktów przy niskiej kadencji, zredukuj bieg i kręć szybciej, aby dotlenić nogi!
                """)

        # --- NEW TAB: NUTRITION ---
        with tab_nutrition:
            st.header("⚡ Kalkulator Spalania Glikogenu (The Bonk Prediction)")
            
            # Interaktywne suwaki
            c1, c2, c3 = st.columns(3)
            carb_intake = c1.number_input("Spożycie Węglowodanów [g/h]", min_value=0, max_value=120, value=60, step=10)
            initial_glycogen = c2.number_input("Początkowy Zapas Glikogenu [g]", min_value=200, max_value=800, value=450, step=50, help="Standardowo: 400-500g dla wytrenowanego sportowca.")
            efficiency_input = c3.number_input("Sprawność Mechaniczna [%]", min_value=18.0, max_value=26.0, value=22.0, step=0.5, help="Amator: 18-21%, Pro: 23%+")
            
            # --- ZMIANA: "MENU KOLARSKIE" (CHEAT SHEET) ---
            with st.expander("🍬 Menu Kolarskie (Ile to węglowodanów?)", expanded=False):
                st.markdown("""
                Aby dostarczyć 90g węgli na godzinę, potrzebujesz np.:
                * **3 x Żel Energetyczny** (standardowo ~25-30g CHO / sztukę)
                * **1.5 Bidonu Izotonika** (standardowo ~40g CHO / 500ml)
                * **3 x Banan** (~25-30g CHO / sztukę)
                * **2 x Baton Energetyczny** (~40-50g CHO / sztukę)
                * **Garść Żelków (100g)** (~75g CHO)
                
                *Pamiętaj: Trening jelita jest równie ważny jak trening nóg! Nie testuj 90g/h pierwszy raz na zawodach.*
                """)
            
            if 'watts' in df_plot.columns:
                intensity_factor = df_plot['watts'] / cp_input
                
                # Model metaboliczny (Logika bez zmian)
                conditions = [
                    (df_plot['watts'] < vt1_watts),
                    (df_plot['watts'] >= vt1_watts) & (df_plot['watts'] < vt2_watts),
                    (df_plot['watts'] >= vt2_watts)
                ]
                choices = [0.3, 0.8, 1.1] 
                carb_fraction = np.select(conditions, choices, default=1.0)
                
                # Obliczenia energii
                energy_kcal_sec = df_plot['watts'] / (efficiency_input/100.0) / 4184.0
                carbs_burned_per_sec = (energy_kcal_sec * carb_fraction) / 4.0
                cumulative_burn = carbs_burned_per_sec.cumsum()
                
                intake_per_sec = carb_intake / 3600.0
                cumulative_intake = np.cumsum(np.full(len(df_plot), intake_per_sec))
                
                glycogen_balance = initial_glycogen - cumulative_burn + cumulative_intake
                
                df_nutri = pd.DataFrame({
                    'Czas [min]': df_plot['time_min'],
                    'Bilans Glikogenu [g]': glycogen_balance,
                    'Spalone [g]': cumulative_burn,
                    'Spożyte [g]': cumulative_intake,
                    'Burn Rate [g/h]': carbs_burned_per_sec * 3600
                })
                
                # --- WYKRES 1: BILANS GLIKOGENU ---
                fig_nutri = go.Figure()
                
                # Linia Balansu
                line_color = '#00cc96' if df_nutri['Bilans Glikogenu [g]'].min() > 0 else '#ef553b'
                
                fig_nutri.add_trace(go.Scatter(
                    x=df_nutri['Czas [min]'], 
                    y=df_nutri['Bilans Glikogenu [g]'], 
                    name='Zapas Glikogenu', 
                    fill='tozeroy', 
                    line=dict(color=line_color, width=2), 
                    hovertemplate="<b>Czas: %{x:.0f} min</b><br>Zapas: %{y:.0f} g<extra></extra>"
                ))
                
                # Linia "Ściana" (Bonk)
                fig_nutri.add_hline(y=0, line_dash="dash", line_color="red", annotation_text="Ściana (Bonk)", annotation_position="bottom right")
                
                fig_nutri.update_layout(
                    template="plotly_dark",
                    title=f"Symulacja Baku Paliwa (Start: {initial_glycogen}g, Intake: {carb_intake}g/h)",
                    hovermode="x unified",
                    yaxis=dict(title="Glikogen [g]"),
                    # ZMIANA TUTAJ: tickformat=".0f" wymusza liczby całkowite
                    xaxis=dict(title="Czas [min]", tickformat=".0f"),
                    margin=dict(l=10, r=10, t=40, b=10),
                    height=400,
                    showlegend=False
                )
                st.plotly_chart(fig_nutri, use_container_width=True)
                
                # --- WYKRES 2: TEMPO SPALANIA (BURN RATE) ---
                st.subheader("🔥 Tempo Spalania (Burn Rate)")
                fig_burn = go.Figure()
                
                burn_rate_smooth = df_nutri['Burn Rate [g/h]'].rolling(window=60, center=True, min_periods=1).mean()
                
                fig_burn.add_trace(go.Scatter(
                    x=df_nutri['Czas [min]'], 
                    y=burn_rate_smooth, 
                    name='Spalanie', 
                    line=dict(color='#ff7f0e', width=2), 
                    fill='tozeroy', 
                    hovertemplate="<b>Czas: %{x:.0f} min</b><br>Spalanie: %{y:.0f} g/h<extra></extra>"
                ))
                
                # Linia Spożycia (Intake)
                fig_burn.add_hline(y=carb_intake, line_dash="dot", line_color="#00cc96", annotation_text=f"Intake: {carb_intake}g/h", annotation_position="top right")
                
                fig_burn.update_layout(
                    template="plotly_dark",
                    title="Zapotrzebowanie na Węglowodany",
                    hovermode="x unified",
                    yaxis=dict(title="Burn Rate [g/h]"),
                    xaxis=dict(title="Czas [min]", tickformat=".0f"),
                    margin=dict(l=10, r=10, t=40, b=10),
                    height=400,
                    showlegend=False
                )
                st.plotly_chart(fig_burn, use_container_width=True)

                # PODSUMOWANIE LICZBOWE
                total_burn = cumulative_burn.iloc[-1]
                total_intake = cumulative_intake[-1]
                final_balance = glycogen_balance.iloc[-1]
                
                n1, n2, n3 = st.columns(3)
                n1.metric("Spalone Węgle", f"{total_burn:.0f} g", help="Suma węglowodanów zużytych na wysiłek")
                n2.metric("Spożyte Węgle", f"{total_intake:.0f} g", help="Suma węglowodanów dostarczonych z jedzenia/napojów")
                n3.metric("Wynik Końcowy", f"{final_balance:.0f} g", delta=f"{final_balance - initial_glycogen:.0f} g", delta_color="inverse" if final_balance < 0 else "normal")
                
                if final_balance < 0:
                    st.error(f"⚠️ **UWAGA:** Według symulacji, Twoje zapasy glikogenu wyczerpały się w okolicach {df_nutri[df_nutri['Bilans Glikogenu [g]'] < 0]['Czas [min]'].iloc[0]:.0f} minuty! To oznacza ryzyko 'odcięcia' (bonk).")
                else:
                    st.success(f"✅ **OK:** Zakończyłeś trening z zapasem {final_balance:.0f}g glikogenu. Strategia żywieniowa wystarczająca dla tej intensywności.")
                
                st.info("""
                **💡 Fizjologia Spalania (Model VT1/VT2):**
                
                * **Strefa Tłuszczowa (< VT1):** Spalasz ok. **20-40g węgli/h**. Reszta to tłuszcz. Tutaj możesz jechać godzinami na samej wodzie.
                * **Strefa Mieszana (VT1 - VT2):** Spalanie węgli skacze do **60-90g/h**. Musisz zacząć jeść (żele/izotonik), żeby nie opróżniać baku.
                * **Strefa Cukrowa (> VT2):** "Turbo". Spalasz **120g/h i więcej**. Twoje jelita nie są w stanie tyle wchłonąć (max ~90g/h). Każda minuta tutaj to "pożyczka", której nie spłacisz w trakcie jazdy.
                
                *Model uwzględnia Twoją wagę, sprawność (Efficiency) oraz progi mocy.*
                """)
            else:
                st.warning("Brak danych mocy (Watts) do obliczenia wydatku energetycznego.")

    # --- TAB SmO2 ---
    with tab_smo2:
        st.header("Analiza Kinetyki SmO2 (LT1 / LT2 Detection)")
        st.markdown("Tutaj szukamy punktów przełamania. Wybierz stabilny odcinek (interwał), a obliczymy trend desaturacji.")

        if 'df_plot' in locals():
            target_df = df_plot
        elif 'df_with_hsi' in locals():
            target_df = df_with_hsi.to_pandas() if hasattr(df_with_hsi, "to_pandas") else df_with_hsi
        elif 'df_clean_pl' in locals():
            target_df = df_clean_pl.to_pandas() if hasattr(df_clean_pl, "to_pandas") else df_clean_pl
        elif 'df_raw' in locals():
            target_df = df_raw.to_pandas() if hasattr(df_raw, "to_pandas") else df_raw
        else:
            st.error("Brak wczytanych danych. Najpierw wgraj plik w sidebar.")
            st.stop()

        if 'time' not in target_df.columns:
            st.error("Brak kolumny 'time' w danych!")
            st.stop()

        target_df['watts_smooth_5s'] = target_df['watts'].rolling(window=5, center=True).mean()
        target_df['smo2_smooth'] = target_df['smo2'].rolling(window=3, center=True).mean()
        target_df['time_str'] = pd.to_datetime(target_df['time'], unit='s').dt.strftime('%H:%M:%S')
        col_inp1, col_inp2 = st.columns(2)
        
        # Inicjalizacja session_state dla zaznaczenia
        if 'smo2_start_sec' not in st.session_state:
            st.session_state.smo2_start_sec = 600  # 10 minut domyślnie
        if 'smo2_end_sec' not in st.session_state:
            st.session_state.smo2_end_sec = 1200  # 20 minut domyślnie
            
        # ===== NOTATKI SmO2 =====
        with st.expander("📝 Dodaj Notatkę do tej Analizy", expanded=False):
            note_col1, note_col2 = st.columns([1, 2])
            with note_col1:
                note_time = st.number_input(
                    "Czas (min)", 
                    min_value=0.0, 
                    max_value=float(len(target_df)/60) if len(target_df) > 0 else 60,
                    value=float(len(target_df)/120) if len(target_df) > 0 else 15,
                    step=0.5,
                    key="smo2_note_time"
                )
            with note_col2:
                note_text = st.text_input(
                    "Notatka",
                    key="smo2_note_text",
                    placeholder="Np. 'Atak 500W', 'Próg beztlenowy', 'Błąd sensoryka'"
                )
            
            if st.button("➕ Dodaj Notatkę", key="smo2_add_note"):
                if note_text:
                    training_notes.add_note(uploaded_file.name, note_time, "smo2", note_text)
                    st.success(f"✅ Notatka: {note_text} @ {note_time:.1f} min")
                else:
                    st.warning("Wpisz tekst notatki!")

        # Wyświetl istniejące notatki SmO2
        existing_notes_smo2 = training_notes.get_notes_for_metric(uploaded_file.name, "smo2")
        if existing_notes_smo2:
            st.subheader("📋 Notatki SmO2")
            for idx, note in enumerate(existing_notes_smo2):
                col_note, col_del = st.columns([4, 1])
                with col_note:
                    st.info(f"⏱️ **{note['time_minute']:.1f} min** | {note['text']}")
                with col_del:
                    if st.button("🗑️", key=f"del_smo2_note_{idx}"):
                        training_notes.delete_note(uploaded_file.name, idx)
                        st.rerun()

        st.markdown("---")
        # ===== KONIEC NOTATEK SmO2 =====

        st.info("💡 **NOWA FUNKCJA:** Zaznacz obszar na wykresie poniżej (kliknij i przeciągnij), aby automatycznie obliczyć metryki!")

        # Opcjonalne: ręczne wprowadzenie czasu (dla precyzji)
        def parse_time_to_seconds(t_str):
            try:
                parts = list(map(int, t_str.split(':')))
                if len(parts) == 3: return parts[0]*3600 + parts[1]*60 + parts[2]
                if len(parts) == 2: return parts[0]*60 + parts[1]
                if len(parts) == 1: return parts[0]
            except:
                return None
            return None

        with st.expander("🔧 Ręczne wprowadzenie zakresu czasowego (opcjonalne)", expanded=False):
            col_inp1, col_inp2 = st.columns(2)
            with col_inp1:
                manual_start = st.text_input("Start Interwału (hh:mm:ss)", value="01:00:00", key="smo2_manual_start")
            with col_inp2:
                manual_end = st.text_input("Koniec Interwału (hh:mm:ss)", value="01:20:00", key="smo2_manual_end")
            
            if st.button("Zastosuj ręczny zakres"):
                manual_start_sec = parse_time_to_seconds(manual_start)
                manual_end_sec = parse_time_to_seconds(manual_end)
                if manual_start_sec is not None and manual_end_sec is not None:
                    st.session_state.smo2_start_sec = manual_start_sec
                    st.session_state.smo2_end_sec = manual_end_sec
                    st.success(f"✅ Zaktualizowano zakres: {manual_start} - {manual_end}")

        # Użyj wartości z session_state
        startsec = st.session_state.smo2_start_sec
        endsec = st.session_state.smo2_end_sec

        start_time_str = st.session_state.get('smo2_manual_start', "01:00:00")
        nd_time_str = st.session_state.get('smo2_manual_end', "01:20:00")
        
        if startsec is not None and endsec is not None:
            if endsec > startsec:
                duration_sec = endsec - startsec
                
                mask = (target_df['time'] >= startsec) & (target_df['time'] <= endsec)
                interval_data = target_df.loc[mask]

                if not interval_data.empty:
                    avg_watts = interval_data['watts'].mean() if 'watts' in interval_data.columns else 0
                    avg_smo2 = interval_data['smo2'].mean() if 'smo2' in interval_data.columns else 0
                    max_smo2 = interval_data['smo2'].max() if 'smo2' in interval_data.columns else 0
                    min_smo2 = interval_data['smo2'].min() if 'smo2' in interval_data.columns else 0
                    
                    if len(interval_data) > 1 and 'smo2' in interval_data.columns:
                        slope, intercept, r_value, p_value, std_err = stats.linregress(interval_data['time'], interval_data['smo2'])
                        trend_desc = f"{slope:.4f} %/s"
                    else:
                        slope = 0
                        intercept = 0
                        trend_desc = "N/A"

                    st.subheader(f"Metryki dla odcinka: {start_time_str} - {nd_time_str} (Czas trwania: {duration_sec}s)")
                    m1, m2, m3, m4, m5 = st.columns(5)
                    m1.metric("Śr. Moc", f"{avg_watts:.0f} W")
                    m2.metric("Śr. SmO2", f"{avg_smo2:.1f} %")
                    m3.metric("Min SmO2", f"{min_smo2:.1f} %", delta_color="inverse")
                    m4.metric("Max SmO2", f"{max_smo2:.1f} %")
                    
                    delta_color = "normal" if slope >= -0.01 else "inverse" 
                    m5.metric("SmO2 Trend (Slope)", trend_desc, delta=trend_desc, delta_color=delta_color)

                    fig_smo2 = go.Figure()

                    fig_smo2.add_trace(go.Scatter(
                        x=target_df['time'], 
                        y=target_df['smo2_smooth'],
                        customdata=target_df['time_str'],
                        mode='lines', 
                        name='SmO2',
                        line=dict(color='#FF4B4B', width=2),
                        hovertemplate="<b>Czas:</b> %{customdata}<br><b>SmO2:</b> %{y:.0f}%<extra></extra>"
                    ))

                    fig_smo2.add_trace(go.Scatter(
                        x=target_df['time'], 
                        y=target_df['watts_smooth_5s'],
                        customdata=target_df['time_str'],
                        mode='lines', 
                        name='Power',
                        line=dict(color='#1f77b4', width=1),
                        yaxis='y2',
                        opacity=0.3,
                        hovertemplate="<b>Czas:</b> %{customdata}<br><b>Moc:</b> %{y:.0f} W<extra></extra>"
                    ))

                    fig_smo2.add_vrect(
                        x0=startsec, x1=endsec,
                        fillcolor="green", opacity=0.1,
                        layer="below", line_width=0,
                        annotation_text="ANALIZA", annotation_position="top left"
                    )
                    
                    if len(interval_data) > 1:
                        trend_line = intercept + slope * interval_data['time']
                        fig_smo2.add_trace(go.Scatter(
                            x=interval_data['time'], 
                            y=trend_line,
                            customdata=interval_data['time_str'],
                            mode='lines', 
                            name='Trend SmO2',
                            line=dict(color='yellow', width=3, dash='dash'),
                            hovertemplate="<b>Czas:</b> %{customdata}<br><b>Trend:</b> %{y:.1f}%<extra></extra>"
                        ))

                    fig_smo2.update_layout(
                        title="Analiza Przebiegu SmO2 vs Power",
                        xaxis_title="Czas",
                        yaxis=dict(title="SmO2 (%)", range=[0, 100]),
                        yaxis2=dict(title="Power (W)", overlaying='y', side='right', showgrid=False),
                        legend=dict(x=0.01, y=0.99),
                        height=500,
                        margin=dict(l=20, r=20, t=40, b=20),
                        hovermode="x unified"
                    )

                    # Wykres z interaktywnym zaznaczaniem
                    selected = st.plotly_chart(fig_smo2, use_container_width=True, key="smo2_chart", on_select="rerun", selection_mode="box")

                    # Obsługa zaznaczenia
                    if selected and 'selection' in selected and 'box' in selected['selection']:
                        box_data = selected['selection']['box']
                        if box_data and len(box_data) > 0:
                            # Pobierz zakres X (czas) z zaznaczenia
                            x_range = box_data[0].get('x', [])
                            if len(x_range) == 2:
                                new_start = min(x_range)
                                new_end = max(x_range)
                                
                                # Aktualizuj session_state
                                if new_start != st.session_state.smo2_start_sec or new_end != st.session_state.smo2_end_sec:
                                    st.session_state.smo2_start_sec = new_start
                                    st.session_state.smo2_end_sec = new_end
                                    st.rerun()

                    # --- PĘTLA HISTEREZY (SmO2 vs WATTS) ---
                    st.divider()
                    st.subheader("🔄 Pętla Histerezy (Opóźnienie Metaboliczne)")
                
                    if 'watts_smooth_5s' in interval_data.columns and 'smo2_smooth' in interval_data.columns:
                        
                        fig_hyst = go.Figure()

                        fig_hyst.add_trace(go.Scatter(
                            x=interval_data['watts_smooth_5s'],
                            y=interval_data['smo2_smooth'],
                            mode='markers+lines',
                            name='Histereza',
                            marker=dict(
                                size=6,
                                color=interval_data['time'], 
                                colorscale='Plasma',
                                showscale=True,
                                colorbar=dict(title="Upływ Czasu", tickmode="array", ticktext=["Start", "Koniec"], tickvals=[interval_data['time'].min(), interval_data['time'].max()])
                            ),
                            line=dict(color='rgba(255,255,255,0.3)', width=1), # Cienka linia łącząca
                            hovertemplate="<b>Moc:</b> %{x:.0f} W<br><b>SmO2:</b> %{y:.1f}%<extra></extra>"
                        ))

                        start_pt = interval_data.iloc[0]
                        end_pt = interval_data.iloc[-1]

                        fig_hyst.add_annotation(
                            x=start_pt['watts_smooth_5s'], y=start_pt['smo2_smooth'],
                            text="START", showarrow=True, arrowhead=2, ax=0, ay=-40, bgcolor="green"
                        )
                        fig_hyst.add_annotation(
                            x=end_pt['watts_smooth_5s'], y=end_pt['smo2_smooth'],
                            text="META", showarrow=True, arrowhead=2, ax=0, ay=-40, bgcolor="red"
                        )

                        fig_hyst.update_layout(
                            template="plotly_dark",
                            title="Kinetyka Tlenowa: Relacja Moc (Wymuszenie) vs SmO2 (Odpowiedź)",
                            xaxis_title="Moc [W]",
                            yaxis_title="SmO2 [%]",
                            height=600,
                            margin=dict(l=20, r=20, t=40, b=20),
                            hovermode="closest"
                        )

                        c_h1, c_h2 = st.columns([3, 1])
                        with c_h1:
                            st.plotly_chart(fig_hyst, use_container_width=True)
                        
                        with c_h2:
                            st.info("""
                            **📚 Interpretacja Kliniczna:**
                            
                            Ten wykres pokazuje **bezwładność** Twojego metabolizmu.
                            
                            * **Oś X:** Co robisz (Waty).
                            * **Oś Y:** Jak reaguje mięsień (Tlen).
                            
                            **Kształt Pętli:**
                            1.  **Wąska (Linia):** Idealne dopasowanie. Podaż tlenu nadąża za popytem w czasie rzeczywistym. Stan "Steady State".
                            2.  **Szeroka Pętla:** Duże opóźnienie. 
                                * Na początku interwału (wzrost mocy) SmO2 spada powoli (korzystasz z zapasów mioglobiny/fosfokreatyny).
                                * Na końcu (spadek mocy) SmO2 rośnie powoli (spłacasz dług tlenowy).
                            
                            **Kierunek (Clockwise):**
                            Typowy dla fizjologii wysiłku. Najpierw rośnie moc, potem spada tlen.
                            """)
                    else:
                        st.warning("Brakuje wygładzonych danych mocy lub SmO2 dla tego interwału.")
                        
                    # 6. SEKCJA TEORII (Rozwijana)
                    with st.expander("📚 TEORIA: Jak wyznaczyć LT1 i LT2 z SmO2? (Kliknij, aby rozwinąć)", expanded=False):
                        st.markdown("""
                        ### 1. Interpretacja Slope (Nachylenia Trendu)
                        Slope mówi nam o równowadze między dostawą a zużyciem tlenu w mięśniu.
                        
                        * **Slope > 0 (Dodatni): "Luksus Tlenowy"**
                            * *Co się dzieje:* Dostawa tlenu przewyższa zużycie.
                            * *Kiedy:* Rozgrzewka, regeneracja, początek interwału (rzut serca rośnie szybciej niż zużycie).
                        
                        * **Slope ≈ 0 (Bliski Zera): "Steady State"**
                            * *Wartości:* Zazwyczaj od **-0.005 do +0.005 %/s**.
                            * *Co się dzieje:* Równowaga. Tyle ile mięsień potrzebuje, tyle krew dostarcza.
                            * *Kiedy:* Jazda w strefie tlenowej (Z2), Sweet Spot (jeśli wytrenowany).
                    
                    * **Slope < 0 (Ujemny): "Desaturacja / Dług Tlenowy"**
                        * *Wartości:* Poniżej **-0.01 %/s** (wyraźny spadek).
                        * *Co się dzieje:* Mitochondria zużywają więcej tlenu niż jest dostarczane. Mioglobina traci tlen.
                        * *Kiedy:* Jazda powyżej progu beztlenowego (LT2), mocne skoki mocy.

                    ---

                    ### 2. Jak znaleźć progi (Breakpoints)?
                    
                    #### 🟢 LT1 (Aerobic Threshold)
                    Szukaj mocy, przy której Slope zmienia się z **dodatniego na płaski (bliski 0)**.
                    * *Przykład:* Przy 180W SmO2 jeszcze rośnie, przy 200W staje w miejscu. **LT1 ≈ 200W**.
                    
                    #### 🔴 LT2 (Anaerobic Threshold / Critical Power)
                    Szukaj mocy, przy której **nie jesteś w stanie ustabilizować SmO2** (brak Steady State).
                    * *Scenariusz:*
                        * 280W: SmO2 spada, ale po minucie się poziomuje (Slope wraca do 0). -> **Jesteś pod progiem.**
                        * 300W: SmO2 leci w dół ciągle i nie chce się zatrzymać (Slope ciągle ujemny). -> **Jesteś nad progiem (powyżej LT2).**
                    
                    ---
                    
                    ### ⚠️ WAŻNE: Pro Tip Biomechaniczny
                    **Uważaj na niską kadencję (Grinding)!**
                    Przy tej samej mocy, niska kadencja = wyższy moment siły (Torque). To powoduje większy ucisk mechaniczny na naczynia krwionośne w mięśniu (okluzja).
                    * *Efekt:* SmO2 może spadać gwałtownie (sztuczna desaturacja) tylko przez mechanikę, mimo że metabolicznie organizm dałby radę.
                    * *Rada:* Testy progowe rób na swojej naturalnej, stałej kadencji.
                    """)
    
                else:
                    st.warning("Brak danych w wybranym zakresie. Sprawdź poprawność wpisanego czasu.")
            else:
                st.error("Czas zakończenia musi być późniejszy niż czas rozpoczęcia!")
        else:
            st.warning("Wprowadź poprawne czasy w formacie h:mm:ss (np. 0:10:00).")

    # --- TAB HEMODYNAMICS (THb vs SmO2) ---
    with tab_hemo:
        st.header("Profil Hemodynamiczny (Mechanika vs Metabolizm)")
        st.markdown("Analiza relacji objętości krwi (THb) do saturacji (SmO2). Pozwala wykryć okluzję (ucisk) i limitery przepływu.")

        if 'df_plot' in locals():
            target_df = df_plot
        elif 'df_with_hsi' in locals():
            target_df = df_with_hsi.to_pandas() if hasattr(df_with_hsi, "to_pandas") else df_with_hsi
        elif 'df_clean_pl' in locals():
            target_df = df_clean_pl.to_pandas() if hasattr(df_clean_pl, "to_pandas") else df_clean_pl
        elif 'df_raw' in locals():
            target_df = df_raw.to_pandas() if hasattr(df_raw, "to_pandas") else df_raw
        else:
            st.error("Brak danych. Najpierw wgraj plik.")
            st.stop()
        col_thb = next((c for c in ['thb', 'total_hemoglobin', 'total_hgb'] if c in target_df.columns), None)
        col_smo2 = 'smo2_smooth' if 'smo2_smooth' in target_df else ('smo2' if 'smo2' in target_df else None)

        if col_thb and col_smo2:
            
            if f"{col_thb}_smooth" not in target_df.columns:
                target_df[f'{col_thb}_smooth'] = target_df[col_thb].rolling(window=10, center=True).mean()
            
            thb_val = f'{col_thb}_smooth'

            if 'smo2' in target_df.columns:
                target_df['smo2_smooth_10s_hemo_trend'] = target_df['smo2'].rolling(window=10, center=True).mean()
                col_smo2_hemo_trend = 'smo2_smooth_10s_hemo_trend'
            else:
                col_smo2_hemo_trend = col_smo2 
            
            # 2. Wykres XY (Scatter) - SmO2 vs THb
            # Kolorujemy punktami Mocy, żeby widzieć co się dzieje na wysokich watach
            
            # Próbkowanie dla szybkości (oryginalne zachowanie)
            df_hemo = target_df.sample(min(len(target_df), 5000))
            
            fig_hemo = px.scatter(
                df_hemo, 
                x=col_smo2, # Revert to original col_smo2 (3s smoothed or raw)
                y=thb_val, 
                color='watts', 
                title="Hemo-Scatter: SmO2 (Oś X) vs THb (Oś Y)", # Revert title
                labels={col_smo2: "SmO2 (Saturacja) [%]", thb_val: "THb (Objętość Krwi) [a.u.]", "watts": "Moc [W]"},
                template="plotly_dark",
                color_continuous_scale='Turbo' # Turbo jest świetne do pokazywania intensywności
            )
            
            # Odwracamy oś X dla SmO2 (zwyczajowo w fizjologii wykresy czyta się od prawej do lewej dla desaturacji)
            fig_hemo.update_xaxes(autorange="reversed")
            
            fig_hemo.update_traces(marker=dict(size=5, opacity=0.6))
            fig_hemo.update_layout(
                height=600,
                margin=dict(l=20, r=20, t=40, b=20)
            )
            
            # Dodajemy adnotacje "ćwiartek" (Uproszczona interpretacja)
            # To wymagałoby znania średnich, ale damy opisy w rogach
            fig_hemo.add_annotation(xref="paper", yref="paper", x=0.05, y=0.95, text="<b>Stres Metaboliczny</b><br>(Wazodylatacja)", showarrow=False, font=dict(color="#00cc96"))
            fig_hemo.add_annotation(xref="paper", yref="paper", x=0.05, y=0.05, text="<b>OKLUZJA / UCISK</b><br>(Limit Przepływu)", showarrow=False, font=dict(color="#ef553b"))
            fig_hemo.add_annotation(xref="paper", yref="paper", x=0.95, y=0.95, text="<b>Regeneracja</b><br>(Napływ)", showarrow=False, font=dict(color="#ffa15a"))
            
            st.plotly_chart(fig_hemo, use_container_width=True)
            
            # 3. Wykres Liniowy w czasie (Dual Axis)
            st.subheader("Trendy w Czasie (Szukanie Rozjazdu)")
            fig_trend = go.Figure()
            
            # SmO2 (Oś Lewa)
            fig_trend.add_trace(go.Scatter(
                x=target_df['time_min'], y=target_df[col_smo2_hemo_trend],
                name='SmO2', line=dict(color='#ab63fa', width=2),
                hovertemplate="SmO2: %{y:.1f}%<extra></extra>"
            ))

            
            # THb (Oś Prawa)
            fig_trend.add_trace(go.Scatter(
                x=target_df['time_min'], y=target_df[thb_val],
                name='THb', line=dict(color='#ffa15a', width=2), yaxis='y2',
                hovertemplate="THb: %{y:.2f}<extra></extra>"
            ))
            
            # Tło - Moc (dla kontekstu)
            if 'watts_smooth_30s' in target_df:
                 fig_trend.add_trace(go.Scatter(
                    x=target_df['time_min'], y=target_df['watts_smooth_30s'],
                    name='Moc', line=dict(color='rgba(255,255,255,0.1)', width=1),
                    fill='tozeroy', fillcolor='rgba(255,255,255,0.05)', yaxis='y3',
                    hoverinfo='skip'
                ))

            # Poprawiony Layout dla fig_trend (bez titlefont)
            fig_trend.update_layout(
                template="plotly_dark",
                title="SmO2 vs THb w Czasie",
                hovermode="x unified",
                yaxis=dict(
                    title=dict(text="SmO2 [%]", font=dict(color='#ab63fa'))
                ),
                yaxis2=dict(
                    title=dict(text="THb [a.u.]", font=dict(color='#ffa15a')),
                    overlaying='y', side='right'
                ),
                yaxis3=dict(title="Moc", overlaying='y', side='right', showgrid=False, showticklabels=False), 
                height=450
            )
            st.plotly_chart(fig_trend, use_container_width=True)

            # 4. Teoria dla Fizjologii
            st.info("""
            **💡 Interpretacja Hemodynamiczna (THb + SmO2):**
            
            THb (Total Hemoglobin) to wskaźnik objętości krwi ("tHb = pompa paliwowa"). SmO2 to wskaźnik zużycia ("SmO2 = bak").
            
            * **Scenariusz 1: Dobra praca (Wazodylatacja)**
                * **SmO2 SPADA 📉 | THb ROŚNIE 📈**
                * *Co to znaczy:* Mięsień pracuje mocno, metabolizm zużywa tlen, ale układ krążenia reaguje prawidłowo, rozszerzając naczynia i pompując więcej krwi. To zdrowy limit metaboliczny.
            
            * **Scenariusz 2: Okluzja / Limit Mechaniczny (UWAGA!)**
                * **SmO2 SPADA 📉 | THb SPADA 📉 (lub płaskie)**
                * *Co to znaczy:* "Wyżymanie gąbki". Napięcie mięśnia jest tak duże (lub kadencja za niska), że ciśnienie wewnątrzmięśniowe blokuje dopływ świeżej krwi.
                * *Działanie:* Zwiększ kadencję, sprawdź siodełko (czy nie uciska tętnic), popraw fit.
            
            * **Scenariusz 3: Venous Pooling (Zastój)**
                * **SmO2 ROŚNIE 📈 | THb ROŚNIE 📈**
                * *Kiedy:* Często podczas nagłego zatrzymania po wysiłku. Krew napływa, ale pompa mięśniowa nie odprowadza jej z powrotem.
            """)

        else:
            st.warning("⚠️ Brak danych THb (Total Hemoglobin). Sensor Moxy/Train.Red powinien dostarczać tę kolumnę (często jako 'thb' lub 'total_hemoglobin'). Bez tego analiza hemodynamiczna jest niemożliwa.")
            st.markdown("Dostępne kolumny w pliku: " + ", ".join(target_df.columns))

    # --- TAB VENT ANALYSIS (VT1 / VT2) ---
    with tab_vent:
        st.header("Analiza Progu Wentylacyjnego (VT1 / VT2 Detection)")
        st.markdown("Analiza dynamiki oddechu. Szukamy nieliniowych przyrostów wentylacji (VE) względem mocy.")

        # 1. Przygotowanie danych
        if 'df_plot' in locals():
            target_df = df_plot
        elif 'df_with_hsi' in locals():
            target_df = df_with_hsi.to_pandas() if hasattr(df_with_hsi, "to_pandas") else df_with_hsi
        elif 'df_clean_pl' in locals():
            target_df = df_clean_pl.to_pandas() if hasattr(df_clean_pl, "to_pandas") else df_clean_pl
        elif 'df_raw' in locals():
            target_df = df_raw.to_pandas() if hasattr(df_raw, "to_pandas") else df_raw
        else:
            st.error("Brak danych.")
            st.stop()

        if 'time' not in target_df.columns or 'tymeventilation' not in target_df.columns:
            st.error("Brak danych wentylacji (tymeventilation) lub czasu!")
            st.stop()

        # Wygładzanie (VE jest szumiące, dajemy 10s smooth)
        target_df['watts_smooth_5s'] = target_df['watts'].rolling(window=5, center=True).mean()
        target_df['ve_smooth'] = target_df['tymeventilation'].rolling(window=10, center=True).mean()
        target_df['rr_smooth'] = target_df['tymebreathrate'].rolling(window=10, center=True).mean() if 'tymebreathrate' in target_df else 0
        
        # Format czasu
        target_df['time_str'] = pd.to_datetime(target_df['time'], unit='s').dt.strftime('%H:%M:%S')

        # 2. Interfejs (START -> KONIEC)
        # Inicjalizacja session_state dla zaznaczenia
        if 'vent_start_sec' not in st.session_state:
                st.session_state.vent_start_sec = 600  # 10 minut domyślnie
        if 'vent_end_sec' not in st.session_state:
                st.session_state.vent_end_sec = 1200  # 20 minut domyślnie
                
        # ===== NOTATKI VENTILATION =====
        with st.expander("📝 Dodaj Notatkę do tej Analizy", expanded=False):
            note_col1, note_col2 = st.columns([1, 2])
            with note_col1:
                note_time_vent = st.number_input(
                    "Czas (min)", 
                    min_value=0.0, 
                    max_value=float(len(target_df)/60) if len(target_df) > 0 else 60,
                    value=float(len(target_df)/120) if len(target_df) > 0 else 15,
                    step=0.5,
                    key="vent_note_time"
                )
            with note_col2:
                note_text_vent = st.text_input(
                    "Notatka",
                    key="vent_note_text",
                    placeholder="Np. 'Próg beztlenowy', 'VE jump', 'Spłycenie oddechu'"
                )
            
            if st.button("➕ Dodaj Notatkę", key="vent_add_note"):
                if note_text_vent:
                    training_notes.add_note(uploaded_file.name, note_time_vent, "ventilation", note_text_vent)
                    st.success(f"✅ Notatka: {note_text_vent} @ {note_time_vent:.1f} min")
                else:
                    st.warning("Wpisz tekst notatki!")

        # Wyświetl istniejące notatki Ventilation
        existing_notes_vent = training_notes.get_notes_for_metric(uploaded_file.name, "ventilation")
        if existing_notes_vent:
            st.subheader("📋 Notatki Wentylacji")
            for idx, note in enumerate(existing_notes_vent):
                col_note, col_del = st.columns([4, 1])
                with col_note:
                    st.info(f"⏱️ **{note['time_minute']:.1f} min** | {note['text']}")
                with col_del:
                    if st.button("🗑️", key=f"del_vent_note_{idx}"):
                        training_notes.delete_note(uploaded_file.name, idx)
                        st.rerun()

        st.markdown("---")
        # ===== KONIEC NOTATEK VENTILATION =====

        st.info("💡 **NOWA FUNKCJA:** Zaznacz obszar na wykresie poniżej (kliknij i przeciągnij), aby automatycznie obliczyć metryki!")

            # Opcjonalne: ręczne wprowadzenie czasu (dla precyzji)
        with st.expander("🔧 Ręczne wprowadzenie zakresu czasowego (opcjonalne)", expanded=False):
                col_inp_1, col_inp_2 = st.columns(2)
                with col_inp_1:
                    manual_start = st.text_input("Start Interwału (hh:mm:ss)", value="01:00:00", key="vent_manual_start")
                with col_inp_2:
                    manual_end = st.text_input("Koniec Interwału (hh:mm:ss)", value="01:20:00", key="vent_manual_end")

                if st.button("Zastosuj ręczny zakres", key="btn_vent_manual"):
                    manual_start_sec = parse_time_to_seconds(manual_start)
                    manual_end_sec = parse_time_to_seconds(manual_end)
                    if manual_start_sec is not None and manual_end_sec is not None:
                        st.session_state.vent_start_sec = manual_start_sec
                        st.session_state.vent_end_sec = manual_end_sec
                        st.success(f"✅ Zaktualizowano zakres: {manual_start} - {manual_end}")

            # Użyj wartości z session_state
        startsec = st.session_state.vent_start_sec
        endsec = st.session_state.vent_end_sec

            
            # 3. Wycinanie
        mask_v = (target_df['time'] >= startsec) & (target_df['time'] <= endsec)
        interval_v = target_df.loc[mask_v]

        if not interval_v.empty:
                # 4. Obliczenia
                avg_w = interval_v['watts'].mean()
                avg_ve = interval_v['tymeventilation'].mean()
                avg_rr = interval_v['tymebreathrate'].mean() if 'tymebreathrate' in interval_v else 0
                max_ve = interval_v['tymeventilation'].max()
                
                # Ve/Power Ratio (Efektywność)
                ve_power_ratio = avg_ve / avg_w if avg_w > 0 else 0
                
                # Trend (Slope) dla VE
                if len(interval_v) > 1:
                    slope_ve, intercept_ve, _, _, _ = stats.linregress(interval_v['time'], interval_v['tymeventilation'])
                    trend_desc_ve = f"{slope_ve:.4f} L/s"
                else:
                    slope_ve = 0; intercept_ve = 0; trend_desc_ve = "N/A"

                # Formatowanie czasu dla wyświetlania
                def fmt_time_v(seconds):
                    try:
                        seconds = int(seconds)
                        h = seconds // 3600
                        m = (seconds % 3600) // 60
                        s = seconds % 60
                        if h > 0:
                            return f"{h:02d}:{m:02d}:{s:02d}"
                        else:
                            return f"{m:02d}:{s:02d}"
                    except:
                        return "-"
                start_time_v = fmt_time_v(startsec)
                end_time_v = fmt_time_v(endsec)
                duration_v = int(endsec - startsec) if (endsec is not None and startsec is not None) else 0

                # Metryki
                st.subheader(f"Metryki Oddechowe: {start_time_v} - {end_time_v} ({duration_v}s)")
                mv1, mv2, mv3, mv4, mv5 = st.columns(5)
                mv1.metric("Śr. Moc", f"{avg_w:.0f} W")
                mv2.metric("Śr. Wentylacja (VE)", f"{avg_ve:.1f} L/min")
                mv3.metric("Częstość (RR)", f"{avg_rr:.1f} /min")
                mv4.metric("Wydajność (VE/W)", f"{ve_power_ratio:.3f}", help="Ile litrów powietrza na 1 Wat mocy. Niżej = lepiej (do pewnego momentu).")
                
                # Kolorowanie trendu (Tu odwrotnie niż w SmO2: Duży wzrost = Czerwony/Ostrzegawczy)
                trend_color = "inverse" if slope_ve > 0.1 else "normal"
                mv5.metric("Trend VE (Slope)", trend_desc_ve, delta=trend_desc_ve, delta_color=trend_color)

                # 5. Wykres
                fig_vent = go.Figure()

                # Lewa Oś: Wentylacja
                fig_vent.add_trace(go.Scatter(
                    x=target_df['time'], y=target_df['ve_smooth'],
                    customdata=target_df['time_str'],
                    mode='lines', name='VE (L/min)',
                    line=dict(color='#ffa15a', width=2),
                    hovertemplate="<b>Czas:</b> %{customdata}<br><b>VE:</b> %{y:.1f} L/min<extra></extra>"
                ))

                # Prawa Oś: Moc
                fig_vent.add_trace(go.Scatter(
                    x=target_df['time'], y=target_df['watts_smooth_5s'],
                    customdata=target_df['time_str'],
                    mode='lines', name='Power',
                    line=dict(color='#1f77b4', width=1),
                    yaxis='y2', opacity=0.3,
                    hovertemplate="<b>Czas:</b> %{customdata}<br><b>Moc:</b> %{y:.0f} W<extra></extra>"
                ))

                # Zaznaczenie
                fig_vent.add_vrect(x0=startsec, x1=endsec, fillcolor="orange", opacity=0.1, layer="below", annotation_text="ANALIZA", annotation_position="top left")

                # Linia trendu VE
                if len(interval_v) > 1:
                    trend_line_ve = intercept_ve + slope_ve * interval_v['time']
                    fig_vent.add_trace(go.Scatter(
                        x=interval_v['time'], y=trend_line_ve,
                        customdata=interval_v['time_str'],
                        mode='lines', name='Trend VE',
                        line=dict(color='white', width=2, dash='dash'),
                        hovertemplate="<b>Trend:</b> %{y:.2f} L/min<extra></extra>"
                    ))

                fig_vent.update_layout(
                    title="Dynamika Wentylacji vs Moc",
                    xaxis_title="Czas",
                    yaxis=dict(title=dict(text="Wentylacja (L/min)", font=dict(color="#ffa15a"))),
                    yaxis2=dict(title=dict(text="Moc (W)", font=dict(color="#1f77b4")), overlaying='y', side='right', showgrid=False),
                    legend=dict(x=0.01, y=0.99),
                    height=500,
                    margin=dict(l=20, r=20, t=40, b=20),
                    hovermode="x unified"
                )
                # Wykres z interaktywnym zaznaczaniem
                selected = st.plotly_chart(fig_vent, use_container_width=True, key="vent_chart", on_select="rerun", selection_mode="box")

                # Obsługa zaznaczenia
                if selected and 'selection' in selected and 'box' in selected['selection']:
                    box_data = selected['selection']['box']
                    if box_data and len(box_data) > 0:
                        # Pobierz zakres X (czas) z zaznaczenia
                        x_range = box_data[0].get('x', [])
                        if len(x_range) == 2:
                            new_start = min(x_range)
                            new_end = max(x_range)
                            
                            # Aktualizuj session_state
                            if new_start != st.session_state.vent_start_sec or new_end != st.session_state.vent_end_sec:
                                st.session_state.vent_start_sec = new_start
                                st.session_state.vent_end_sec = new_end
                                st.rerun()

                # 6. TEORIA ODDECHOWA
                with st.expander("🫁 TEORIA: Jak znaleźć VT1 i VT2 na podstawie Slope?", expanded=False):
                    st.markdown("""
                    ### Interpretacja Slope (Nachylenia VE)
                    Wentylacja rośnie nieliniowo. Szukamy punktów załamania krzywej ("Kinks").

                    #### 🟢 1. Strefa Tlenowa (Poniżej VT1)
                    * **Zachowanie:** VE rośnie proporcjonalnie do mocy (liniowo).
                    * **Slope:** Stabilny, umiarkowanie dodatni (np. 0.02 - 0.05 L/s).
                    * **RR (Oddechy):** Stabilne, wolne pogłębianie oddechu.

                    #### 🟡 2. Próg VT1 (Aerobic Threshold) - "Pierwsze Przełamanie"
                    * **Co szukać:** Pierwszy moment, gdzie Slope wyraźnie wzrasta, mimo że moc rośnie liniowo.
                    * **Fizjologia:** Buforowanie kwasu mlekowego wodorowęglanami -> powstaje ekstra CO2 -> musisz go wydychać.
                    * **Test mowy:** Tutaj zaczynasz urywać zdania.

                    #### 🔴 3. Próg VT2 (Respiratory Compensation Point) - "Drugie Przełamanie"
                    * **Co szukać:** Slope wystrzeliwuje w górę ("Vertical spike"). VE rośnie wykładniczo.
                    * **Wartości Slope:** Bardzo wysokie (np. > 0.15 L/s).
                    * **RR (Oddechy):** Gwałtowny wzrost częstości (tachypnoe).
                    * **Fizjologia:** Hiperwentylacja. Organizm nie nadąża z usuwaniem CO2. Koniec równowagi.
                    ---
                    **Pro Tip:** Porównaj Slope VE ze Slope Mocy. Jeśli Moc rośnie o 5%, a VE o 15% -> właśnie przekroczyłeś próg.
                    """)
        else:
            st.warning("Brak danych w tym zakresie.")
    
    # --- TAB LIMITERS (RADAR CHART) ---
    with tab_limiters:
        st.header("Analiza Limiterów Fizjologicznych (Radar)")
        st.markdown("Sprawdzamy, który układ (Serce, Płuca, Mięśnie) był 'wąskim gardłem' podczas najcięższych momentów treningu.")

        # Sprawdzamy dostępność danych
        has_hr = 'heartrate' in df_plot.columns
        has_ve = any(c in df_plot.columns for c in ['tymeventilation', 've', 'ventilation'])
        has_smo2 = 'smo2' in df_plot.columns
        has_watts = 'watts' in df_plot.columns

        if has_watts and (has_hr or has_ve or has_smo2):
            
            # 1. Wybór okna czasowego (Peak Power)
            window_options = {
                "1 min (Anaerobic)": 60, 
                "5 min (VO2max)": 300, 
                "20 min (FTP)": 1200,
                "60 min (Endurance)": 3600
            }
            selected_window_name = st.selectbox("Wybierz okno analizy (MMP):", list(window_options.keys()), index=1)
            window_sec = window_options[selected_window_name]

            # Znajdujemy indeks startu dla najlepszej średniej mocy w tym oknie
            # Rolling musi mieć min_periods=window_sec, żeby nie liczyć "połówek" na początku
            df_plot['rolling_watts'] = df_plot['watts'].rolling(window=window_sec, min_periods=window_sec).mean()

            if df_plot['rolling_watts'].isna().all():
                st.warning(f"Trening jest krótszy niż {window_sec/60:.0f} min. Wybierz krótsze okno.")
                st.stop()

            peak_idx = df_plot['rolling_watts'].idxmax()

            # Sprawdzamy, czy znaleziono peak (czy trening był wystarczająco długi)
            if not pd.isna(peak_idx):
                # Wycinamy ten fragment danych
                start_idx = max(0, peak_idx - window_sec + 1)
                df_peak = df_plot.iloc[start_idx:peak_idx+1]
                
                # 2. Obliczamy % wykorzystania potencjału (Estymacja Maxów)
                
                # HR (Centralny)
                peak_hr_avg = df_peak['heartrate'].mean() if has_hr else 0
                max_hr_user = df_plot['heartrate'].max() 
                pct_hr = (peak_hr_avg / max_hr_user * 100) if max_hr_user > 0 else 0
                
                # VE (Oddechowy)
                col_ve_nm = next((c for c in ['tymeventilation', 've', 'ventilation'] if c in df_plot.columns), None)
                peak_ve_avg = df_peak[col_ve_nm].mean() if col_ve_nm else 0
                # Estymujemy Max VE jako 110% VT2 (bezpieczny margines dla RCP)
                max_ve_user = vt2_vent * 1.1 
                pct_ve = (peak_ve_avg / max_ve_user * 100) if max_ve_user > 0 else 0
                
                # SmO2 (Lokalny) - Odwrócona logika (im mniej tym "więcej" pracy)
                peak_smo2_avg = df_peak['smo2'].mean() if has_smo2 else 100
                # Używamy 100 - SmO2 jako "stopnia ekstrakcji tlenu"
                pct_smo2_util = 100 - peak_smo2_avg
                
                # Power (Mechaniczny) vs CP
                peak_w_avg = df_peak['watts'].mean()
                pct_power = (peak_w_avg / cp_input * 100) if cp_input > 0 else 0

                # 3. Rysujemy Radar
                categories = ['Serce (% HRmax)', 'Płuca (% VEmax)', 'Mięśnie (% Desat)', 'Moc (% CP)']
                values = [pct_hr, pct_ve, pct_smo2_util, pct_power]
                
                # Zamykamy koło dla wykresu radarowego
                values += [values[0]]
                categories += [categories[0]]

                fig_radar = go.Figure()
                fig_radar.add_trace(go.Scatterpolar(
                    r=values,
                    theta=categories,
                    fill='toself',
                    name=selected_window_name,
                    line=dict(color='#00cc96'),
                    fillcolor='rgba(0, 204, 150, 0.3)',
                    hovertemplate="%{theta}: <b>%{r:.1f}%</b><extra></extra>"
                ))

                # Dynamiczna skala - jeśli moc wyskoczy poza 120% (np. przy 1 min), zwiększamy zakres
                max_val = max(values)
                range_max = 100 if max_val < 100 else (max_val + 10)

                fig_radar.update_layout(
                    polar=dict(
                        radialaxis=dict(
                            visible=True,
                            range=[0, range_max] 
                        )
                    ),
                    template="plotly_dark",
                    title=f"Profil Obciążenia: {selected_window_name} ({peak_w_avg:.0f} W)",
                    height=500
                )
                
                st.plotly_chart(fig_radar, use_container_width=True)
                
                # 4. Interpretacja
                st.info(f"""
                **🔍 Diagnoza dla odcinka {selected_window_name}:**
                
                * **Serce (Central):** {pct_hr:.1f}% Maxa. (Wysokie tętno = koszt transportu).
                * **Płuca (Oddech):** {pct_ve:.1f}% Szacowanego Maxa. (Wysokie VE = koszt usunięcia CO2).
                * **Mięśnie (Lokalne):** {pct_smo2_util:.1f}% Wykorzystania tlenu (Średnie SmO2: {peak_smo2_avg:.1f}%).
                * **Moc:** {pct_power:.0f}% Twojego CP/FTP.
                
                **Co Cię zatrzymało?**
                Patrz, który "wierzchołek" jest najdalej od środka.
                * Jeśli **Serce > Mięśnie**: Ograniczenie centralne (układ krążenia nie nadąża z dostawą).
                * Jeśli **Mięśnie > Serce**: Ograniczenie peryferyjne (mięśnie zużywają wszystko, co dostają, albo jest okluzja mechaniczna).
                """)
            else:
                st.warning(f"Twój trening jest krótszy niż {window_sec/60:.0f} min, więc nie możemy wyznaczyć tego okna.")
        else:
            st.error("Brakuje kluczowych danych (Watts + HR/VE/SmO2) do stworzenia radaru.")

    # --- ZAKŁADKA AI / MODEL ---
    with tab_ai:
        st.header("🧠 AI Neural Coach (Powered by Apple MLX)")
        st.caption("Analiza 'Bazy Tlenowej' (280W) oraz 'Silnika' (360W)")

        if MLX_AVAILABLE:
            col_ai_1, col_ai_2 = st.columns([1, 2])
            
            with col_ai_1:
                st.info("System Neuralny Gotowy.")
                if st.button("🚀 Trenuj Mózg (Aktualizuj)", type="primary"):
                    pass 
                
                last_base, last_thresh = "-", "-"
                
                if os.path.exists(HISTORY_FILE):
                    try:
                        with open(HISTORY_FILE, 'r') as f:
                            h_data = json.load(f)
                            
                            if h_data:
                                for entry in reversed(h_data):
                                    val = entry.get('hr_base')
                                    if val is not None and val != "None":
                                        last_base = f"{float(val):.1f}"
                                        break
                                
                                for entry in reversed(h_data):
                                    val = entry.get('hr_thresh')
                                    if val is not None and val != "None":
                                        last_thresh = f"{float(val):.1f}"
                                        break
                                        
                    except Exception as e:
                        print(f"Błąd odczytu historii: {e}")
                
                st.markdown("### Aktualna Forma")
                k1, k2 = st.columns(2)
                k1.metric("Baza (280W)", f"{last_base} bpm", help="Oczekiwane tętno przy 280W @ 80rpm")
                k2.metric("Próg (360W)", f"{last_thresh} bpm", help="Oczekiwane tętno przy 360W @ 80rpm")

            with col_ai_2:
                # --- NOWY WYKRES DWULINIOWY (POPRAWIONY - OŚ X TO NUMER SESJI) ---
                if os.path.exists(HISTORY_FILE):
                    try:
                        with open(HISTORY_FILE, 'r') as f:
                            hist_data = json.load(f)
                        
                        if len(hist_data) > 0:
                            hist_df = pd.DataFrame(hist_data)
                            
                            hist_df = hist_df.reset_index()
                            hist_df['session_nr'] = hist_df.index + 1
                            
                            hover_text_base = hist_df.apply(lambda row: f"Plik: {row.get('source_file', 'N/A')}<br>Baza: {row['hr_base']:.1f} bpm", axis=1)
                            hover_text_thresh = hist_df.apply(lambda row: f"Plik: {row.get('source_file', 'N/A')}<br>Próg: {row['hr_thresh']:.1f} bpm", axis=1)

                            fig_evo = go.Figure()
                            
                            # Linia 1: Baza (280W)
                            fig_evo.add_trace(go.Scatter(
                                x=hist_df['session_nr'], 
                                y=hist_df['hr_base'], 
                                mode='lines+markers',
                                name='Baza (280W)',
                                line=dict(color='#00cc96', width=3), # Zielony
                                marker=dict(size=6),
                                hovertext=hover_text_base,
                                hoverinfo="text"
                            ))
                            
                            # Linia 2: Próg (360W)
                            fig_evo.add_trace(go.Scatter(
                                x=hist_df['session_nr'], 
                                y=hist_df['hr_thresh'], 
                                mode='lines+markers',
                                name='Próg (360W)',
                                line=dict(color='#ef553b', width=3), # Czerwony
                                marker=dict(size=6),
                                hovertext=hover_text_thresh,
                                hoverinfo="text"
                            ))
                            
                            fig_evo.update_layout(
                                template="plotly_dark",
                                title="Ewolucja Formy: Baza vs Próg",
                                xaxis_title="Kolejne Treningi (Sesja #)",
                                yaxis_title="HR [bpm] (Im niżej tym lepiej)",
                                hovermode="x unified",
                                legend=dict(orientation="h", y=1.1, x=0),
                                height=350
                            )
                            st.plotly_chart(fig_evo, use_container_width=True)
                    except Exception as e:
                        st.error(f"Błąd wykresu historii: {e}")

            st.divider()
            
            if 'ai_hr' in df_plot_resampled.columns:
                st.subheader("Analiza: Rzeczywistość vs AI")
                fig_ai_comp = go.Figure()
                fig_ai_comp.add_trace(go.Scatter(x=df_plot_resampled['time_min'], y=df_plot_resampled['heartrate_smooth'], 
                                             name='Rzeczywiste HR', line=dict(color='#ef553b', width=2)))
                fig_ai_comp.add_trace(go.Scatter(x=df_plot_resampled['time_min'], y=df_plot_resampled['ai_hr'], 
                                             name='AI Model HR (Oczekiwane)', line=dict(color='#00cc96', dash='dot', width=2)))
                
                fig_ai_comp.update_layout(template="plotly_dark", title="Czy serce reagowało zgodnie z planem?", hovermode="x unified")
                st.plotly_chart(fig_ai_comp, use_container_width=True)
                
                diff = df_plot_resampled['heartrate_smooth'] - df_plot_resampled['ai_hr']
                avg_diff = diff.mean()
                
                if avg_diff > 3:
                    st.warning(f"⚠️ **Wysoki Dryf Dnia (+{avg_diff:.1f} bpm):** Twoje tętno było wyższe niż model oczekiwał dla tej mocy. Możliwe zmęczenie, choroba lub upał.")
                elif avg_diff < -3:
                    st.success(f"✅ **Dzień Konia ({avg_diff:.1f} bpm):** Tętno niższe niż zazwyczaj. Świetna dyspozycja!")
                else:
                    st.info(f"🆗 **Norma ({avg_diff:.1f} bpm):** Reakcja serca zgodna z Twoim profilem historycznym.")

        else:
            st.warning("⚠️ Moduł AI wymaga procesora Apple Silicon i biblioteki `mlx`. Zainstaluj: `pip install mlx`")

    # --- TAB MODEL CP (PREDICTION) ---
    with tab_model:
        st.header("Matematyczny Model CP (Critical Power Estimation)")
        st.markdown("Estymacja Twojego CP i W' na podstawie krzywej mocy (MMP) z tego treningu. Używamy modelu liniowego: `Praca = CP * t + W'`.")

        if 'watts' in df_plot.columns and len(df_plot) > 1200: # Minimum 20 minut danych
            
            # 1. Wybieramy punkty czasowe do modelu (standardowe dla modelu 2-parametrowego)
            # Unikamy bardzo krótkich czasów (< 2-3 min), bo tam dominuje Pmax/AC
            durations = [180, 300, 600, 900, 1200] # 3min, 5min, 10min, 15min, 20min
            
            # Filtrujemy czasy dłuższe niż długość treningu
            valid_durations = [d for d in durations if d < len(df_plot)]
            
            if len(valid_durations) >= 3: # Potrzebujemy min. 3 punktów do sensownej regresji
                
                mmp_values = []
                work_values = []
                
                # Liczymy MMP i Pracę dla każdego punktu
                for d in valid_durations:
                    # Rolling mean max
                    p = df_plot['watts'].rolling(window=d).mean().max()
                    if not pd.isna(p):
                        mmp_values.append(p)
                        # Praca [J] = Moc [W] * Czas [s]
                        work_values.append(p * d)
                
                # 2. Regresja Liniowa (Work vs Time)
                # Y = Work, X = Time
                # Slope = CP, Intercept = W'
                slope, intercept, r_value, p_value, std_err = stats.linregress(valid_durations, work_values)
                
                modeled_cp = slope
                modeled_w_prime = intercept
                r_squared = r_value**2

                # 3. Wyświetlenie Wyników
                c_res1, c_res2, c_res3 = st.columns(3)
                
                c_res1.metric("Estymowane CP (z pliku)", f"{modeled_cp:.0f} W", 
                              delta=f"{modeled_cp - cp_input:.0f} W vs Ustawienia",
                              help="Moc Krytyczna wyliczona z Twoich najmocniejszych odcinków w tym pliku.")
                
                c_res2.metric("Estymowane W'", f"{modeled_w_prime:.0f} J",
                              delta=f"{modeled_w_prime - w_prime_input:.0f} J vs Ustawienia",
                              help="Pojemność beztlenowa wyliczona z modelu.")
                
                c_res3.metric("Jakość Dopasowania (R²)", f"{r_squared:.4f}", 
                              delta_color="normal" if r_squared > 0.98 else "inverse",
                              help="Jak bardzo Twoje wyniki pasują do teoretycznej krzywej. >0.98 = Bardzo wiarygodne.")

                st.markdown("---")

                # 4. Wizualizacja: Krzywa MMP vs Krzywa Modelowa
                # Generujemy punkty teoretyczne dla zakresu 1 min - 30 min
                x_theory = np.arange(60, 1800, 60) # co minutę
                y_theory = [modeled_cp + (modeled_w_prime / t) for t in x_theory]
                
                # Rzeczywiste MMP z pliku dla tych samych czasów
                y_actual = []
                x_actual = []
                for t in x_theory:
                    if t < len(df_plot):
                        val = df_plot['watts'].rolling(t).mean().max()
                        y_actual.append(val)
                        x_actual.append(t)

                fig_model = go.Figure()
                
                # Rzeczywiste MMP
                fig_model.add_trace(go.Scatter(
                    x=np.array(x_actual)/60, y=y_actual,
                    mode='markers', name='Twoje MMP (Actual)',
                    marker=dict(color='#00cc96', size=8)
                ))
                
                # Model Teoretyczny
                fig_model.add_trace(go.Scatter(
                    x=x_theory/60, y=y_theory,
                    mode='lines', name=f'Model CP ({modeled_cp:.0f}W)',
                    line=dict(color='#ef553b', dash='dash')
                ))

                fig_model.update_layout(
                    template="plotly_dark",
                    title="Power Duration Curve: Rzeczywistość vs Model",
                    xaxis_title="Czas trwania [min]",
                    yaxis_title="Moc [W]",
                    hovermode="x unified",
                    height=500
                )
                st.plotly_chart(fig_model, use_container_width=True)
                
                # 5. Interpretacja
                st.info(f"""
                **📊 Interpretacja Modelu:**
                
                Ten algorytm próbuje dopasować Twoje wysiłki do fizjologicznego prawa mocy krytycznej.
                
                * **Jeśli Estymowane CP > Ustawione CP:** Brawo! W tym treningu pokazałeś, że jesteś mocniejszy niż myślisz. Rozważ aktualizację ustawień w sidebarze.
                * **Jeśli Estymowane CP < Ustawione CP:** To normalne, jeśli nie jechałeś "do odciny" (All-Out) na odcinkach 3-20 min. Model pokazuje tylko to, co *zademonstrowałeś*, a nie Twój absolutny potencjał.
                * **R² (R-kwadrat):** Jeśli jest niskie (< 0.95), oznacza to, że Twoja jazda była nieregularna i model nie może znaleźć jednej linii, która pasuje do Twoich wyników.
                """)

            else:
                st.warning("Trening jest zbyt krótki lub brakuje mocnych odcinków, by zbudować wiarygodny model CP (wymagane wysiłki > 3 min i > 10 min).")
        else:
            st.warning("Za mało danych (wymagane min. 20 minut jazdy z pomiarem mocy).")

       # --- EXPORT DO PDF (Wersja CLEAN & STABLE) ---
from fpdf import FPDF
from datetime import datetime

# 1. Funkcja czyszcząca tekst (niezbędna dla FPDF bez zewnętrznych czcionek)
def clean_text(text):
    if text is None: return ""
    text = str(text)
    replacements = {
        'ą': 'a', 'ć': 'c', 'ę': 'e', 'ł': 'l', 'ń': 'n', 'ó': 'o', 'ś': 's', 'ź': 'z', 'ż': 'z',
        'Ą': 'A', 'Ć': 'C', 'Ę': 'E', 'Ł': 'L', 'Ń': 'N', 'Ó': 'O', 'Ś': 'S', 'Ź': 'Z', 'Ż': 'Z',
        '²': '2', '³': '3', '°': 'st.', '≈': '~', 'Δ': 'delta', 'Średnia': 'Srednia'
    }
    for k, v in replacements.items():
        text = text.replace(k, v)
    # Usuwamy znaki, których standardowy font nie ogarnie
    return text.encode('latin-1', 'replace').decode('latin-1')

def fmt_time(seconds):
    try:
        seconds = int(seconds)
        h = seconds // 3600
        m = (seconds % 3600) // 60
        s = seconds % 60
        if h > 0: return f"{h}h {m}m"
        return f"{m}m {s}s"
    except: return "-"

class ProPDF(FPDF):
    def header(self):
        # Pasek koloru na górze
        self.set_fill_color(0, 204, 150) # Streamlit Green
        self.rect(0, 0, 210, 5, 'F')
        
        self.ln(5)
        self.set_font('Arial', 'B', 18)
        self.set_text_color(40, 40, 40)
        self.cell(0, 10, 'RAPORT TRENINGOWY', 0, 1, 'L')
        
        self.set_font('Arial', '', 10)
        self.set_text_color(100, 100, 100)
        self.cell(0, 5, 'Pro Athlete Dashboard Analysis', 0, 1, 'L')
        
        self.ln(5)
        self.set_draw_color(200, 200, 200)
        self.line(10, 30, 200, 30)
        self.ln(10)

    def footer(self):
        self.set_y(-15)
        self.set_font('Arial', 'I', 8)
        self.set_text_color(150)
        self.cell(0, 10, f'Data generowania: {datetime.now().strftime("%Y-%m-%d %H:%M")} | Strona {self.page_no()}', 0, 0, 'C')

    def section_header(self, title):
        self.ln(5)
        self.set_font('Arial', 'B', 12)
        self.set_fill_color(240, 240, 240)
        self.set_text_color(0, 0, 0)
        # Pełna szerokość paska
        self.cell(0, 8, f"  {clean_text(title)}", 0, 1, 'L', 1)
        self.ln(3)

    def kpi_box(self, label, value, unit, x, y, w=45):
        # Tło
        self.set_xy(x, y)
        self.set_fill_color(255, 255, 255)
        self.set_draw_color(220, 220, 220)
        self.rect(x, y, w, 18, 'DF')
        
        # Label (Góra, mniejsza czcionka)
        self.set_xy(x, y+2)
        self.set_font('Arial', '', 7)
        self.set_text_color(100)
        self.cell(w, 4, clean_text(label), 0, 2, 'C')
        
        # Value (Środek, duża czcionka)
        self.set_font('Arial', 'B', 11)
        self.set_text_color(0)
        self.cell(w, 6, clean_text(value), 0, 2, 'C')
        
        # Unit (Dół, mała czcionka)
        self.set_font('Arial', '', 7)
        self.set_text_color(150)
        self.cell(w, 4, clean_text(unit), 0, 0, 'C')

    def create_table_row(self, data, widths, fill=False):
        """Tworzy wiersz tabeli z podanymi danymi i szerokościami kolumn."""
        self.set_font('Arial', '', 9)
        self.set_text_color(40, 40, 40)
        if fill:
            self.set_fill_color(245, 245, 245)
        else:
            self.set_fill_color(255, 255, 255)
        
        for i, datum in enumerate(data):
            w = widths[i] if i < len(widths) else widths[-1]
            self.cell(w, 7, clean_text(str(datum)), 1, 0, 'C', fill)
        self.ln()

    def create_table_header(self, headers, widths):
        """Tworzy nagłówek tabeli."""
        self.set_font('Arial', 'B', 9)
        self.set_fill_color(60, 60, 60)
        self.set_text_color(255, 255, 255)
        
        for i, header in enumerate(headers):
            w = widths[i] if i < len(widths) else widths[-1]
            self.cell(w, 8, clean_text(str(header)), 1, 0, 'C', 1)
        self.ln()
        self.set_text_color(0, 0, 0)

# ===== DOCX EXPORT BUTTON =====
st.sidebar.markdown("---")
st.sidebar.header("📄 Export Raportu")

if 'df_plot' in locals() and uploaded_file is not None:
    # Kolumny dla przycisków
    col_docx, col_pdf, col_png = st.sidebar.columns(3)
    
    with col_docx:
        # Generuj DOCX
        try:
            # AKTUALIZACJA: Dodano w_prime_input na końcu
            docx_doc = generate_docx_report(
                metrics, df_plot, df_plot_resampled, uploaded_file, cp_input,
                vt1_watts, vt2_watts, rider_weight, vt1_vent, vt2_vent, w_prime_input
            )
            
            # Zapisz do BytesIO
            docx_buffer = BytesIO()
            docx_doc.save(docx_buffer)
            docx_buffer.seek(0)
            
            st.sidebar.download_button(
                label="📥 Pobierz Raport DOCX",
                data=docx_buffer.getvalue(),
                file_name=f"Raport_{uploaded_file.name.split('.')[0]}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
        except Exception as e:
            st.sidebar.error(f"Błąd DOCX: {e}")
    
    with col_pdf:
        # Stary PDF button zostaje
        st.sidebar.info("PDF deprecated - używaj DOCX")
        
    with col_png:
    # Generuj PNG ZIP
        try:
            png_zip = export_all_charts_as_png(
                df_plot, df_plot_resampled, cp_input, vt1_watts, vt2_watts,
                metrics, rider_weight, uploaded_file
            )
            
            st.sidebar.download_button(
                label="📸 Pobierz Wykresy PNG (ZIP)",
                data=png_zip,
                file_name=f"Wykresy_{uploaded_file.name.split('.')[0]}.zip",
                mime="application/zip",
                use_container_width=True
            )
        except Exception as e:
            st.sidebar.error(f"Błąd PNG: {e}")

        
else:
    st.sidebar.info("Wgraj plik aby pobrać raport.")
# ===== KONIEC DOCX =====

st.sidebar.markdown("---")
st.sidebar.header("🖨️ Export Raportu")

if 'df_plot' in locals() and uploaded_file is not None:
    
    def generate_final_pdf():
        pdf = ProPDF()
        pdf.add_page()
        
        # --- OBLICZENIA POMOCNICZE (Pre-Calc) ---
        
        # 1. VO2Max (MMP 5min)
        vo2_max_est = 0
        if 'watts' in df_plot.columns:
            mmp_5m = df_plot['watts'].rolling(300).mean().max()
            if not pd.isna(mmp_5m):
                vo2_max_est = (10.8 * mmp_5m / rider_weight) + 7
        
        # 2. WĘGLOWODANY (Dokładna metoda strefowa)
        carbs_total = 0
        if 'watts' in df_plot.columns:
            energy_kcal_series = (df_plot['watts'] / 0.22) / 4184.0
            conditions = [
                (df_plot['watts'] < vt1_watts),
                (df_plot['watts'] >= vt1_watts) & (df_plot['watts'] < vt2_watts),
                (df_plot['watts'] >= vt2_watts)
            ]
            choices = [0.3, 0.8, 1.1]
            carb_fraction = np.select(conditions, choices, default=1.0)
            carbs_series = (energy_kcal_series * carb_fraction) / 4.0
            carbs_total = carbs_series.sum()

        # 3. Pulse Power Interpretacja
        pp_trend_txt = "Brak danych"
        pp_interpret = "N/A"
        if 'watts_smooth' in df_plot.columns and 'heartrate_smooth' in df_plot.columns:
            mask = (df_plot['watts_smooth'] > 50) & (df_plot['heartrate_smooth'] > 90)
            df_sub = df_plot[mask].copy()
            if not df_sub.empty:
                df_sub['pp'] = df_sub['watts_smooth'] / df_sub['heartrate_smooth']
                slope, intercept, _, _, _ = stats.linregress(np.arange(len(df_sub)), df_sub['pp'].values)
                trend_pct = ((intercept + slope*len(df_sub)) - intercept) / intercept * 100 if intercept != 0 else 0
                pp_trend_txt = f"{trend_pct:.1f}%"
                
                if trend_pct < -10: pp_interpret = "Krytyczny Dryf"
                elif trend_pct < -5: pp_interpret = "Umiarkowany Dryf"
                elif trend_pct > 5: pp_interpret = "Wzrost Wydajnosci"
                else: pp_interpret = "Stabilna Wydolnosc"

        # --- SEKCJA 1: PODSUMOWANIE (GRID 4x2) ---
        pdf.section_header("1. Podsumowanie Wykonania (KPI)")
        
        avg_w = metrics.get('avg_watts', 0)
        avg_hr = metrics.get('avg_hr', 0)
        work_kj = metrics.get('work_above_cp_kj', 0)
        
        start_x = 10
        start_y = pdf.get_y()
        w_box = 45
        gap = 2
        
        # Rząd 1
        pdf.kpi_box("Srednia Moc", f"{avg_w:.0f}", "W", start_x, start_y, w_box)
        pdf.kpi_box("Srednie HR", f"{avg_hr:.0f}", "bpm", start_x + w_box + gap, start_y, w_box)
        pdf.kpi_box("Szac. VO2Max", f"{vo2_max_est:.1f}", "ml/kg/min", start_x + (w_box + gap)*2, start_y, w_box)
        pdf.kpi_box("Spalone Wegle", f"{carbs_total:.0f}", "g", start_x + (w_box + gap)*3, start_y, w_box)
        
        # Rząd 2
        start_y += 20
        pdf.kpi_box("Decoupling (Dryf)", f"{decoupling_percent:.1f}", "%", start_x, start_y, w_box)
        pdf.kpi_box("Efficiency (EF)", f"{metrics.get('ef_factor', 0):.2f}", "W/bpm", start_x + w_box + gap, start_y, w_box)
        vent_val = metrics.get('avg_vent', 0)
        pdf.kpi_box("Wentylacja (VE)", f"{vent_val:.1f}", "L/min", start_x + (w_box + gap)*2, start_y, w_box)
        
        if 'smo2' in df_plot.columns:
            pdf.kpi_box("Srednie SmO2", f"{df_plot['smo2'].mean():.1f}", "%", start_x + (w_box + gap)*3, start_y, w_box)
        else:
            pdf.kpi_box("Srednie SmO2", "-", "%", start_x + (w_box + gap)*3, start_y, w_box)
            
        pdf.set_y(start_y + 25)

        # --- SEKCJA 2: ANALIZA FIZJOLOGICZNA (Tabela) ---
        pdf.section_header("2. Fizjologia i Pulse Power")
        
        pdf.set_fill_color(220)
        pdf.set_font('Arial', 'B', 9)
        pdf.cell(60, 8, "Parametr", 1, 0, 'L', 1)
        pdf.cell(40, 8, "Wartosc", 1, 0, 'C', 1)
        pdf.cell(90, 8, "Interpretacja / Status", 1, 1, 'L', 1)
        
        pdf.set_font('Arial', '', 9)
        
        pdf.cell(60, 8, clean_text("Pulse Power Trend"), 1)
        pdf.cell(40, 8, clean_text(pp_trend_txt), 1, 0, 'C')
        if "Stabilna" in pp_interpret: pdf.set_text_color(0, 150, 0)
        elif "Dryf" in pp_interpret: pdf.set_text_color(200, 0, 0)
        pdf.cell(90, 8, clean_text(pp_interpret), 1, 1, 'L')
        pdf.set_text_color(0)
        
        pdf.cell(60, 8, clean_text("Praca Beztlenowa (>CP)"), 1)
        pdf.cell(40, 8, f"{work_kj:.0f} kJ", 1, 0, 'C')
        w_status = "Bezpiecznie" if work_kj < w_prime_input else "Przekroczono W' (Ryzykowne)"
        pdf.cell(90, 8, clean_text(w_status), 1, 1, 'L')
        
        if 'core_temperature' in df_plot.columns:
            max_t = df_plot['core_temperature'].max()
            t_status = "Komfort" if max_t < 38.5 else ("Stres Cieplny" if max_t < 39.0 else "PRZEGRZANIE")
            pdf.cell(60, 8, clean_text("Temp. Maksymalna"), 1)
            pdf.cell(40, 8, f"{max_t:.2f} C", 1, 0, 'C')
            pdf.cell(90, 8, clean_text(t_status), 1, 1, 'L')

        pdf.ln(5)

        # --- SEKCJA 3: HRV & VT1 ESTIMATION ---
        pdf.section_header("3. HRV: Progi i Geometria")
        
        vt1_w_est = "-"
        vt1_hr_est = "-"
        alpha_avg_txt = "-"
        
        rr_col = next((c for c in df_clean_pl.columns if any(x in c.lower() for x in ['rr', 'hrv', 'ibi', 'r-r'])), None)
        
        if rr_col:
            try:
                temp_dfa, _ = calculate_dynamic_dfa(df_clean_pl, window_sec=120)
                if temp_dfa is not None and not temp_dfa.empty:
                    alpha_avg_txt = f"{temp_dfa['alpha1'].mean():.2f}"
                    
                    orig_times = df_plot['time'].values
                    orig_watts = df_plot['watts_smooth'].values
                    orig_hr = df_plot['heartrate_smooth'].values
                    dfa_times = temp_dfa['time'].values
                    dfa_watts = np.interp(dfa_times, orig_times, orig_watts)
                    dfa_hr = np.interp(dfa_times, orig_times, orig_hr)
                    
                    valid_mask = (temp_dfa['time'] > 300) & (temp_dfa['alpha1'] < 0.75)
                    stress_df = temp_dfa[valid_mask]
                    
                    if not stress_df.empty:
                        idx = stress_df.index[0]
                        vt1_w_est = f"{dfa_watts[idx]:.0f} W"
                        vt1_hr_est = f"{dfa_hr[idx]:.0f} bpm"
                    else:
                        vt1_w_est = "Alpha > 0.75"
                        vt1_hr_est = "-"
            except: pass

        pdf.set_fill_color(240)
        pdf.cell(50, 8, "Parametr HRV", 1, 0, 'L', 1)
        pdf.cell(40, 8, "Wynik", 1, 0, 'C', 1)
        pdf.cell(100, 8, "Opis", 1, 1, 'L', 1)
        
        pdf.cell(50, 8, "Est. VT1 (Moc)", 1)
        pdf.set_font('Arial', 'B', 9)
        pdf.cell(40, 8, clean_text(vt1_w_est), 1, 0, 'C')
        pdf.set_font('Arial', '', 8)
        pdf.cell(100, 8, "Moc przy ktorej Alpha-1 spada ponizej 0.75", 1, 1)
        
        pdf.set_font('Arial', '', 9)
        pdf.cell(50, 8, "Est. VT1 (HR)", 1)
        pdf.set_font('Arial', 'B', 9)
        pdf.cell(40, 8, clean_text(vt1_hr_est), 1, 0, 'C')
        pdf.set_font('Arial', '', 8)
        pdf.cell(100, 8, "Tetno na pierwszym progu wentylacyjnym", 1, 1)
        
        pdf.set_font('Arial', '', 9)
        pdf.cell(50, 8, "Srednie DFA Alpha-1", 1)
        pdf.cell(40, 8, alpha_avg_txt, 1, 0, 'C')
        pdf.set_font('Arial', '', 8)
        pdf.cell(100, 8, "Korelacja fraktalna (1.0 = Baza, 0.5 = Prog beztlenowy)", 1, 1)

        pdf.ln(5)

        # --- SEKCJA 4: CZAS W STREFACH ---
        pdf.section_header("4. Czas w Strefach")
        if 'watts' in df_plot.columns:
            pdf.set_fill_color(0, 204, 150)
            pdf.set_text_color(255)
            z_headers = ['Strefa', 'Zakres', 'Czas', '%']
            w_z = [60, 45, 45, 40]
            for i, h in enumerate(z_headers):
                pdf.cell(w_z[i], 7, h, 1, 0, 'C', 1)
            pdf.ln()
            
            pdf.set_text_color(0)
            zones = [
                ("Z1 Recovery", 0, 0.55*cp_input),
                ("Z2 Endurance", 0.56*cp_input, 0.75*cp_input),
                ("Z3 Tempo", 0.76*cp_input, 0.90*cp_input),
                ("Z4 Threshold", 0.91*cp_input, 1.05*cp_input),
                ("Z5 VO2Max", 1.06*cp_input, 1.20*cp_input),
                ("Z6 Anaerobic", 1.21*cp_input, 2000)
            ]
            total_t = len(df_plot)
            
            for i, (name, low, high) in enumerate(zones):
                count = len(df_plot[(df_plot['watts'] >= low) & (df_plot['watts'] < high)])
                pct = count/total_t*100 if total_t>0 else 0
                range_s = f"{low:.0f}-{high:.0f} W" if high < 1999 else f"> {low:.0f} W"
                fill = (i % 2 == 1)
                pdf.create_table_row([name, range_s, fmt_time(count), f"{pct:.1f}%"], w_z, fill=fill)

        # Pomocniczy parser (lokalny dla pewności)
        def _local_parse(t_str):
            try:
                parts = list(map(int, t_str.split(':')))
                if len(parts) == 3: return parts[0]*3600 + parts[1]*60 + parts[2]
                if len(parts) == 2: return parts[0]*60 + parts[1]
                if len(parts) == 1: return parts[0]
            except: return None
            return None

        # --- SEKCJA 5: ANALIZA ODCINKA (SmO2) ---
        pdf.ln(5)
        pdf.section_header("5. Analiza Wybranego Odcinka (SmO2)")
        
        s_int_str = start_time_str if 'start_time_str' in globals() else "Brak"
        e_int_str = end_time_str if 'end_time_str' in globals() else "Brak"
        
        s_sec_val = _local_parse(s_int_str)
        e_sec_val = _local_parse(e_int_str)
        
        found_smo2 = False
        if s_sec_val is not None and e_sec_val is not None:
            col_s = 'smo2_smooth' if 'smo2_smooth' in df_plot.columns else ('smo2' if 'smo2' in df_plot.columns else None)
            col_w = 'watts_smooth' if 'watts_smooth' in df_plot.columns else 'watts'
            
            if col_s and col_w:
                mask_int = (df_plot['time'] >= s_sec_val) & (df_plot['time'] <= e_sec_val)
                df_int = df_plot[mask_int].copy()
                if not df_int.empty:
                    found_smo2 = True
                    dur_int = e_sec_val - s_sec_val
                    
                    avg_w_int = df_int[col_w].mean()
                    avg_s_int = df_int[col_s].mean()
                    min_s_int = df_int[col_s].min()
                    
                    slope_s, _, _, _, _ = stats.linregress(df_int['time'], df_int[col_s])
                    
                    pdf.set_font('Arial', '', 10)
                    pdf.cell(0, 5, f"Zakres: {s_int_str} - {e_int_str} (Czas: {dur_int}s)", 0, 1)
                    start_x = 10; start_y = pdf.get_y() + 5; w_box = 45; gap = 2
                    
                    pdf.kpi_box("Srednia Moc", f"{avg_w_int:.0f}", "W", start_x, start_y, w_box)
                    pdf.kpi_box("Srednie SmO2", f"{avg_s_int:.1f}", "%", start_x + w_box + gap, start_y, w_box)
                    pdf.kpi_box("Min SmO2", f"{min_s_int:.1f}", "%", start_x + (w_box + gap)*2, start_y, w_box)
                    pdf.kpi_box("Slope SmO2", f"{slope_s:.4f}", "%/s", start_x + (w_box + gap)*3, start_y, w_box)
                    pdf.ln(25)
        
        if not found_smo2:
            pdf.set_font('Arial', 'I', 9)
            pdf.cell(0, 10, clean_text("Brak danych lub nie wybrano odcinka w zakladce SmO2."), 0, 1)

        # --- SEKCJA 6: ANALIZA ODCINKA (WENTYLACJA) ---
        # NOWOŚĆ: To jest ten blok, o który prosiłeś
        pdf.section_header("6. Analiza Wybranego Odcinka (Wentylacja)")
        
        # Pobieramy zmienne z zakładki Vent (muszą być globalne)
        s_vent_str = start_time_v if 'start_time_v' in globals() else "Brak"
        e_vent_str = end_time_v if 'end_time_v' in globals() else "Brak"
        
        s_v_sec = _local_parse(s_vent_str)
        e_v_sec = _local_parse(e_vent_str)
        
        found_vent = False
        if s_v_sec is not None and e_v_sec is not None:
            col_ve = 'tymeventilation' if 'tymeventilation' in df_plot.columns else None
            col_w = 'watts' if 'watts' in df_plot.columns else None
            
            if col_ve and col_w:
                mask_v = (df_plot['time'] >= s_v_sec) & (df_plot['time'] <= e_v_sec)
                df_v = df_plot[mask_v].copy()
                
                if not df_v.empty:
                    found_vent = True
                    dur_v = e_v_sec - s_v_sec
                    
                    avg_w_v = df_v[col_w].mean()
                    avg_ve_v = df_v[col_ve].mean()
                    ve_w_ratio = avg_ve_v / avg_w_v if avg_w_v > 0 else 0
                    
                    slope_ve, _, _, _, _ = stats.linregress(df_v['time'], df_v[col_ve])
                    
                    pdf.set_font('Arial', '', 10)
                    pdf.cell(0, 5, f"Zakres: {s_vent_str} - {e_vent_str} (Czas: {dur_v}s)", 0, 1)
                    
                    start_x = 10; start_y = pdf.get_y() + 5; w_box = 45; gap = 2
                    
                    pdf.kpi_box("Srednia Moc", f"{avg_w_v:.0f}", "W", start_x, start_y, w_box)
                    pdf.kpi_box("Srednie VE", f"{avg_ve_v:.1f}", "L/min", start_x + w_box + gap, start_y, w_box)
                    pdf.kpi_box("Wydajnosc (VE/W)", f"{ve_w_ratio:.3f}", "L/W", start_x + (w_box + gap)*2, start_y, w_box)
                    pdf.kpi_box("Slope VE", f"{slope_ve:.4f}", "L/s", start_x + (w_box + gap)*3, start_y, w_box)
                    
                    pdf.ln(25)

        if not found_vent:
            pdf.set_font('Arial', 'I', 9)
            pdf.cell(0, 10, clean_text("Brak danych lub nie wybrano odcinka w zakladce Ventilation."), 0, 1)

        return pdf.output(dest='S').encode('latin-1', 'replace')

    pdf_bytes = generate_final_pdf()
    
    st.sidebar.download_button(
        label="📄 Pobierz Raport PRO (PDF)",
        data=pdf_bytes,
        file_name=f"Raport_PRO_{uploaded_file.name.split('.')[0]}.pdf",
        mime="application/pdf"
    )
else:
    st.sidebar.info("Wgraj plik aby pobrac PDF.")