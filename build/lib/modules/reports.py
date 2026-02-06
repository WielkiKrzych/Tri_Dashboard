from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import zipfile
from io import BytesIO
import numpy as np
import pandas as pd
from datetime import datetime
from typing import Dict, Any, Optional, Tuple, List


def _calculate_fallback_metrics(metrics: Dict[str, Any], df_plot: pd.DataFrame) -> Dict[str, Any]:
    """
    Calculate fallback metrics if not provided in metrics dict.

    Args:
        metrics: Dictionary of pre-calculated metrics
        df_plot: DataFrame with training data

    Returns:
        Updated metrics dictionary with fallback values
    """
    import logging

    result = metrics.copy()

    # Normalized Power (NP)
    if result.get("np", 0) == 0 and "watts" in df_plot.columns:
        logging.warning("NP not pre-calculated - using fallback calculation")
        rolling_30s = df_plot["watts"].rolling(window=30, min_periods=1).mean()
        result["np"] = np.power(np.mean(np.power(rolling_30s, 4)), 0.25)

    # Total work
    if result.get("work_kj", 0) == 0 and "watts" in df_plot.columns:
        logging.warning("work_kj not pre-calculated - using fallback")
        result["work_kj"] = df_plot["watts"].sum() / 1000

    # Core temperature
    if result.get("max_core", 0) == 0 and "core_temperature" in df_plot.columns:
        result["max_core"] = df_plot["core_temperature"].max()
        result["avg_core"] = df_plot["core_temperature"].mean()

    # Heat Strain Index
    if result.get("max_hsi", 0) == 0 and "hsi" in df_plot.columns:
        result["max_hsi"] = df_plot["hsi"].max()

    return result


def _add_header_section(doc: Document, uploaded_file: Any) -> None:
    """Add report header with title, date and source."""
    # Style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # Title
    title = doc.add_heading("Pro Athlete Dashboard - Raport Treningowy", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Date
    date_para = doc.add_paragraph(f"Data: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.runs[0].font.color.rgb = RGBColor(100, 100, 100)

    # Source file
    source_para = doc.add_paragraph(f"Plik: {uploaded_file.name if uploaded_file else 'Brak'}")
    source_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()


def _add_kpi_section(doc: Document, metrics: Dict[str, Any], df_plot: pd.DataFrame) -> None:
    """Add KPI summary section with metrics table."""
    doc.add_heading("1. Podsumowanie KPI", level=1)

    # Prepare KPI data
    kpi_data = [
        ("Średnia Moc", f"{metrics.get('avg_watts', 0):.0f} W"),
        ("Normalized Power (NP)", f"{metrics.get('np', 0):.0f} W"),
        ("Praca Całkowita", f"{metrics.get('work_kj', 0):.0f} kJ"),
        ("Średnie Tętno", f"{metrics.get('avg_hr', 0):.0f} bpm"),
        ("Średnia Kadencja", f"{metrics.get('avg_cadence', 0):.0f} rpm"),
        ("Średnia Wentylacja (VE)", f"{metrics.get('avg_vent', 0):.1f} L/min"),
        ("Średnie Oddechy (RR)", f"{metrics.get('avg_rr', 0):.1f} /min"),
        ("Średnie SmO2", f"{df_plot['smo2'].mean() if 'smo2' in df_plot.columns else 0:.1f}%"),
        ("Min SmO2", f"{df_plot['smo2'].min() if 'smo2' in df_plot.columns else 0:.1f}%"),
        ("Max SmO2", f"{df_plot['smo2'].max() if 'smo2' in df_plot.columns else 0:.1f}%"),
        ("Power/HR", f"{metrics.get('power_hr', 0):.2f}"),
        ("Efficiency Factor (EF)", f"{metrics.get('ef_factor', 0):.2f}"),
        ("Średnie Pulse Power", f"{metrics.get('avg_pp', 0):.2f} W/bpm"),
        ("Średnie RMSSD", f"{metrics.get('avg_rmssd', 0):.1f} ms"),
        ("Max HSI (Indeks Ciepła)", f"{metrics.get('max_hsi', 0):.1f}"),
        ("Max Temperatura Ciała", f"{metrics.get('max_core', 0):.2f} °C"),
        ("Średnia Temperatura Ciała", f"{metrics.get('avg_core', 0):.2f} °C"),
        ("Spalone Węglowodany", f"{metrics.get('carbs_total', 0):.0f} g"),
    ]

    # Create table
    table = doc.add_table(rows=len(kpi_data) + 1, cols=2)
    table.style = "Light Grid Accent 1"

    # Headers
    table.rows[0].cells[0].text = "Metryka"
    table.rows[0].cells[1].text = "Wartość"

    # Fill data
    for i, (label, val) in enumerate(kpi_data):
        row_cells = table.rows[i + 1].cells
        row_cells[0].text = label
        row_cells[1].text = val

    doc.add_paragraph()


def _add_thresholds_section(
    doc: Document,
    metrics: Dict[str, Any],
    vt1_watts: float,
    vt2_watts: float,
    vt1_vent: float,
    vt2_vent: float,
    cp_input: float,
    w_prime_input: float,
    rider_weight: float,
) -> None:
    """Add training thresholds and physiological parameters section."""
    doc.add_heading("2. Progi i Parametry Fizjologiczne", level=1)

    p = doc.add_paragraph()
    p.add_run(f"VT1 (Próg Tlenowy): ").bold = True
    p.add_run(f"{vt1_watts} W @ {vt1_vent} L/min\n")

    p.add_run(f"VT2 (Próg Beztlenowy): ").bold = True
    p.add_run(f"{vt2_watts} W @ {vt2_vent} L/min\n")

    p.add_run(f"CP (Moc Krytyczna): ").bold = True
    p.add_run(f"{cp_input} W\n")

    p.add_run(f"W' (Pojemność Beztlenowa): ").bold = True
    p.add_run(f"{w_prime_input} J\n")

    vo2_max = metrics.get("vo2_max_est", 0)
    p.add_run(f"Szacunkowe VO2max (MMP 5'): ").bold = True
    p.add_run(f"{vo2_max:.1f} ml/kg/min\n")

    p.add_run(f"Waga zawodnika: ").bold = True
    p.add_run(f"{rider_weight} kg")

    doc.add_paragraph()


def _add_power_zones_section(doc: Document, df_plot: pd.DataFrame, cp_input: float) -> None:
    """Add power zones time distribution section."""
    doc.add_heading("3. Czas w Strefach Mocy", level=1)

    if "watts" not in df_plot.columns:
        doc.add_paragraph("Brak danych mocy.")
        return

    # Zone definitions
    bins = [
        0,
        0.55 * cp_input,
        0.75 * cp_input,
        0.90 * cp_input,
        1.05 * cp_input,
        1.20 * cp_input,
        10000,
    ]
    labels = [
        "Z1 Recovery",
        "Z2 Endurance",
        "Z3 Tempo",
        "Z4 Threshold",
        "Z5 VO2Max",
        "Z6 Anaerobic",
    ]

    # Calculate time in zones
    dfz = df_plot.copy()
    dfz["Zone"] = pd.cut(dfz["watts"], bins=bins, labels=labels, right=False)

    # Create table
    table = doc.add_table(rows=7, cols=3)
    table.style = "Light Grid Accent 1"

    table.rows[0].cells[0].text = "Strefa"
    table.rows[0].cells[1].text = "Zakres"
    table.rows[0].cells[2].text = "Czas"

    for i, (label, low, high) in enumerate(zip(labels, bins[:-1], bins[1:])):
        count = len(dfz[dfz["Zone"] == label])
        time_min = count / 60

        # Display formatting
        display_high = 2000 if high == 10000 else int(high)
        display_low = int(low)

        table.rows[i + 1].cells[0].text = label
        table.rows[i + 1].cells[1].text = f"{display_low}-{display_high} W"
        table.rows[i + 1].cells[2].text = f"{time_min:.1f} min"

    doc.add_paragraph()


def _add_notes_section(doc: Document) -> None:
    """Add coach/athlete notes section."""
    doc.add_heading("4. Notatki Trenera / Zawodnika", level=1)

    # Editable space
    note_p = doc.add_paragraph("[Miejsce na Twoje notatki...]")
    note_p.runs[0].font.italic = True
    note_p.runs[0].font.color.rgb = RGBColor(150, 150, 150)

    # Empty lines for visual space
    for _ in range(5):
        doc.add_paragraph("")


def _add_footer(doc: Document) -> None:
    """Add report footer."""
    doc.add_paragraph("---")
    footer = doc.add_paragraph("Raport wygenerowany przez Pro Athlete Dashboard | Streamlit App")
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


def generate_docx_report(
    metrics: Dict[str, Any],
    df_plot: pd.DataFrame,
    df_plot_resampled: pd.DataFrame,
    uploaded_file: Any,
    cp_input: float,
    vt1_watts: float,
    vt2_watts: float,
    rider_weight: float,
    vt1_vent: float,
    vt2_vent: float,
    w_prime_input: float,
) -> Optional[Document]:
    """
    Generuje raport .docx z rozszerzonymi KPI, W', VO2max i czystymi notatkami.

    Args:
        metrics: Dictionary of pre-calculated training metrics
        df_plot: DataFrame with training data
        df_plot_resampled: Resampled DataFrame for charts
        uploaded_file: Uploaded file object
        cp_input: Critical Power in watts
        vt1_watts: VT1 power in watts
        vt2_watts: VT2 power in watts
        rider_weight: Rider weight in kg
        vt1_vent: VT1 ventilation in L/min
        vt2_vent: VT2 ventilation in L/min
        w_prime_input: W' capacity in joules

    Returns:
        Document object or None if generation not requested
    """
    import streamlit as st

    if not st.session_state.get("report_generation_requested", False):
        return None

    # Calculate fallback metrics if needed
    metrics = _calculate_fallback_metrics(metrics, df_plot)

    # Create document
    doc = Document()

    # Add sections
    _add_header_section(doc, uploaded_file)
    _add_kpi_section(doc, metrics, df_plot)
    _add_thresholds_section(
        doc,
        metrics,
        vt1_watts,
        vt2_watts,
        vt1_vent,
        vt2_vent,
        cp_input,
        w_prime_input,
        rider_weight,
    )
    _add_power_zones_section(doc, df_plot, cp_input)
    _add_notes_section(doc)
    _add_footer(doc)

    return doc


def export_all_charts_as_png(
    df_plot,
    df_plot_resampled,
    cp_input,
    vt1_watts,
    vt2watts,
    metrics,
    rider_weight,
    uploaded_file,
    vent_start_sec,
    vent_end_sec,
    smo2_start_sec,
    smo2_end_sec,
):
    """
    Export wykresów PNG z pełną legendą statystyczną (Ghost Traces).

    SOLID (OCP): Używa wzorca Registry - nowe wykresy można dodawać
    jako klasy w chart_exporters.py bez modyfikacji tej funkcji.
    """
    import streamlit as st

    if not st.session_state.get("report_generation_requested", False):
        return None

    from .chart_exporters import CHART_REGISTRY, ChartContext

    zip_buffer = BytesIO()

    # Utwórz kontekst danych dla eksporterów (ISP - jeden obiekt zamiast wielu parametrów)
    ctx = ChartContext(
        df_plot=df_plot,
        df_plot_resampled=df_plot_resampled,
        rider_weight=rider_weight,
        cp_input=cp_input,
        vt1_watts=vt1_watts,
        vt2_watts=vt2watts,
        metrics=metrics,
        smo2_start_sec=smo2_start_sec,
        smo2_end_sec=smo2_end_sec,
        vent_start_sec=vent_start_sec,
        vent_end_sec=vent_end_sec,
    )

    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zipf:
        # SOLID (OCP): Iteracja po zarejestrowanych eksporterach
        # Dodanie nowego wykresu = tylko nowa klasa z dekoratorem @register_chart
        for exporter in CHART_REGISTRY:
            try:
                if exporter.can_export(ctx):
                    png_bytes = exporter.export(ctx)
                    zipf.writestr(exporter.filename, png_bytes)
            except Exception as e:
                # Loguj błąd ale kontynuuj z innymi wykresami
                print(f"Error exporting {exporter.filename}: {e}")

        # README
        readme = f"""RAPORT WYKRESÓW Z PEŁNĄ ANALIZĄ
Data: {datetime.now().strftime("%Y-%m-%d %H:%M")}
Plik: {uploaded_file.name}

Wszystkie wykresy zawierają statystyki w legendzie (Avg, Min, Max, Trends).
Wykresy 9 i 10 zawierają analizę odcinków wybranych w aplikacji.

SOLID: Ten eksport używa wzorca Registry.
Dodawanie nowych wykresów: utwórz klasę w modules/chart_exporters.py
"""
        zipf.writestr("00_README.txt", readme)

    zip_buffer.seek(0)
    return zip_buffer.getvalue()
