from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import zipfile
import io
from io import BytesIO
import numpy as np
import pandas as pd
import plotly.graph_objects as go
from scipy import stats
from datetime import datetime
from .plots import add_stats_to_legend

def generate_docx_report(metrics, df_plot, df_plot_resampled, uploaded_file, cp_input, 
                        vt1_watts, vt2_watts, rider_weight, vt1_vent, vt2_vent, w_prime_input):
    """
    Generuje raport .docx z rozszerzonymi KPI, W', VO2max i czystymi notatkami.
    """
    
    doc = Document()
    
    # --- OBLICZENIA POBIERANE Z METRICS (Code Centralization) ---
    # Wszystkie kluczowe obliczenia powinny być wykonane w app.py lub modules/calculations.py
    # i przekazane w słowniku metrics.
    
    np_val = metrics.get('np', 0)
    total_work_kj = metrics.get('work_kj', 0)
    avg_pp = metrics.get('avg_pp', 0)
    max_hsi = metrics.get('max_hsi', 0)
    max_core = metrics.get('max_core', 0)
    avg_core = metrics.get('avg_core', 0)
    avg_rmssd = metrics.get('avg_rmssd', 0)
    carbs_total = metrics.get('carbs_total', 0)
    vo2_max_est = metrics.get('vo2_max_est', 0)

    # Metrics should be pre-calculated in app.py
    # Fallback warnings for debugging if metrics are missing
    import logging
    
    if np_val == 0 and 'watts' in df_plot.columns:
        logging.warning("NP not pre-calculated in app.py - using fallback calculation")
        rolling_30s = df_plot['watts'].rolling(window=30, min_periods=1).mean()
        np_val = np.power(np.mean(np.power(rolling_30s, 4)), 0.25)

    if total_work_kj == 0 and 'watts' in df_plot.columns:
        logging.warning("work_kj not pre-calculated in app.py - using fallback")
        total_work_kj = df_plot['watts'].sum() / 1000

    if carbs_total == 0:
        logging.warning("carbs_total not pre-calculated in app.py")

    if max_core == 0 and 'core_temperature' in df_plot.columns:
        max_core = df_plot['core_temperature'].max()
        avg_core = df_plot['core_temperature'].mean()

    if max_hsi == 0 and 'hsi' in df_plot.columns:
        max_hsi = df_plot['hsi'].max()
    
    # --- END FALLBACK CALCULATIONS ---

    # STYLE
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    
    # NAGŁÓWEK
    title = doc.add_heading('Pro Athlete Dashboard - Raport Treningowy', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    date_para = doc.add_paragraph(f"Data: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.runs[0].font.color.rgb = RGBColor(100, 100, 100)
    
    source_para = doc.add_paragraph(f"Plik: {uploaded_file.name if uploaded_file else 'Brak'}")
    source_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph() 
    
    # SEKCJA 1: PODSUMOWANIE KPI (ROZSZERZONE)
    doc.add_heading('1. Podsumowanie KPI', level=1)
    
    # Tabela z dużą ilością danych (18 wierszy)
    kpi_data = [
        ("Średnia Moc", f"{metrics.get('avg_watts', 0):.0f} W"),
        ("Normalized Power (NP)", f"{np_val:.0f} W"),
        ("Praca Całkowita", f"{total_work_kj:.0f} kJ"),
        ("Średnie Tętno", f"{metrics.get('avg_hr', 0):.0f} bpm"),
        ("Średnia Kadencja", f"{metrics.get('avg_cadence', 0):.0f} rpm"),
        ("Średnia Wentylacja (VE)", f"{metrics.get('avg_vent', 0):.1f} L/min"),
        ("Średnie Oddechy (RR)", f"{metrics.get('avg_rr', 0):.1f} /min"),
        ("Średnie SmO2", f"{df_plot['smo2'].mean() if 'smo2' in df_plot.columns else 0:.1f}%"),
        ("Min SmO2", f"{df_plot['smo2'].min() if 'smo2' in df_plot.columns else 0:.1f}%"),
        ("Max SmO2", f"{df_plot['smo2'].max() if 'smo2' in df_plot.columns else 0:.1f}%"),
        ("Power/HR", f"{metrics.get('power_hr', 0):.2f}"),
        ("Efficiency Factor (EF)", f"{metrics.get('ef_factor', 0):.2f}"),
        ("Średnie Pulse Power", f"{avg_pp:.2f} W/bpm"),
        ("Średnie RMSSD", f"{avg_rmssd:.1f} ms"),
        ("Max HSI (Indeks Ciepła)", f"{max_hsi:.1f}"),
        ("Max Temperatura Ciała", f"{max_core:.2f} °C"),
        ("Średnia Temperatura Ciała", f"{avg_core:.2f} °C"),
        ("Spalone Węglowodany", f"{carbs_total:.0f} g")
    ]

    table = doc.add_table(rows=len(kpi_data)+1, cols=2)
    table.style = 'Light Grid Accent 1'
    
    # Nagłówki tabeli
    table.rows[0].cells[0].text = 'Metryka'
    table.rows[0].cells[1].text = 'Wartość'
    
    # Wypełnianie danymi
    for i, (label, val) in enumerate(kpi_data):
        row_cells = table.rows[i+1].cells
        row_cells[0].text = label
        row_cells[1].text = val
    
    doc.add_paragraph()
    
    # SEKCJA 2: PROGI TRENINGOWE (DODANO W' i VO2max)
    doc.add_heading('2. Progi i Parametry Fizjologiczne', level=1)
    
    p = doc.add_paragraph()
    p.add_run(f"VT1 (Próg Tlenowy): ").bold = True
    p.add_run(f"{vt1_watts} W @ {vt1_vent} L/min\n")
    
    p.add_run(f"VT2 (Próg Beztlenowy): ").bold = True
    p.add_run(f"{vt2_watts} W @ {vt2_vent} L/min\n")
    
    p.add_run(f"CP (Moc Krytyczna): ").bold = True
    p.add_run(f"{cp_input} W\n")
    
    p.add_run(f"W' (Pojemność Beztlenowa): ").bold = True
    p.add_run(f"{w_prime_input} J\n")
    
    p.add_run(f"Szacunkowe VO2max (MMP 5'): ").bold = True
    p.add_run(f"{vo2_max_est:.1f} ml/kg/min\n")
    
    p.add_run(f"Waga zawodnika: ").bold = True
    p.add_run(f"{rider_weight} kg")
    
    doc.add_paragraph()
    
    # SEKCJA 3: STREFY MOCY (POPRAWKA Z6)
    doc.add_heading('3. Czas w Strefach Mocy', level=1)
    
    if 'watts' in df_plot.columns:
        # Definicja stref
        bins = [0, 0.55*cp_input, 0.75*cp_input, 0.90*cp_input, 1.05*cp_input, 1.20*cp_input, 10000]
        labels = ['Z1 Recovery', 'Z2 Endurance', 'Z3 Tempo', 'Z4 Threshold', 'Z5 VO2Max', 'Z6 Anaerobic']
        
        # Obliczenia
        dfz = df_plot.copy()
        dfz['Zone'] = pd.cut(dfz['watts'], bins=bins, labels=labels, right=False)
        
        table2 = doc.add_table(rows=7, cols=3)
        table2.style = 'Light Grid Accent 1'
        
        table2.rows[0].cells[0].text = 'Strefa'
        table2.rows[0].cells[1].text = 'Zakres'
        table2.rows[0].cells[2].text = 'Czas'
        
        for i, (label, low, high) in enumerate(zip(labels, bins[:-1], bins[1:])):
            count = len(dfz[dfz['Zone'] == label])
            time_min = count / 60
            
            # POPRAWKA: Wyświetlanie 2000 W zamiast 10000 W dla estetyki
            display_high = 2000 if high == 10000 else int(high)
            display_low = int(low)
            
            table2.rows[i+1].cells[0].text = label
            table2.rows[i+1].cells[1].text = f"{display_low}-{display_high} W"
            table2.rows[i+1].cells[2].text = f"{time_min:.1f} min"
    
    doc.add_paragraph()
    
    # SEKCJA 4: NOTATKI (WYCZYSZCZONE)
    doc.add_heading('4. Notatki Trenera / Zawodnika', level=1)
    
    # Pusty paragraf jako przestrzeń edytowalna w Pages/Word
    note_p = doc.add_paragraph("[Miejsce na Twoje notatki...]")
    note_p.runs[0].font.italic = True
    note_p.runs[0].font.color.rgb = RGBColor(150, 150, 150)
    
    # Dodajemy kilka pustych linii, żeby wizualnie zrobić miejsce, ale bez "___"
    for _ in range(5):
        doc.add_paragraph("")
    
    # STOPKA
    doc.add_paragraph("---")
    footer = doc.add_paragraph("Raport wygenerowany przez Pro Athlete Dashboard | Streamlit App")
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    return doc

def export_all_charts_as_png(df_plot, df_plot_resampled, cp_input, vt1_watts, vt2watts,
                            metrics, rider_weight, uploaded_file, vent_start_sec, vent_end_sec, smo2_start_sec, smo2_end_sec):
    """
    Export wykresów PNG z pełną legendą statystyczną (Ghost Traces).
    
    SOLID (OCP): Używa wzorca Registry - nowe wykresy można dodawać
    jako klasy w chart_exporters.py bez modyfikacji tej funkcji.
    """
    from .chart_exporters import CHART_REGISTRY, ChartContext
    
    zip_buffer = BytesIO()
    
    # Utwórz kontekst danych dla eksporterów (ISP - jeden obiekt zamiast wielu parametrów)
    ctx = ChartContext(
        df_plot=df_plot,
        df_plot_resampled=df_plot_resampled,
        rider_weight=rider_weight,
        cp_input=cp_input,
        vt1_watts=vt1_watts,
        vt2_watts=vt2watts,
        metrics=metrics,
        smo2_start_sec=smo2_start_sec,
        smo2_end_sec=smo2_end_sec,
        vent_start_sec=vent_start_sec,
        vent_end_sec=vent_end_sec
    )

    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zipf:
        
        # SOLID (OCP): Iteracja po zarejestrowanych eksporterach
        # Dodanie nowego wykresu = tylko nowa klasa z dekoratorem @register_chart
        for exporter in CHART_REGISTRY:
            try:
                if exporter.can_export(ctx):
                    png_bytes = exporter.export(ctx)
                    zipf.writestr(exporter.filename, png_bytes)
            except Exception as e:
                # Loguj błąd ale kontynuuj z innymi wykresami
                print(f"Error exporting {exporter.filename}: {e}")

        # README
        readme = f"""RAPORT WYKRESÓW Z PEŁNĄ ANALIZĄ
Data: {datetime.now().strftime('%Y-%m-%d %H:%M')}
Plik: {uploaded_file.name}

Wszystkie wykresy zawierają statystyki w legendzie (Avg, Min, Max, Trends).
Wykresy 9 i 10 zawierają analizę odcinków wybranych w aplikacji.

SOLID: Ten eksport używa wzorca Registry.
Dodawanie nowych wykresów: utwórz klasę w modules/chart_exporters.py
"""
        zipf.writestr('00_README.txt', readme)
    
    zip_buffer.seek(0)
    return zip_buffer.getvalue()
