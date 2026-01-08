"""
DOCX Builder Module (v3.0 - Premium Styling).

Generates editable Word documents for Ramp Test reports.
Structure matches PDF 1:1 with COACH NOTES after each section.

Key Features:
- Premium title page with DarkGlass style
- Section numbering matching PDF
- Colored backgrounds: Navy (recommendations), Red (warnings), Green (positives)
- Author and contact info
- Single Source of Truth: uses map_ramp_json_to_pdf_data()
"""
import os
import logging
from typing import Dict, Any, Optional, List
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm, Twips
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import OxmlElement, parse_xml
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

from .pdf.builder import map_ramp_json_to_pdf_data

logger = logging.getLogger("Tri_Dashboard.DOCXBuilder")

# =============================================================================
# COLOR CONSTANTS
# =============================================================================
COLORS = {
    "navy": RGBColor(26, 82, 118),       # #1A5276 - recommendations/training
    "dark_blue": RGBColor(23, 37, 42),   # #17252A - DarkGlass header
    "red": RGBColor(192, 57, 43),        # #C0392B - warnings/limitations
    "green": RGBColor(39, 174, 96),      # #27AE60 - positives/strengths
    "gray": RGBColor(127, 140, 141),     # #7F8C8D - neutral
    "white": RGBColor(255, 255, 255),
    "light_gray": RGBColor(236, 240, 241),
}


# =============================================================================
# HELPER FUNCTIONS
# =============================================================================

def _set_cell_shading(cell, color_hex: str):
    """Set cell background color using shading."""
    shading_elm = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading_elm)


def _add_colored_box(doc, text: str, bg_color: str = "navy", text_color: RGBColor = None):
    """Add a text box with colored background using a single-cell table.
    
    Args:
        doc: Document object
        text: Text content
        bg_color: "navy", "red", "green" or hex color
        text_color: RGBColor for text (default: white)
    """
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    
    # Set background color
    color_map = {
        "navy": "1A5276",
        "red": "C0392B", 
        "green": "27AE60",
        "dark": "17252A",
        "gray": "7F8C8D",
    }
    hex_color = color_map.get(bg_color, bg_color.lstrip("#"))
    _set_cell_shading(cell, hex_color)
    
    # Add text
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = text_color or COLORS["white"]
    
    # Padding
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    
    doc.add_paragraph()  # Spacing


def _add_coach_notes(doc, title: str = "📝 NOTATKI TRENERA", lines: int = 6):
    """Add Coach Notes block after each section."""
    p_sep = doc.add_paragraph()
    p_sep.add_run("─" * 60)
    p_sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p_title = doc.add_paragraph()
    run_title = p_title.add_run(title)
    run_title.bold = True
    run_title.font.size = Pt(11)
    run_title.font.color.rgb = COLORS["navy"]
    
    p_instr = doc.add_paragraph()
    run_instr = p_instr.add_run("[Miejsce na Twoje obserwacje i komentarze]")
    run_instr.italic = True
    run_instr.font.color.rgb = COLORS["gray"]
    
    for _ in range(lines):
        doc.add_paragraph()
    
    p_sep2 = doc.add_paragraph()
    p_sep2.add_run("─" * 60)
    p_sep2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


def _add_section_header(doc, number: int, title: str, description: str = ""):
    """Add numbered section header with description.
    
    Args:
        doc: Document object
        number: Section number (1-16)
        title: Section title (will be uppercased)
        description: 1-2 sentence description in 10pt
    """
    # Title - uppercase, with number
    full_title = f"{number}. {title.upper()}"
    heading = doc.add_heading(full_title, 1)
    
    # Description under title (10pt, italic)
    if description:
        p_desc = doc.add_paragraph()
        run_desc = p_desc.add_run(description)
        run_desc.italic = True
        run_desc.font.size = Pt(10)
        run_desc.font.color.rgb = COLORS["gray"]
        p_desc.paragraph_format.space_after = Pt(12)


def _add_subsection_header(doc, title: str):
    """Add subsection header (uppercase)."""
    doc.add_heading(title.upper(), 2)


def _add_metric_table(doc, data: List[List[str]], has_header: bool = True):
    """Add a simple editable table."""
    if not data or len(data) == 0:
        return
    
    rows = len(data)
    cols = len(data[0]) if data else 0
    if cols == 0:
        return
    
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'
    
    for i, row_data in enumerate(data):
        cells = table.rows[i].cells
        for j, val in enumerate(row_data):
            cells[j].text = str(val) if val is not None else "---"
            if i == 0 and has_header:
                _set_cell_shading(cells[j], "2C3E50")  # Dark header
                for paragraph in cells[j].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = COLORS["white"]
    
    doc.add_paragraph()


def _add_recommendation_box(doc, text: str):
    """Add training recommendation with navy background."""
    _add_colored_box(doc, text, "navy")


def _add_warning_box(doc, text: str):
    """Add warning/limitation with red background."""
    _add_colored_box(doc, text, "red")


def _add_positive_box(doc, text: str):
    """Add positive/strength with green background."""
    _add_colored_box(doc, text, "green")


# =============================================================================
# SECTION BUILDERS
# =============================================================================

def _build_title_page(doc, meta: Dict):
    """Build premium title page with DarkGlass styling."""
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Main title block with dark background
    title_table = doc.add_table(rows=1, cols=1)
    title_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = title_table.rows[0].cells[0]
    _set_cell_shading(cell, "17252A")  # DarkGlass
    
    # Title text
    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(30)
    p1.paragraph_format.space_after = Pt(10)
    
    run1 = p1.add_run("BADANIA WYDOLNOŚCIOWE")
    run1.bold = True
    run1.font.size = Pt(24)
    run1.font.color.rgb = COLORS["white"]
    
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("w oparciu o Wentylację Minutową (VE)")
    run2.font.size = Pt(16)
    run2.font.color.rgb = RGBColor(189, 195, 199)  # Light gray
    
    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("i Natlenienie Mięśniowe (SmO₂)")
    run3.font.size = Pt(16)
    run3.font.color.rgb = RGBColor(189, 195, 199)
    p3.paragraph_format.space_after = Pt(30)
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Test info
    test_date = meta.get('test_date', '---')
    session_id = meta.get('session_id', '')[:8] if meta.get('session_id') else ''
    
    p_info = doc.add_paragraph()
    p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_info = p_info.add_run(f"Data testu: {test_date}")
    run_info.font.size = Pt(14)
    
    if session_id:
        p_id = doc.add_paragraph()
        p_id.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_id = p_id.add_run(f"ID sesji: {session_id}")
        run_id.font.size = Pt(11)
        run_id.font.color.rgb = COLORS["gray"]
    
    # Spacer to push footer to bottom
    for _ in range(15):
        doc.add_paragraph()
    
    # Author at bottom
    p_author = doc.add_paragraph()
    p_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_author = p_author.add_run("Opracowanie: Krzysztof Kubicz")
    run_author.bold = True
    run_author.font.size = Pt(12)
    
    # Generation date
    gen_date = datetime.now().strftime("%d.%m.%Y, %H:%M")
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_date = p_date.add_run(f"Data generowania raportu: {gen_date}")
    run_date.font.size = Pt(10)
    run_date.font.color.rgb = COLORS["gray"]
    
    doc.add_page_break()


def _build_section_toc(doc, section_titles: List[Dict[str, str]]):
    """Section: Table of Contents."""
    _add_section_header(doc, 0, "SPIS TREŚCI", 
                       "Przegląd wszystkich sekcji raportu z numerami stron.")
    
    toc_data = [["Nr", "Sekcja", "Str."]]
    for section in section_titles:
        toc_data.append([
            section.get("num", ""),
            section.get("title", ""),
            section.get("page", "")
        ])
    
    _add_metric_table(doc, toc_data)
    _add_coach_notes(doc, lines=3)
    doc.add_page_break()


def _build_section_cover(doc, data: Dict, meta: Dict):
    """Section 1: Summary Page."""
    _add_section_header(doc, 1, "BADANIA WYDOLNOŚCIOWE - RAPORT POTESTOWY",
                       "Podsumowanie kluczowych wyników testu wydolnościowego z progami i mocą krytyczną.")
    
    thresholds = data.get("thresholds", {})
    cp = data.get("cp_model", {})
    
    results = [
        ["Parametr", "Wartość", "Opis"],
        ["VT1 (Próg tlenowy)", f"{thresholds.get('vt1_watts', '---')} W", "Górna granica Z2"],
        ["VT2 (Próg beztlenowy)", f"{thresholds.get('vt2_watts', '---')} W", "Początek Z4"],
        ["Critical Power (CP)", f"{cp.get('cp_watts', '---')} W", "Moc krytyczna"],
        ["W' (Rezerwa)", f"{cp.get('w_prime_kj', '---')} kJ", "Pojemność anaerobowa"],
    ]
    _add_metric_table(doc, results)
    
    _add_coach_notes(doc)
    doc.add_page_break()


def _build_section_zones(doc, thresholds: Dict):
    """Section 2: Training Zones."""
    _add_section_header(doc, 2, "REKOMENDOWANE STREFY TRENINGOWE",
                       "Indywidualne strefy treningowe wyznaczone na podstawie progów wentylacyjnych.")
    
    vt1_str = str(thresholds.get("vt1_watts", "0"))
    vt2_str = str(thresholds.get("vt2_watts", "0"))
    
    if vt1_str.isdigit() and vt2_str.isdigit() and int(vt1_str) > 0:
        vt1, vt2 = float(vt1_str), float(vt2_str)
        
        zones = [
            ["Strefa", "Zakres [W]", "Intensywność", "Cel adaptacji"],
            ["Z1 Recovery", f"< {int(vt1 * 0.8)}", "Bardzo łatwy", "Regeneracja aktywna"],
            ["Z2 Endurance", f"{int(vt1 * 0.8)}–{int(vt1)}", "Komfortowy", "Baza tlenowa, spalanie tłuszczu"],
            ["Z3 Tempo", f"{int(vt1)}–{int(vt2)}", "Umiarkowany", "Próg, efektywność"],
            ["Z4 Threshold", f"{int(vt2)}–{int(vt2 * 1.05)}", "Ciężki", "Wytrzymałość progowa"],
            ["Z5 VO₂max", f"> {int(vt2 * 1.05)}", "Maksymalny", "Pułap tlenowy"],
        ]
        _add_metric_table(doc, zones)
        
        _add_recommendation_box(doc, "💡 ZALECENIE: 80% czasu trenuj w Z1-Z2, 20% w Z4-Z5. Unikaj 'szarej strefy' Z3.")
    else:
        _add_warning_box(doc, "⚠️ Brak danych do wyznaczenia stref (wymagane VT1 i VT2).")
    
    _add_coach_notes(doc)
    doc.add_page_break()


def _build_section_thresholds(doc, thresholds: Dict, figure_paths: Dict):
    """Section 3: VT1/VT2 Threshold Details."""
    _add_section_header(doc, 3, "SZCZEGÓŁY PROGÓW VT1 / VT2",
                       "Progi wentylacyjne to kluczowe punkty określające intensywność treningu.")
    
    data = [
        ["Próg", "Moc [W]", "HR [bpm]", "VE [L/min]", "Interpretacja"],
        ["VT1", thresholds.get('vt1_watts', '---'), thresholds.get('vt1_hr', '---'), 
         thresholds.get('vt1_ve', '---'), "Górna granica treningu bazowego"],
        ["VT2", thresholds.get('vt2_watts', '---'), thresholds.get('vt2_hr', '---'), 
         thresholds.get('vt2_ve', '---'), "Próg mleczanowy / tempo race"],
    ]
    _add_metric_table(doc, data)
    
    if "ve_profile" in figure_paths and os.path.exists(figure_paths["ve_profile"]):
        doc.add_picture(figure_paths["ve_profile"], width=Inches(5.5))
    
    _add_coach_notes(doc)
    doc.add_page_break()


def _build_section_interpretation(doc, thresholds: Dict, cp: Dict):
    """Section 4: Results Interpretation."""
    _add_section_header(doc, 4, "CO OZNACZAJĄ TE WYNIKI?",
                       "Interpretacja fizjologiczna wyników testu i ich znaczenie dla treningu.")
    
    _add_subsection_header(doc, "INTERPRETACJA VT1")
    vt1_w = thresholds.get('vt1_watts', '---')
    doc.add_paragraph(f"VT1 = {vt1_w} W — to górna granica treningu bazowego (Z1-Z2). "
                      "Poniżej tego progu dominuje metabolizm tlenowy i spalanie tłuszczów.")
    
    _add_subsection_header(doc, "INTERPRETACJA VT2")
    vt2_w = thresholds.get('vt2_watts', '---')
    doc.add_paragraph(f"VT2 = {vt2_w} W — próg mleczanowy, granica tempo/threshold. "
                      "Powyżej VT2 akumulacja mleczanu przewyższa jego utylizację.")
    
    _add_subsection_header(doc, "INTERPRETACJA CP / W'")
    cp_w = cp.get('cp_watts', '---')
    wprime = cp.get('w_prime_kj', '---')
    doc.add_paragraph(f"CP = {cp_w} W, W' = {wprime} kJ — Twoja moc krytyczna i rezerwa beztlenowa. "
                      "CP to najwyższa moc utrzymywana teoretycznie bez końca.")
    
    _add_recommendation_box(doc, "💡 TIP: Znajomość tych progów pozwala precyzyjnie targetować adaptacje treningowe.")
    
    _add_coach_notes(doc, lines=8)
    doc.add_page_break()


def _build_section_ventilation(doc, vent_data: Dict):
    """Section 5: Breathing Control."""
    _add_section_header(doc, 5, "KONTROLA ODDYCHANIA I METABOLIZMU",
                       "Analiza wzorców oddechowych i efektywności wentylacyjnej podczas testu.")
    
    if not vent_data:
        _add_warning_box(doc, "⚠️ Brak danych wentylacyjnych dla tej sesji.")
        _add_coach_notes(doc)
        doc.add_page_break()
        return
    
    metrics = vent_data.get("metrics", {})
    ve_data = [
        ["Metryka", "Wartość", "Interpretacja"],
        ["VE Peak", f"{metrics.get('ve_peak', '---')} L/min", "Szczytowa wentylacja"],
        ["VE/VCO₂ @ VT1", f"{metrics.get('ve_vco2_vt1', '---')}", "Efektywność oddychania"],
        ["Breathing Reserve", f"{metrics.get('breathing_reserve', '---')}%", "Margines wentylacyjny"],
    ]
    _add_metric_table(doc, ve_data)
    
    _add_coach_notes(doc)
    doc.add_page_break()


def _build_section_metabolic_engine(doc, metabolic_data: Dict):
    """Section 6: Metabolic Engine."""
    _add_section_header(doc, 6, "SILNIK METABOLICZNY I STRATEGIA TRENINGOWA",
                       "Analiza profilu metabolicznego i identyfikacja głównego limitatora wydajności.")
    
    if not metabolic_data:
        _add_warning_box(doc, "⚠️ Brak danych o profilu metabolicznym.")
        _add_coach_notes(doc)
        doc.add_page_break()
        return
    
    profile = metabolic_data.get("profile", {})
    limiter = profile.get('limiter', 'nieznany')
    
    doc.add_paragraph(f"Zidentyfikowany limiter wydajności: {limiter.upper()}")
    
    if limiter.lower() == "aerobic":
        _add_warning_box(doc, "⚠️ LIMITER: System tlenowy - priorytet: trening Z2 i tempo (Z3)")
    elif limiter.lower() == "glycolytic":
        _add_recommendation_box(doc, "💡 LIMITER: System glikolityczny - priorytet: interwały VO2max")
    else:
        _add_positive_box(doc, "✓ Profil zbalansowany - możesz pracować nad wszystkimi systemami")
    
    _add_coach_notes(doc)
    doc.add_page_break()


def _build_section_pdc(doc, cp: Dict, figure_paths: Dict):
    """Section 7: Power Duration Curve."""
    _add_section_header(doc, 7, "KRZYWA MOCY (PDC) I CRITICAL POWER",
                       "Model CP/W' opisuje Twoją zdolność do utrzymywania mocy w czasie.")
    
    pdc_data = [
        ["Parametr", "Wartość", "Znaczenie"],
        ["Critical Power (CP)", f"{cp.get('cp_watts', '---')} W", "Granica mocy aerobowej"],
        ["W' (W-prime)", f"{cp.get('w_prime_kj', '---')} kJ", "Rezerwa anaerobowa"],
    ]
    _add_metric_table(doc, pdc_data)
    
    if "pdc_curve" in figure_paths and os.path.exists(figure_paths["pdc_curve"]):
        doc.add_picture(figure_paths["pdc_curve"], width=Inches(5.5))
    
    _add_recommendation_box(doc, "💡 STRATEGIA: Ataki powyżej CP zużywają W'. Regeneracja W' wymaga zejścia poniżej CP.")
    
    _add_coach_notes(doc)
    doc.add_page_break()


def _build_section_smo2(doc, smo2_data: Dict, smo2_manual: Dict, figure_paths: Dict):
    """Section 8: Muscle Oxygenation."""
    _add_section_header(doc, 8, "DIAGNOSTYKA OKSYGENACJI MIĘŚNIOWEJ (SMO₂)",
                       "SmO₂ pokazuje balans między dostawą a zużyciem tlenu bezpośrednio w mięśniu.")
    
    smo2_table = [
        ["Punkt", "Moc [W]", "HR [bpm]", "Znaczenie"],
        ["SmO₂ LT1", smo2_manual.get('lt1_watts', '---'), smo2_manual.get('lt1_hr', '---'), 
         "Początek desaturacji"],
        ["SmO₂ LT2", smo2_manual.get('lt2_watts', '---'), smo2_manual.get('lt2_hr', '---'), 
         "Głęboka desaturacja"],
    ]
    _add_metric_table(doc, smo2_table)
    
    if "smo2_power" in figure_paths and os.path.exists(figure_paths["smo2_power"]):
        doc.add_picture(figure_paths["smo2_power"], width=Inches(5.5))
    
    _add_coach_notes(doc)
    doc.add_page_break()


def _build_section_cardiovascular(doc, cardio_data: Dict):
    """Section 9: Cardiovascular Diagnostic."""
    _add_section_header(doc, 9, "DIAGNOSTYKA SERCOWO-NACZYNIOWA",
                       "Analiza efektywności układu krążenia i kosztu sercowego wysiłku.")
    
    if not cardio_data:
        _add_warning_box(doc, "⚠️ Brak danych sercowo-naczyniowych.")
        _add_coach_notes(doc)
        doc.add_page_break()
        return
    
    pp = cardio_data.get("pulse_power", 0)
    ef = cardio_data.get("efficiency_factor", 0)
    drift = cardio_data.get("hr_drift_pct", 0)
    
    cardio_table = [
        ["Metryka", "Wartość", "Interpretacja"],
        ["Moc Pulsowa", f"{pp:.2f} W/bpm", "Moc generowana na jedno uderzenie serca"],
        ["Wsp. Efektywności", f"{ef:.2f} W/bpm", "Wydajność krążenia"],
        ["Dryf HR", f"{drift:.1f}%", "Stabilność tętna podczas wysiłku"],
    ]
    _add_metric_table(doc, cardio_table)
    
    if drift < 3:
        _add_positive_box(doc, "✓ DOSKONALE: Niski dryf HR świadczy o dobrej wydolności i nawodnieniu")
    elif drift > 6:
        _add_warning_box(doc, "⚠️ UWAGA: Wysoki dryf HR może wskazywać na problemy z termoregulacją")
    
    _add_coach_notes(doc)
    doc.add_page_break()


def _build_section_limiter_radar(doc, limiter_data: Dict, figure_paths: Dict):
    """Section 10: System Load Radar."""
    _add_section_header(doc, 10, "RADAR OBCIĄŻENIA SYSTEMÓW",
                       "Wizualizacja obciążenia poszczególnych systemów fizjologicznych.")
    
    if "limiters_radar" in figure_paths and os.path.exists(figure_paths["limiters_radar"]):
        doc.add_picture(figure_paths["limiters_radar"], width=Inches(5.5))
    
    doc.add_paragraph(
        "Radar pokazuje względne obciążenie systemów: Serca, Płuc, Mięśni i Mocy. "
        "System z najwyższym obciążeniem jest Twoim głównym limitatorem."
    )
    
    _add_coach_notes(doc)
    doc.add_page_break()


def _build_section_biomech(doc, biomech_data: Dict, figure_paths: Dict):
    """Section 11: Biomechanical Analysis."""
    _add_section_header(doc, 11, "ANALIZA BIOMECHANICZNA",
                       "Ocena sposobu generowania mocy: kadencja vs moment obrotowy.")
    
    if "biomech_summary" in figure_paths and os.path.exists(figure_paths["biomech_summary"]):
        doc.add_picture(figure_paths["biomech_summary"], width=Inches(5.5))
    
    if "biomech_torque_smo2" in figure_paths and os.path.exists(figure_paths["biomech_torque_smo2"]):
        doc.add_picture(figure_paths["biomech_torque_smo2"], width=Inches(5.5))
        _add_warning_box(doc, "⚠️ Spadek SmO₂ przy wysokim momencie może świadczyć o okluzji mechanicznej")
    
    _add_coach_notes(doc)
    doc.add_page_break()


def _build_section_drift(doc, kpi: Dict, figure_paths: Dict):
    """Section 12: Physiological Drift."""
    _add_section_header(doc, 12, "DRYF FIZJOLOGICZNY",
                       "Analiza zmian parametrów w czasie trwania testu.")
    
    if "drift_heatmap_hr" in figure_paths and os.path.exists(figure_paths["drift_heatmap_hr"]):
        _add_subsection_header(doc, "MAPA DRYFU HR VS POWER")
        doc.add_picture(figure_paths["drift_heatmap_hr"], width=Inches(5.5))
    
    if "drift_heatmap_smo2" in figure_paths and os.path.exists(figure_paths["drift_heatmap_smo2"]):
        _add_subsection_header(doc, "MAPA OKSYDACJI SMO₂ VS POWER")
        doc.add_picture(figure_paths["drift_heatmap_smo2"], width=Inches(5.5))
    
    _add_coach_notes(doc)
    doc.add_page_break()


def _build_section_kpi(doc, kpi: Dict):
    """Section 13: Key Performance Indicators."""
    _add_section_header(doc, 13, "KLUCZOWE WSKAŹNIKI WYDAJNOŚCI (KPI)",
                       "Zbiorcze metryki charakteryzujące Twoją wydolność i efektywność.")
    
    kpi_table = [
        ["Metryka", "Wartość", "Co oznacza"],
        ["Efficiency Factor", str(kpi.get("ef", "---")), "Moc na uderzenie serca"],
        ["Pa:Hr Decoupling", f"{kpi.get('pa_hr', '---')}%", "Stabilność układu krążenia"],
        ["SmO₂ Drift", f"{kpi.get('smo2_drift', '---')}%", "Zmęczenie lokalne mięśni"],
        ["VO₂max Estimate", f"{kpi.get('vo2max_est', '---')} ml/kg/min", "Szacowany pułap tlenowy"],
    ]
    _add_metric_table(doc, kpi_table)
    
    _add_coach_notes(doc)
    doc.add_page_break()


def _build_section_thermal(doc, thermo_data: Dict, figure_paths: Dict):
    """Section 14: Thermoregulation."""
    _add_section_header(doc, 14, "ANALIZA TERMOREGULACJI",
                       "Wpływ temperatury ciała na wydajność i efektywność wysiłku.")
    
    if "thermal_hsi" in figure_paths and os.path.exists(figure_paths["thermal_hsi"]):
        _add_subsection_header(doc, "HEAT STRAIN INDEX (HSI)")
        doc.add_picture(figure_paths["thermal_hsi"], width=Inches(5.5))
    
    if "thermal_efficiency" in figure_paths and os.path.exists(figure_paths["thermal_efficiency"]):
        _add_subsection_header(doc, "EFEKTYWNOŚĆ VS TEMPERATURA")
        doc.add_picture(figure_paths["thermal_efficiency"], width=Inches(5.5))
    
    metrics = thermo_data.get("metrics", {})
    if metrics:
        thermal_table = [
            ["Metryka", "Wartość"],
            ["Max Temp", f"{metrics.get('max_core_temp', '---')} °C"],
            ["Peak HSI", f"{metrics.get('peak_hsi', '---')}"],
        ]
        _add_metric_table(doc, thermal_table)
    
    _add_coach_notes(doc)
    doc.add_page_break()


def _build_section_executive_summary(doc, exec_data: Dict, meta: Dict):
    """Section 15: Executive Summary with Training Decisions."""
    _add_section_header(doc, 15, "PODSUMOWANIE FIZJOLOGICZNE",
                       "Synteza wyników i decyzje treningowe oparte na danych z testu.")
    
    limiter = exec_data.get("limiter", {})
    training_cards = exec_data.get("training_cards", [])
    confidence_panel = exec_data.get("confidence_panel", {})
    
    # Primary Limiter
    _add_subsection_header(doc, "DOMINUJĄCY LIMITER")
    limiter_name = limiter.get("name", "NIEZNANY")
    verdict = limiter.get("verdict", "")
    
    p = doc.add_paragraph()
    p.add_run(f"{limiter.get('icon', '⚖️')} {limiter_name}").bold = True
    
    if verdict:
        doc.add_paragraph(verdict)
    
    for line in limiter.get("interpretation", [])[:3]:
        doc.add_paragraph(f"• {line}")
    
    # Confidence
    _add_subsection_header(doc, "PEWNOŚĆ TESTU")
    overall = confidence_panel.get("overall_score", 0)
    label = confidence_panel.get("label", "---")
    
    if overall >= 80:
        _add_positive_box(doc, f"✓ Pewność testu: {overall}% ({label}) - WYSOKA WIARYGODNOŚĆ")
    elif overall >= 60:
        _add_recommendation_box(doc, f"Pewność testu: {overall}% ({label}) - AKCEPTOWALNA WIARYGODNOŚĆ")
    else:
        _add_warning_box(doc, f"⚠️ Pewność testu: {overall}% ({label}) - NISKA WIARYGODNOŚĆ")
    
    # Training Decisions (navy boxes)
    _add_subsection_header(doc, "DECYZJE TRENINGOWE")
    
    for i, card in enumerate(training_cards[:3], 1):
        strategy = card.get('strategy_name', '---')
        power_range = card.get('power_range', '---')
        volume = card.get('volume', '---')
        goal = card.get('adaptation_goal', '---')
        response = card.get('expected_response', '---')
        
        _add_recommendation_box(doc, f"STRATEGIA {i}: {strategy}")
        doc.add_paragraph(f"Moc: {power_range} | Objętość: {volume}")
        doc.add_paragraph(f"Cel: {goal}")
        doc.add_paragraph(f"Spodziewany efekt: {response}")
        doc.add_paragraph()
    
    _add_coach_notes(doc, lines=10)
    doc.add_page_break()


def _build_section_limitations(doc, conf: Dict):
    """Section 16: Limitations (with red backgrounds)."""
    _add_section_header(doc, 16, "OGRANICZENIA INTERPRETACJI",
                       "Ważne zastrzeżenia dotyczące interpretacji wyników tego raportu.")
    
    limitations = [
        "To nie jest badanie medyczne. Wyniki są szacunkami algorytmicznymi.",
        "Dokładność zależy od jakości danych. Kalibracja czujników ma znaczenie.",
        "Progi są przybliżeniami. VT1/VT2 mogą się różnić od pomiarów laboratoryjnych.",
        "Wyniki są jednorazowe. Powtarzaj testy co 6-8 tygodni.",
    ]
    
    for lim in limitations:
        _add_warning_box(doc, f"⚠️ {lim}")
    
    # Warnings from confidence
    warnings = conf.get("warnings", [])
    if warnings:
        _add_subsection_header(doc, "OSTRZEŻENIA Z ANALIZY")
        for w in warnings:
            _add_warning_box(doc, f"⚠️ {w}")
    
    _add_coach_notes(doc, lines=5)


def _build_contact_footer(doc):
    """Add contact info at the very end."""
    doc.add_paragraph()
    doc.add_paragraph()
    
    p_sep = doc.add_paragraph()
    p_sep.add_run("─" * 60)
    p_sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p_contact = doc.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_name = p_contact.add_run("Krzysztof Kubicz")
    run_name.bold = True
    run_name.font.size = Pt(12)
    
    p_email = doc.add_paragraph()
    p_email.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_email = p_email.add_run("kubiczk@icloud.com")
    run_email.font.size = Pt(11)
    run_email.font.color.rgb = COLORS["navy"]


# =============================================================================
# MAIN BUILDER FUNCTION
# =============================================================================

def build_ramp_docx(
    report_data: Dict[str, Any],
    figure_paths: Optional[Dict[str, str]] = None,
    output_path: Optional[str] = None
) -> Optional[str]:
    """Build complete Ramp Test DOCX report with premium styling.
    
    Features:
    - Premium title page with DarkGlass style
    - Numbered sections with descriptions
    - Colored boxes: Navy (recommendations), Red (warnings), Green (positives)
    - COACH NOTES after each section
    - Contact info at the end
    """
    if not HAS_DOCX:
        logger.error("python-docx library not installed. Cannot generate DOCX.")
        return None

    try:
        figure_paths = figure_paths or {}
        doc = Document()
        
        # Map data using SAME function as PDF
        data = map_ramp_json_to_pdf_data(report_data)
        
        meta = data['metadata']
        thresholds = data['thresholds']
        cp = data['cp_model']
        smo2 = data['smo2']
        smo2_manual = data['smo2_manual']
        conf = data['confidence']
        kpi = data['kpi']
        exec_summary = data.get('executive_summary', {})
        vent_data = data.get('vent_advanced', {})
        metabolic_data = data.get('metabolic_strategy', {})
        cardio_data = data.get('cardio_advanced', {})
        limiter_data = data.get('limiter_analysis', {})
        biomech_data = data.get('biomech_occlusion', {})
        thermo_data = data.get('thermo_analysis', {})
        
        # Section titles for TOC
        section_titles = [
            {"num": "1", "title": "Badania Wydolnościowe - Raport Potestowy", "page": "3"},
            {"num": "2", "title": "Rekomendowane Strefy Treningowe", "page": "4"},
            {"num": "3", "title": "Szczegóły Progów VT1 / VT2", "page": "5"},
            {"num": "4", "title": "Co Oznaczają Te Wyniki?", "page": "6"},
            {"num": "5", "title": "Kontrola Oddychania i Metabolizmu", "page": "7"},
            {"num": "6", "title": "Silnik Metaboliczny i Strategia", "page": "8"},
            {"num": "7", "title": "Krzywa Mocy (PDC) i Critical Power", "page": "9"},
            {"num": "8", "title": "Diagnostyka Oksygenacji Mięśniowej", "page": "10"},
            {"num": "9", "title": "Diagnostyka Sercowo-Naczyniowa", "page": "11"},
            {"num": "10", "title": "Radar Obciążenia Systemów", "page": "12"},
            {"num": "11", "title": "Analiza Biomechaniczna", "page": "13"},
            {"num": "12", "title": "Dryf Fizjologiczny", "page": "14"},
            {"num": "13", "title": "Kluczowe Wskaźniki Wydajności (KPI)", "page": "15"},
            {"num": "14", "title": "Analiza Termoregulacji", "page": "16"},
            {"num": "15", "title": "Podsumowanie Fizjologiczne", "page": "17"},
            {"num": "16", "title": "Ograniczenia Interpretacji", "page": "18"},
        ]
        
        # =======================================================================
        # BUILD ALL SECTIONS
        # =======================================================================
        
        # Title Page
        _build_title_page(doc, meta)
        
        # Table of Contents
        _build_section_toc(doc, section_titles)
        
        # All content sections
        _build_section_cover(doc, data, meta)
        _build_section_zones(doc, thresholds)
        _build_section_thresholds(doc, thresholds, figure_paths)
        _build_section_interpretation(doc, thresholds, cp)
        _build_section_ventilation(doc, vent_data)
        _build_section_metabolic_engine(doc, metabolic_data)
        _build_section_pdc(doc, cp, figure_paths)
        _build_section_smo2(doc, smo2, smo2_manual, figure_paths)
        _build_section_cardiovascular(doc, cardio_data)
        _build_section_limiter_radar(doc, limiter_data, figure_paths)
        _build_section_biomech(doc, biomech_data, figure_paths)
        _build_section_drift(doc, kpi, figure_paths)
        _build_section_kpi(doc, kpi)
        _build_section_thermal(doc, thermo_data, figure_paths)
        _build_section_executive_summary(doc, exec_summary, meta)
        _build_section_limitations(doc, conf)
        
        # Contact footer
        _build_contact_footer(doc)
        
        # =======================================================================
        # SAVE
        # =======================================================================
        if output_path:
            Path(output_path).parent.mkdir(parents=True, exist_ok=True)
            doc.save(output_path)
            logger.info(f"DOCX saved to: {output_path}")
            return output_path
            
    except Exception as e:
        logger.error(f"DOCX Generation failed: {e}")
        import traceback
        logger.error(traceback.format_exc())
        return None
    
    return None
